from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from werkzeug.utils import secure_filename
import uuid
import os
import traceback
import json
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

from config import Config
from models.database import TemplateDB
from models.persistent_database import get_persistent_template_db, get_persistent_cai_contact_db
from utils.advanced_template_analyzer import analyze_template
from utils.advanced_resume_parser import parse_resume
from utils.doc_converter import convert_doc_to_docx, needs_conversion
from utils.cai_contact_extractor import extract_cai_contact_from_template
from utils.azure_storage import get_storage_manager

# Import routes
from routes.onlyoffice_routes import onlyoffice_bp

# Import Azure Monitor tracker (supports Live Metrics)
try:
    from utils.azure_monitor_tracker import AzureMonitorTracker
    insights_tracker = AzureMonitorTracker()
    insights_available = True
except Exception as e:
    print(f"[WARN] Azure Monitor not available: {e}")
    insights_tracker = None
    insights_available = False

# Try to import enhanced formatter, fallback to standard if not available
try:
    from utils.enhanced_formatter_integration import format_resume_intelligent
    print("[OK] Enhanced intelligent formatter loaded")
except ImportError:
    from utils.intelligent_formatter import format_resume_intelligent
    print("[WARN] Using standard formatter (enhanced version not available)")

# Configure Flask to use React build output
# Check if running in Docker container
if os.path.exists('/app/frontend'):
    # Docker container path
    frontend_dir = '/app/frontend'
else:
    # Local development path
    frontend_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'frontend', 'build')

static_dir = os.path.join(frontend_dir, 'static')

app = Flask(__name__,
           static_folder=static_dir,
           template_folder=frontend_dir)
app.config.from_object(Config)
Config.init_app(app)

# Initialize Azure Monitor
if insights_available and insights_tracker:
    insights_tracker.init_app(app)
    print("[OK] Azure Monitor tracking enabled")
else:
    print("[WARN] Azure Monitor not initialized")

# Enable CORS for React frontend and OnlyOffice Document Server
# CRITICAL: Must allow Docker container IPs for OnlyOffice callbacks
import socket

# Get local IP for CORS
try:
    s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
    s.connect(("8.8.8.8", 80))
    local_ip = s.getsockname()[0]
    s.close()
except:
    local_ip = "192.168.0.104"

CORS(
    app,
    resources={r"/api/*": {
        "origins": [
            "https://resume-formatter.reddesert-f6724e64.centralus.azurecontainerapps.io",  # Frontend
            "https://onlyoffice.reddesert-f6724e64.centralus.azurecontainerapps.io",  # OnlyOffice
            "http://localhost:3000",
            "http://localhost:3001",
            "http://127.0.0.1:3000",
            "http://127.0.0.1:3001",
            "http://localhost:8080",  # OnlyOffice Document Server
            "http://host.docker.internal",  # Docker Desktop
            "http://host.docker.internal:5000",
            f"http://{local_ip}",  # Local network IP
            f"http://{local_ip}:5000",
            f"http://{local_ip}:3000",
            "http://192.168.65.254",  # Docker internal gateway
            "*"  # Allow all for OnlyOffice (it uses various IPs)
        ]
    }},
    supports_credentials=True,
    allow_headers=["Content-Type", "Authorization"],
    methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"]
)

# Register blueprints
app.register_blueprint(onlyoffice_bp)

# Initialize databases
db = TemplateDB()  # Keep for backward compatibility
persistent_db = get_persistent_template_db()  # New persistent storage
cai_db = get_persistent_cai_contact_db()  # Persistent CAI contacts
storage_manager = get_storage_manager()  # Azure storage manager

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in Config.ALLOWED_EXTENSIONS


# ===== Persistent CAI Contact storage used by both backend and formatter =====
def _cai_store_path():
    home = os.path.expanduser("~")
    return os.path.join(home, ".resume_formatter_cai_contact.json")


@app.route('/api/cai-contact', methods=['GET'])
def get_cai_contact():
    """Return stored CAI contact from persistent storage. If none, return empty fields."""
    try:
        # Try persistent storage first
        contact_data = cai_db.get_contact()
        print(f"✅ CAI contact retrieved from persistent storage: {contact_data.get('name', 'No name')}")
        return jsonify({"success": True, "contact": contact_data})
    except Exception as e:
        print(f"⚠️ Error reading CAI contact from persistent storage: {e}")
        # Fallback to local storage
        try:
            path = _cai_store_path()
            data = {"name": "", "phone": "", "email": ""}
            if os.path.exists(path):
                with open(path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                print(f"✅ CAI contact retrieved from local fallback: {data.get('name', 'No name')}")
            return jsonify({"success": True, "contact": data})
        except Exception as e2:
            print(f"❌ Error reading CAI contact from fallback: {e2}")
            return jsonify({"success": True, "contact": {"name": "", "phone": "", "email": ""}})


@app.route('/api/cai-contact', methods=['POST'])
def save_cai_contact():
    """Persist CAI contact to persistent storage. Overwrites stored values. Body: JSON {name, phone, email}."""
    try:
        payload = request.get_json(silent=True) or {}
        data = {
            "name": str(payload.get("name", "")).strip(),
            "phone": str(payload.get("phone", "")).strip(),
            "email": str(payload.get("email", "")).strip(),
        }
        
        # Save to persistent storage
        success = cai_db.save_contact(data)
        
        if success:
            print(f"✅ CAI contact saved to persistent storage: {data.get('name', 'No name')}")
            
            # Also save to local fallback for backward compatibility
            try:
                path = _cai_store_path()
                with open(path, 'w', encoding='utf-8') as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)
                print(f"✅ CAI contact also saved to local fallback")
            except Exception as e:
                print(f"⚠️ Failed to save to local fallback: {e}")
            
            return jsonify({"success": True, "contact": data})
        else:
            print(f"❌ Failed to save CAI contact to persistent storage")
            return jsonify({"success": False, "message": "Failed to save to persistent storage"}), 500
            
    except Exception as e:
        print(f"❌ Error saving CAI contact: {e}")
        traceback.print_exc()
        return jsonify({"success": False, "message": str(e)}), 500


@app.route('/api/cai-contact', methods=['DELETE'])
def delete_cai_contact():
    """Delete CAI contact from persistent storage."""
    try:
        # Clear the contact data
        empty_data = {"name": "", "phone": "", "email": ""}
        
        # Save empty contact to persistent storage (effectively deleting it)
        success = cai_db.save_contact(empty_data)
        
        if success:
            print(f"✅ CAI contact deleted from persistent storage")
            
            # Also clear local fallback
            try:
                path = _cai_store_path()
                if os.path.exists(path):
                    os.remove(path)
                print(f"✅ CAI contact also deleted from local fallback")
            except Exception as e:
                print(f"⚠️ Failed to delete from local fallback: {e}")
            
            return jsonify({"success": True, "message": "Contact deleted"})
        else:
            print(f"❌ Failed to delete CAI contact from persistent storage")
            return jsonify({"success": False, "message": "Failed to delete from persistent storage"}), 500
            
    except Exception as e:
        print(f"❌ Error deleting CAI contact: {e}")
        traceback.print_exc()
        return jsonify({"success": False, "message": str(e)}), 500


@app.route('/api/templates/<template_id>/cai-contacts', methods=['GET'])
def get_template_cai_contacts(template_id):
    """Get CAI contacts associated with a specific template"""
    try:
        # Get template with CAI contact
        template = persistent_db.get_template(template_id)
        if not template:
            template = db.get_template(template_id)

        if template and template.get('cai_contact'):
            # Return the CAI contact extracted from the template
            cai_contact = template['cai_contact']
            print(f"✅ CAI Contact from template {template_id}: {cai_contact.get('name', 'N/A')} ({cai_contact.get('state', 'N/A')})")

            return jsonify({
                "success": True,
                "contacts": [{"id": 1, **cai_contact}],
                "contact_ids": [1],
                "template_name": template.get('name', 'Unknown')
            })
        else:
            # No CAI contact in template, return empty
            print(f"⚠️ No CAI contact found in template {template_id}")
            return jsonify({
                "success": True,
                "contacts": [],
                "contact_ids": [],
                "template_name": template.get('name', 'Unknown') if template else 'Unknown'
            })
    except Exception as e:
        print(f"❌ Error getting template CAI contacts: {e}")
        return jsonify({"success": False, "message": str(e)}), 500

@app.route('/api/templates/<template_id>/cai-contacts', methods=['POST'])
def save_template_cai_contacts(template_id):
    """Save CAI contacts for a specific template"""
    try:
        # For now, this is a placeholder since we use a single global CAI contact
        # In the future, this could be expanded to support template-specific contacts
        data = request.get_json()
        contact_ids = data.get('contact_ids', [])
        
        print(f"📝 Template {template_id} CAI contact mapping saved: {contact_ids}")
        return jsonify({"success": True, "message": "Template contact mapping saved"})
    except Exception as e:
        print(f"❌ Error saving template CAI contacts: {e}")
        return jsonify({"success": False, "message": str(e)}), 500

# ===== Plural endpoint compatibility =====
@app.route('/api/cai-contacts', methods=['GET', 'POST'])
def cai_contacts_plural():
    """Handle both GET and POST for plural endpoint - compatibility with frontend"""
    try:
        if request.method == 'GET':
            # Inline implementation to avoid function call issues
            try:
                contact_data = cai_db.get_contact()
                return jsonify({"success": True, "contact": contact_data})
            except Exception as e:
                return jsonify({"success": True, "contact": {"name": "", "phone": "", "email": ""}})
        else:  # POST
            # Inline implementation to avoid function call issues
            try:
                payload = request.get_json(silent=True) or {}
                data = {
                    "name": str(payload.get("name", "")).strip(),
                    "phone": str(payload.get("phone", "")).strip(),
                    "email": str(payload.get("email", "")).strip(),
                }
                success = cai_db.save_contact(data)
                if success:
                    return jsonify({"success": True, "contact": data})
                else:
                    return jsonify({"success": False, "message": "Failed to save"}), 500
            except Exception as e:
                return jsonify({"success": False, "message": str(e)}), 500
    except Exception as e:
        return jsonify({"error": f"Route error: {str(e)}"}), 500

@app.route('/api/health', methods=['GET'])
def health():
    return jsonify({'status': 'ok'})

# Test routes removed - using multiple decorators approach instead

@app.route('/api/storage-status', methods=['GET'])
def storage_status():
    """Get Azure Storage connection status"""
    try:
        status = storage_manager.test_connection()
        return jsonify({
            'success': True,
            'storage_status': status
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'storage_status': {
                'connected': False,
                'storage_type': 'unknown',
                'message': f'Error checking storage: {str(e)}',
                'containers': []
            }
        })

@app.route('/api/templates', methods=['GET'])
def get_templates():
    """Get all templates from persistent storage"""
    try:
        # PRIMARY: Use Azure persistent storage
        templates = persistent_db.get_all_templates()
        if templates and len(templates) > 0:
            print(f"✅ Retrieved {len(templates)} templates from Azure persistent storage")
            for template in templates:
                print(f"  📋 Template: {template['id']} - {template['name']}")
            return jsonify({'success': True, 'templates': templates})

        # FALLBACK: Try local database only if Azure is empty or unavailable
        templates = db.get_all_templates()
        if templates and len(templates) > 0:
            print(f"⚠️  Retrieved {len(templates)} templates from local fallback database")
            for template in templates:
                print(f"  📋 Template: {template['id']} - {template['name']}")
            return jsonify({'success': True, 'templates': templates})

        print(f"📭 No templates found in any database")
        return jsonify({'success': True, 'templates': []})
    except Exception as e:
        print(f"❌ Error getting templates: {e}")
        return jsonify({'success': True, 'templates': []})

@app.route('/api/templates', methods=['POST'])
def upload_template():
    """Upload and analyze template"""
    try:
        if 'template_file' not in request.files or 'template_name' not in request.form:
            return jsonify({'success': False, 'message': 'Missing file or name'}), 400
        
        file = request.files['template_file']
        name = request.form['template_name'].strip()
        
        if file.filename == '' or not name:
            return jsonify({'success': False, 'message': 'Invalid input'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'success': False, 'message': 'Invalid file type'}), 400
        
        # Save file
        filename = secure_filename(file.filename)
        file_type = filename.rsplit('.', 1)[1].lower()
        template_id = str(uuid.uuid4())
        saved_filename = f"{template_id}_{filename}"
        file_path = os.path.join(Config.TEMPLATE_FOLDER, saved_filename)
        file.save(file_path)
        
        print(f"\n{'='*70}")
        print(f"📤 UPLOADING TEMPLATE: {name}")
        print(f"📄 Original file type: {file_type}")
        print(f"{'='*70}\n")
        
        # Handle .doc files - try conversion or accept as-is
        final_file_path = file_path
        final_file_type = file_type

        if needs_conversion(filename):
            print(f"🔄 Attempting to convert .doc to .docx...")
            converted_path = convert_doc_to_docx(file_path)

            if converted_path and os.path.exists(converted_path):
                print(f"✅ Successfully converted to .docx")
                # Update file info to point to converted file
                final_file_path = converted_path
                final_file_type = 'docx'

                # Update saved filename to reflect the conversion
                converted_filename = os.path.basename(converted_path)
                saved_filename = converted_filename

                # Remove original .doc file to save space
                try:
                    os.remove(file_path)
                    print(f"🗑️  Removed original .doc file")
                except:
                    pass
            else:
                print(f"❌ .doc conversion not available - LibreOffice/Pandoc not installed")
                # Clean up the uploaded file
                try:
                    os.remove(file_path)
                except:
                    pass
                return jsonify({
                    'success': False,
                    'message': 'Cannot process .doc files. Please save as .docx format in Microsoft Word (File → Save As → Word Document (.docx)) and try again.'
                }), 400
        
        # Analyze template with advanced analyzer
        format_data = analyze_template(final_file_path)

        # Extract CAI contact information from template
        cai_contact = extract_cai_contact_from_template(final_file_path)
        if cai_contact:
            print(f"✅ CAI Contact detected in template:")
            print(f"   Name: {cai_contact.get('name', 'N/A')}")
            print(f"   State: {cai_contact.get('state', 'N/A')}")

        # Save to persistent storage
        persistent_success = persistent_db.add_template(template_id, name, saved_filename, final_file_type, format_data, cai_contact)
        
        if persistent_success:
            # Upload file to persistent storage
            file_upload_success = persistent_db.upload_template_file(template_id, final_file_path, saved_filename)
            
            if file_upload_success:
                print(f"✅ Template '{name}' saved to persistent storage")
                
                # Also save to local database for backward compatibility
                try:
                    db.add_template(template_id, name, saved_filename, final_file_type, format_data)
                    print(f"✅ Template '{name}' also saved to local database")
                except Exception as e:
                    print(f"⚠️ Failed to save to local database: {e}")
            else:
                print(f"❌ Failed to upload template file to persistent storage")
                return jsonify({'success': False, 'message': 'Failed to upload template file'}), 500
        else:
            print(f"❌ Failed to save template metadata to persistent storage")
            # Fallback to local database
            try:
                db.add_template(template_id, name, saved_filename, final_file_type, format_data)
                print(f"✅ Template '{name}' saved to local database (fallback)")
            except Exception as e:
                print(f"❌ Failed to save to local database: {e}")
                return jsonify({'success': False, 'message': 'Failed to save template'}), 500
        
        return jsonify({
            'success': True,
            'id': template_id,
            'name': name,
            'message': 'Template uploaded and analyzed successfully'
        })
    
    except Exception as e:
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/format', methods=['POST'])
def format_resumes():
    """Format resumes using selected template"""
    try:
        if 'template_id' not in request.form or 'resume_files' not in request.files:
            return jsonify({'success': False, 'message': 'Missing template or files'}), 400
        
        template_id = request.form['template_id']
        
        # Try to get template from persistent storage first
        template = persistent_db.get_template(template_id)
        
        if not template:
            # Fallback to local database
            template = db.get_template(template_id)
            print(f"✅ Template retrieved from local fallback: {template_id}")
        else:
            print(f"✅ Template retrieved from persistent storage: {template_id}")
        
        if not template:
            return jsonify({'success': False, 'message': 'Template not found'}), 404
        
        # Download template file from persistent storage if needed
        template_filename = template['filename']
        local_template_path = os.path.join(Config.TEMPLATE_FOLDER, template_filename)
        
        if not os.path.exists(local_template_path):
            print(f"📥 Downloading template file from persistent storage...")
            download_success = persistent_db.download_template_file(template_id, template_filename, local_template_path)
            
            if not download_success:
                print(f"❌ Failed to download template file from persistent storage")
                return jsonify({'success': False, 'message': 'Template file not available'}), 404
            else:
                print(f"✅ Template file downloaded successfully")
        
        files = request.files.getlist('resume_files')
        formatted_files = []
        
        # Get template analysis
        template_analysis = template['format_data']
        
        # Ensure paths are set
        template_file_path = os.path.join(Config.TEMPLATE_FOLDER, template['filename'])
        template_analysis['template_path'] = template_file_path
        template_analysis['template_type'] = template['file_type']
        
        print(f"\n{'='*70}")
        print(f"🎯 FORMATTING SESSION")
        print(f"{'='*70}")
        print(f"📋 Template: {template['name']}")
        print(f"📁 Template Path: {template_file_path}")
        print(f"✓ Template Exists: {os.path.exists(template_file_path)}")
        print(f"📊 Resumes to Process: {len(files)}")
        print(f"{'='*70}\n")
        
        # Use ThreadPoolExecutor for parallel processing
        from concurrent.futures import ThreadPoolExecutor, as_completed
        import time
        
        start_time = time.time()
        
        # Extract CAI contact data BEFORE threading (request context not available in threads)
        cai_contact_data = None
        cai_contacts_data = None
        edit_cai_contact = False
        
        # Check for multiple contacts first (new format)
        if 'cai_contacts' in request.form:
            try:
                import json
                cai_contacts_data = json.loads(request.form['cai_contacts'])
                edit_cai_contact = request.form.get('edit_cai_contact') == 'true'
                print(f"  ✏️  CAI Contacts (multiple) edit enabled: {len(cai_contacts_data)} contact(s)")
            except Exception as e:
                print(f"  ⚠️  Error parsing CAI contacts data: {e}")
        # Backward compatibility: single contact
        elif 'cai_contact' in request.form:
            try:
                import json
                cai_contact_data = json.loads(request.form['cai_contact'])
                edit_cai_contact = request.form.get('edit_cai_contact') == 'true'
                print(f"  ✏️  CAI Contact (single) edit enabled: {cai_contact_data}")
            except Exception as e:
                print(f"  ⚠️  Error parsing CAI contact data: {e}")

        # Extract skills data
        skills_data = None
        if 'skills' in request.form:
            try:
                import json
                skills_data = json.loads(request.form['skills'])
                print(f"  📊 Skills provided: {len(skills_data)} skill(s)")
                for skill in skills_data[:5]:  # Show first 5
                    print(f"     - {skill['name']}: Level {skill['level']}")
                if len(skills_data) > 5:
                    print(f"     ... and {len(skills_data) - 5} more")
            except Exception as e:
                print(f"  ⚠️  Error parsing skills data: {e}")

        def process_single_resume(file, idx, total, cai_data, cai_contacts, edit_cai, skills):
            """Process a single resume file"""
            if file.filename == '' or not allowed_file(file.filename):
                return None
            
            # Save resume
            filename = secure_filename(file.filename)
            file_type = filename.rsplit('.', 1)[1].lower()
            resume_id = str(uuid.uuid4())
            saved_filename = f"{resume_id}_{filename}"
            file_path = os.path.join(Config.RESUME_FOLDER, saved_filename)
            file.save(file_path)
            
            print(f"\n{'─'*70}")
            print(f"📄 Processing Resume {idx}/{total}: {filename}")
            print(f"📄 Original file type: {file_type}")
            print(f"{'─'*70}")
            
            # Convert .doc to .docx if needed
            final_file_path = file_path
            final_file_type = file_type

            if needs_conversion(filename):
                print(f"🔄 Converting resume .doc to .docx...")
                converted_path = convert_doc_to_docx(file_path)

                if converted_path and os.path.exists(converted_path):
                    print(f"✅ Successfully converted resume to .docx")
                    final_file_path = converted_path
                    final_file_type = 'docx'

                    # Remove original .doc file to save space
                    try:
                        os.remove(file_path)
                        print(f"🗑️  Removed original .doc file")
                    except:
                        pass
                else:
                    print(f"⚠️  Resume .doc conversion not available, treating as .docx")
                    # Try renaming to .docx
                    try:
                        docx_path = file_path.replace('.doc', '.docx')
                        os.rename(file_path, docx_path)
                        final_file_path = docx_path
                        final_file_type = 'docx'
                        print(f"✅ Renamed resume .doc to .docx")
                    except Exception as e:
                        print(f"⚠️  Could not rename resume, continuing anyway: {e}")
                        pass
            
            # Parse resume with advanced parser (with timing)
            parse_start = time.time()
            resume_data = parse_resume(final_file_path, final_file_type)
            parse_time = time.time() - parse_start
            print(f"  ⏱️  Parsing took: {parse_time:.2f}s")
            
            # Add CAI contact data if provided (multiple contacts preferred)
            if cai_contacts:
                resume_data['cai_contacts'] = cai_contacts
                resume_data['edit_cai_contact'] = edit_cai
            elif cai_data:
                resume_data['cai_contact'] = cai_data
                resume_data['edit_cai_contact'] = edit_cai

            # Add skills data if provided
            if skills:
                resume_data['selected_skills'] = skills
                print(f"  📊 Added {len(skills)} skills to resume data")

            if resume_data:
                # Format resume with intelligent formatter
                # Create DOCX only (NO PDF for speed!)
                docx_filename = f"formatted_{resume_id}.docx"
                docx_path = os.path.join(Config.OUTPUT_FOLDER, docx_filename)
                
                format_start = time.time()
                if format_resume_intelligent(resume_data, template_analysis, docx_path):
                    format_time = time.time() - format_start
                    print(f"  ⏱️  Formatting took: {format_time:.2f}s")
                    # Check if DOCX was created
                    if os.path.exists(docx_path):
                        # NO PDF CONVERSION FOR SPEED!
                        result = {
                            'filename': docx_filename,
                            'original': filename,
                            'name': resume_data['name'],
                            'template_name': template.get('name', 'resume')
                        }
                        print(f"✅ Successfully formatted: {filename} → {docx_filename}\n")
                        
                        # Cleanup
                        try:
                            os.remove(file_path)
                        except:
                            pass
                        
                        return result
                    else:
                        print(f"⚠️  Formatting completed but output file not found\n")
                else:
                    print(f"❌ Failed to format: {filename}\n")
            else:
                print(f"❌ Failed to parse resume: {filename}\n")
            
            # Cleanup on failure
            try:
                os.remove(file_path)
            except:
                pass
            
            return None
        
        # Process all resumes in parallel for speed
        from concurrent.futures import ThreadPoolExecutor, as_completed
        
        with ThreadPoolExecutor(max_workers=min(4, len(files))) as executor:
            # Submit all tasks with CAI contact data and skills
            future_to_file = {
                executor.submit(process_single_resume, file, idx, len(files), cai_contact_data, cai_contacts_data, edit_cai_contact, skills_data): file
                for idx, file in enumerate(files, 1)
            }
            
            # Collect results as they complete
            for future in as_completed(future_to_file):
                result = future.result()
                if result:
                    formatted_files.append(result)
        
        elapsed_time = time.time() - start_time
        print(f"{'='*70}")
        print(f"✅ FORMATTING COMPLETE: {len(formatted_files)}/{len(files)} successful")
        print(f"⏱️  Total Time: {elapsed_time:.2f} seconds ({elapsed_time/len(files):.2f}s per resume)")
        print(f"{'='*70}\n")

        # Track successful output generation in Application Insights
        if insights_available and insights_tracker:
            try:
                # Get user info from request (if using Azure AD authentication)
                user_id = request.headers.get('X-User-Id', 'anonymous')
                user_email = request.headers.get('X-User-Email', 'anonymous@example.com')

                insights_tracker.track_output_generated(
                    user_id=user_id,
                    template_id=template_id,
                    template_name=template.get('name', 'Unknown'),
                    input_count=len(files),
                    output_count=len(formatted_files),
                    processing_time_ms=int(elapsed_time * 1000),
                    success=True
                )
                print(f"[INSIGHTS] Tracked output generation: {len(formatted_files)} outputs")
            except Exception as track_error:
                print(f"[WARN] Failed to track event: {track_error}")

        return jsonify({
            'success': True,
            'files': formatted_files,
            'message': f'Formatted {len(formatted_files)} resume(s)'
        })
    
    except Exception as e:
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/download/<filename>')
def download_file(filename):
    """Download formatted resume with proper filename and auto-cleanup"""
    file_path = os.path.join(Config.OUTPUT_FOLDER, filename)
    
    if not os.path.exists(file_path):
        return jsonify({'error': 'File not found'}), 404
    
    # Try to get a better filename from request args
    candidate_name = request.args.get('name', '')
    template_name = request.args.get('template', 'resume')
    
    # Clean names for filename
    if candidate_name:
        # Remove special characters and spaces
        clean_name = ''.join(c if c.isalnum() or c in (' ', '-', '_') else '' for c in candidate_name)
        clean_name = clean_name.replace(' ', '_')
        clean_template = ''.join(c if c.isalnum() or c in (' ', '-', '_') else '' for c in template_name)
        clean_template = clean_template.replace(' ', '_')
        download_name = f"{clean_name}_{clean_template}.docx"
    else:
        download_name = filename
    
    # Send file
    response = send_from_directory(
        Config.OUTPUT_FOLDER, 
        filename, 
        as_attachment=True,
        download_name=download_name
    )
    
    # Schedule file deletion after download (cleanup to prevent storage buildup)
    @response.call_on_close
    def cleanup_file():
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                print(f"🗑️  Auto-deleted: {filename}")
        except Exception as e:
            print(f"⚠️  Failed to delete {filename}: {e}")
    
    return response

@app.route('/api/preview/<filename>')
def preview_file(filename):
    """Convert DOCX to HTML for fast preview - no PDF needed"""
    try:
        # Security: validate filename
        if '..' in filename or '/' in filename or '\\' in filename:
            return jsonify({'success': False, 'error': 'Invalid filename'}), 400
        
        # Handle both .docx and .pdf requests (convert .pdf to .docx)
        if filename.endswith('.pdf'):
            filename = filename.replace('.pdf', '.docx')
        
        # Look for DOCX file in output directory
        docx_path = os.path.join(Config.OUTPUT_FOLDER, filename)
        
        if not os.path.exists(docx_path):
            return jsonify({'success': False, 'error': 'File not found'}), 404
        
        # Convert DOCX to HTML using mammoth (fast!)
        try:
            import mammoth
        except ImportError:
            return jsonify({
                'success': False, 
                'error': 'mammoth library not installed. Run: pip install mammoth'
            }), 500
        
        print(f"📄 Converting DOCX to HTML preview: {filename}")
        with open(docx_path, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html_body = result.value
        
        # Wrap HTML with proper styling for resume display
        html_content = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <style>
        body {{
            font-family: 'Calibri', 'Arial', sans-serif;
            font-size: 11pt;
            line-height: 1.4;
            max-width: 850px;
            margin: 20px auto;
            padding: 20px;
            background: white;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 10px 0;
        }}
        td, th {{
            border: 1px solid #333;
            padding: 8px;
            text-align: left;
        }}
        th {{
            background-color: #f0f0f0;
            font-weight: bold;
        }}
        h1, h2, h3 {{
            color: #333;
            margin-top: 15px;
            margin-bottom: 10px;
        }}
        ul, ol {{
            margin-left: 20px;
        }}
        p {{
            margin: 5px 0;
        }}
    </style>
</head>
<body>
{html_body}
</body>
</html>
"""
        
        print(f"✅ HTML preview generated ({len(html_content)} chars)")
        return jsonify({
            'success': True,
            'html': html_content,
            'filename': filename
        })
        
    except Exception as e:
        print(f"❌ Preview error: {e}")
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/templates/<template_id>/thumbnail', methods=['DELETE'])
def delete_template_thumbnail(template_id):
    """Delete template thumbnail to force regeneration"""
    try:
        # Delete from Azure Storage
        if storage_manager.delete_thumbnail(template_id):
            print(f"✅ Thumbnail deleted from storage: {template_id}")
        
        # Delete local cache if exists
        thumbnail_filename = f"{template_id}_thumb.png"
        thumbnail_path = os.path.join(Config.OUTPUT_FOLDER, thumbnail_filename)
        if os.path.exists(thumbnail_path):
            os.remove(thumbnail_path)
            print(f"✅ Local thumbnail cache cleared: {template_id}")
        
        return jsonify({'success': True, 'message': 'Thumbnail deleted'}), 200
    except Exception as e:
        print(f"❌ Error deleting thumbnail: {e}")
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/templates/<template_id>/thumbnail', methods=['GET'])
def get_template_thumbnail(template_id):
    """Generate and return template thumbnail image with Azure Storage caching"""
    try:
        import platform
        
        # Get template from persistent storage first, then fallback to memory DB
        template = persistent_db.get_template(template_id)
        if not template:
            # Fallback to local database
            template = db.get_template(template_id)
            print(f"✅ Template retrieved from local fallback for thumbnail: {template_id}")
        else:
            print(f"✅ Template retrieved from persistent storage for thumbnail: {template_id}")
        
        if not template:
            return jsonify({'success': False, 'message': 'Template not found'}), 404
        
        # Local thumbnail path
        thumbnail_filename = f"{template_id}_thumb.png"
        thumbnail_path = os.path.join(Config.OUTPUT_FOLDER, thumbnail_filename)
        
        # Check if thumbnail exists in Azure Storage first
        if storage_manager.thumbnail_exists(template_id):
            # Download from Azure to local cache
            if storage_manager.download_thumbnail(template_id, thumbnail_path):
                print(f"✅ Thumbnail served from Azure Storage: {template_id}")
                response = send_from_directory(Config.OUTPUT_FOLDER, thumbnail_filename, mimetype='image/png')
                response.headers['Cache-Control'] = 'public, max-age=86400, immutable'
                response.headers['ETag'] = template_id
                return response
        
        # Thumbnail doesn't exist in Azure - need to generate it
        # Thumbnails require Windows-only tooling (COM/docx2pdf). On non-Windows, use fallback.
        if platform.system().lower() != 'windows' or os.name != 'nt':
            print(f"⚠️ Thumbnail generation not supported on {platform.system()}, using fallback")
            
            # Try to create document preview thumbnail
            try:
                from utils.document_thumbnail import save_placeholder_thumbnail, create_docx_preview_thumbnail
                print(f"🎨 Creating document preview thumbnail for: {template['name']}")
                
                # First try to download the actual DOCX file for preview
                temp_template_path = os.path.join(Config.TEMPLATE_FOLDER, template['filename'])
                docx_preview_created = False
                
                try:
                    # Check if template is from persistent storage or memory DB
                    template_from_persistent = persistent_db.get_template(template_id) is not None
                    
                    if template_from_persistent:
                        # Try to download template file from Azure Storage
                        if persistent_db.download_template_file(template_id, template['filename'], temp_template_path):
                            print(f"📄 Downloaded template file from Azure for preview: {template['filename']}")
                            if create_docx_preview_thumbnail(temp_template_path, thumbnail_path):
                                print(f"✅ Created DOCX preview thumbnail: {template_id}")
                                docx_preview_created = True
                            # Clean up downloaded file
                            try:
                                os.remove(temp_template_path)
                            except:
                                pass
                    else:
                        # Template is from memory DB - file should be in local TEMPLATE_FOLDER
                        local_template_path = os.path.join(Config.TEMPLATE_FOLDER, template['filename'])
                        if os.path.exists(local_template_path):
                            print(f"📄 Using local template file for preview: {template['filename']}")
                            if create_docx_preview_thumbnail(local_template_path, thumbnail_path):
                                print(f"✅ Created DOCX preview thumbnail from local file: {template_id}")
                                docx_preview_created = True
                        else:
                            print(f"⚠️ Local template file not found: {local_template_path}")
                except Exception as download_e:
                    print(f"⚠️ Could not access template for preview: {download_e}")
                
                # If DOCX preview failed, use enhanced placeholder
                if not docx_preview_created:
                    if save_placeholder_thumbnail(template['name'], template_id, thumbnail_path):
                        print(f"✅ Created enhanced document thumbnail: {template_id}")
                    else:
                        print(f"❌ Failed to create document thumbnail for: {template_id}")
                        return jsonify({'error': 'Failed to create thumbnail'}), 500
                
                # Upload thumbnail to Azure Storage for caching
                try:
                    if storage_manager.upload_thumbnail(template_id, thumbnail_path):
                        print(f"✅ Document thumbnail uploaded to Azure: {template_id}")
                except Exception as upload_e:
                    print(f"⚠️ Failed to upload thumbnail to Azure: {upload_e}")
                
                # Return the thumbnail
                if os.path.exists(thumbnail_path):
                    response = send_from_directory(Config.OUTPUT_FOLDER, thumbnail_filename, mimetype='image/png')
                    response.headers['Cache-Control'] = 'public, max-age=3600'
                    response.headers['ETag'] = f"{template_id}-document"
                    return response
                else:
                    print(f"❌ Thumbnail file not found after creation: {thumbnail_path}")
                    return jsonify({'error': 'Thumbnail file not created'}), 500
                    
            except ImportError as ie:
                print(f"❌ Import error for document_thumbnail: {ie}")
                import traceback
                traceback.print_exc()
                return jsonify({'error': 'Document thumbnail not available'}), 500
            except Exception as e:
                print(f"❌ Document thumbnail creation failed: {e}")
                import traceback
                traceback.print_exc()
                return jsonify({'error': f'Thumbnail generation failed: {str(e)}'}), 500
        
        # Get template file - handle both persistent storage and memory DB
        temp_template_path = os.path.join(Config.TEMPLATE_FOLDER, template['filename'])
        template_from_persistent = persistent_db.get_template(template_id) is not None
        
        if template_from_persistent:
            # Download from Azure Storage
            if not persistent_db.download_template_file(template_id, template['filename'], temp_template_path):
                print(f"❌ Failed to download template file from Azure: {template_id}")
                return jsonify({'success': False, 'message': 'Template file not found in storage'}), 404
        else:
            # Template is from memory DB - file should already be in local TEMPLATE_FOLDER
            local_template_path = os.path.join(Config.TEMPLATE_FOLDER, template['filename'])
            if not os.path.exists(local_template_path):
                print(f"❌ Local template file not found: {local_template_path}")
                return jsonify({'success': False, 'message': 'Template file not found locally'}), 404
            temp_template_path = local_template_path
        
        # Generate thumbnail from the downloaded template
        try:
            import pythoncom
            from docx2pdf import convert
            import fitz  # PyMuPDF for PDF to image conversion
            from PIL import Image
            
            print(f"🖼️ Generating thumbnail for template: {template_id}")
            
            # Convert DOCX to PDF first
            temp_pdf = os.path.join(Config.OUTPUT_FOLDER, f"{template_id}_temp.pdf")
            
            pythoncom.CoInitialize()
            try:
                convert(temp_template_path, temp_pdf)
            finally:
                pythoncom.CoUninitialize()
            
            # Convert first page of PDF to image
            if os.path.exists(temp_pdf):
                pdf_document = fitz.open(temp_pdf)
                first_page = pdf_document[0]
                # Render at 120 DPI for faster loading
                pix = first_page.get_pixmap(matrix=fitz.Matrix(120/72, 120/72))
                
                # Save as PNG first
                temp_png = thumbnail_path.replace('.png', '_temp.png')
                pix.save(temp_png)
                pdf_document.close()
                
                # Optimize with PIL for smaller file size
                img = Image.open(temp_png)
                img.save(thumbnail_path, 'PNG', optimize=True, quality=85)
                
                # Clean up temp files
                os.remove(temp_pdf)
                os.remove(temp_png)
                
                # Upload thumbnail to Azure Storage for persistence
                if storage_manager.upload_thumbnail(template_id, thumbnail_path):
                    print(f"✅ Thumbnail uploaded to Azure Storage: {template_id}")
                else:
                    print(f"⚠️ Failed to upload thumbnail to Azure Storage: {template_id}")
                
                # Clean up downloaded template file
                try:
                    os.remove(temp_template_path)
                except:
                    pass
                
                print(f"✅ Thumbnail generated successfully: {template_id}")
            else:
                return jsonify({'success': False, 'message': 'PDF conversion failed'}), 500
                
        except Exception as e:
            print(f"⚠️ Thumbnail generation failed: {e}")
            traceback.print_exc()
            return jsonify({'success': False, 'message': f'Thumbnail generation failed: {str(e)}'}), 500
        
        # Return the thumbnail image with aggressive caching
        response = send_from_directory(Config.OUTPUT_FOLDER, thumbnail_filename, mimetype='image/png')
        response.headers['Cache-Control'] = 'public, max-age=86400, immutable'  # Cache for 24 hours
        response.headers['ETag'] = template_id  # Use template ID as ETag
        return response
        
    except Exception as e:
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)}), 500

# ===== Static/SPA routes =====
@app.route('/')
def serve_index():
    # Serve the built React index.html
    try:
        return send_from_directory(frontend_dir, 'index.html')
    except Exception:
        return '<h1>Resume Formatter</h1>', 200

@app.route('/manifest.json')
def serve_manifest():
    try:
        return send_from_directory(frontend_dir, 'manifest.json', mimetype='application/json')
    except Exception:
        return jsonify({'error': 'manifest not found'}), 404

@app.route('/favicon.ico')
def serve_favicon():
    try:
        return send_from_directory(frontend_dir, 'favicon.ico')
    except Exception:
        return '', 204

@app.route('/favicon.svg')
def serve_favicon_svg():
    try:
        return send_from_directory(frontend_dir, 'favicon.svg')
    except Exception:
        return '', 204

@app.route('/static/<path:filename>')
def serve_static(filename):
    """Serve static files (CSS, JS, etc.)"""
    try:
        return send_from_directory(static_dir, filename)
    except Exception:
        return '', 404

@app.route('/api/templates/<template_id>', methods=['DELETE'])
def delete_template(template_id):
    """Delete template from both local storage and Azure"""
    print(f"🗑️ DELETE request received for template: {template_id}")
    try:
        # Get template info - try persistent storage first, then fallback
        template = persistent_db.get_template(template_id)
        if not template:
            print(f"⚠️ Template not found in persistent storage, trying fallback...")
            template = db.get_template(template_id)
        
        print(f"📋 Template found: {template}")
        if not template:
            print(f"❌ Template not found in any database: {template_id}")
            return jsonify({'success': False, 'message': 'Template not found'}), 404
        
        # Delete from local storage
        file_path = os.path.join(Config.TEMPLATE_FOLDER, template['filename'])
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
                print(f"✅ Deleted local template file: {file_path}")
            except Exception as e:
                print(f"⚠️ Failed to delete local file: {e}")
        
        # Delete from Azure Storage
        try:
            storage_manager.delete_template_file(template_id, template['filename'])
            print(f"✅ Deleted template from Azure Storage: {template_id}")
        except Exception as e:
            print(f"⚠️ Failed to delete from Azure Storage: {e}")
        
        # Delete thumbnail from Azure Storage
        try:
            storage_manager.delete_thumbnail(template_id)
            print(f"✅ Deleted thumbnail from Azure Storage: {template_id}")
        except Exception as e:
            print(f"⚠️ Failed to delete thumbnail from Azure Storage: {e}")
        
        # Delete from database (both databases to be safe)
        try:
            db.delete_template(template_id)
            print(f"✅ Deleted from local database: {template_id}")
        except Exception as e:
            print(f"⚠️ Failed to delete from local database: {e}")
        
        try:
            persistent_db.delete_template(template_id)
            print(f"✅ Deleted from persistent database: {template_id}")
        except Exception as e:
            print(f"⚠️ Failed to delete from persistent database: {e}")
        
        print(f"✅ Template deleted successfully: {template_id}")
        return jsonify({'success': True, 'message': 'Template deleted successfully'})
    except Exception as e:
        print(f"❌ Error deleting template: {e}")
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/templates/<template_id>', methods=['PUT'])
def rename_template(template_id):
    """Rename template in both local and persistent storage"""
    print(f"✏️ PUT request received to rename template: {template_id}")
    try:
        data = request.get_json()
        if not data or 'name' not in data:
            return jsonify({'success': False, 'message': 'Missing template name'}), 400

        new_name = data['name'].strip()
        if not new_name:
            return jsonify({'success': False, 'message': 'Template name cannot be empty'}), 400

        # Get template from persistent storage first
        template = persistent_db.get_template(template_id)
        if not template:
            # Fallback to local database
            template = db.get_template(template_id)
            print(f"✅ Template retrieved from local fallback: {template_id}")
        else:
            print(f"✅ Template retrieved from persistent storage: {template_id}")

        if not template:
            print(f"❌ Template not found: {template_id}")
            return jsonify({'success': False, 'message': 'Template not found'}), 404

        # Update template name in persistent storage
        try:
            # Update the template name
            template['name'] = new_name

            # Save to persistent storage
            persistent_success = persistent_db.add_template(
                template_id,
                new_name,
                template['filename'],
                template['file_type'],
                template.get('format_data', {}),
                template.get('cai_contact')
            )

            if persistent_success:
                print(f"✅ Template renamed in persistent storage: {template_id} -> {new_name}")

                # Also update local database for backward compatibility
                try:
                    db.add_template(
                        template_id,
                        new_name,
                        template['filename'],
                        template['file_type'],
                        template.get('format_data', {})
                    )
                    print(f"✅ Template renamed in local database: {template_id} -> {new_name}")
                except Exception as e:
                    print(f"⚠️ Failed to update local database: {e}")

                return jsonify({
                    'success': True,
                    'message': 'Template renamed successfully',
                    'template': {
                        'id': template_id,
                        'name': new_name
                    }
                })
            else:
                print(f"❌ Failed to rename template in persistent storage")
                return jsonify({'success': False, 'message': 'Failed to rename template'}), 500

        except Exception as e:
            print(f"❌ Error renaming template: {e}")
            traceback.print_exc()
            return jsonify({'success': False, 'message': str(e)}), 500

    except Exception as e:
        print(f"❌ Error processing rename request: {e}")
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/templates/<template_id>/content', methods=['GET'])
def get_template_content(template_id):
    """Get template content for editing"""
    try:
        print(f"📖 GET /api/templates/{template_id}/content - Request received")

        # Get template from persistent storage first, then fallback
        template = persistent_db.get_template(template_id)
        if not template:
            print(f"⚠️ Template not in persistent storage, checking local DB...")
            template = db.get_template(template_id)

        if not template:
            print(f"❌ Template not found: {template_id}")
            return jsonify({'success': False, 'message': 'Template not found in database'}), 404

        print(f"✅ Template found: {template['name']} (file: {template['filename']})")

        # Get template file
        template_filename = template['filename']
        local_template_path = os.path.join(Config.TEMPLATE_FOLDER, template_filename)

        print(f"📁 Looking for template file at: {local_template_path}")

        # Download from persistent storage if not local
        if not os.path.exists(local_template_path):
            print(f"📥 Template not local, attempting download from persistent storage...")
            try:
                download_success = persistent_db.download_template_file(template_id, template_filename, local_template_path)
                if not download_success:
                    print(f"❌ Failed to download template file from persistent storage")
                    return jsonify({'success': False, 'message': 'Template file not available in storage'}), 404
                print(f"✅ Template downloaded successfully")
            except Exception as dl_error:
                print(f"❌ Download error: {dl_error}")
                return jsonify({'success': False, 'message': f'Failed to download template: {str(dl_error)}'}), 500

        # Verify file exists
        if not os.path.exists(local_template_path):
            print(f"❌ Template file still not found after download attempt: {local_template_path}")
            return jsonify({'success': False, 'message': 'Template file not found on disk'}), 404

        print(f"✅ Template file exists, extracting content...")

        # Extract text content from DOCX using Mammoth (more reliable)
        try:
            import mammoth

            print(f"📄 Extracting content using Mammoth...")

            with open(local_template_path, "rb") as docx_file:
                result = mammoth.convert_to_html(docx_file)
                html_content = result.value
                messages = result.messages

                # Log any warnings
                for message in messages:
                    print(f"  ⚠️  {message}")

            if not html_content or html_content.strip() == "":
                html_content = "<p>No content found in template</p>"

            print(f"✅ Template content extracted successfully")
            return jsonify({
                'success': True,
                'content': html_content,
                'template_name': template['name']
            })

        except Exception as e:
            print(f"❌ Error extracting template content: {e}")
            traceback.print_exc()

            # Return user-friendly error message
            error_msg = str(e)
            if "not a Word file" in error_msg or "content type" in error_msg:
                error_msg = "This file appears to be corrupted or is not a valid DOCX file. Please re-upload or save as .docx format."

            return jsonify({
                'success': False,
                'message': f'Cannot edit this template: {error_msg}'
            }), 500

    except Exception as e:
        print(f"❌ Error getting template content: {e}")
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/templates/<template_id>/content', methods=['PUT'])
def update_template_content(template_id):
    """Update template content from editor"""
    try:
        data = request.get_json()
        if not data or 'content' not in data:
            return jsonify({'success': False, 'message': 'No content provided'}), 400

        content_html = data['content']

        # Get template
        template = persistent_db.get_template(template_id)
        if not template:
            template = db.get_template(template_id)

        if not template:
            return jsonify({'success': False, 'message': 'Template not found'}), 404

        # Get template file
        template_filename = template['filename']
        local_template_path = os.path.join(Config.TEMPLATE_FOLDER, template_filename)

        # Download from persistent storage if not local
        if not os.path.exists(local_template_path):
            print(f"📥 Downloading template file for updating...")
            download_success = persistent_db.download_template_file(template_id, template_filename, local_template_path)
            if not download_success:
                return jsonify({'success': False, 'message': 'Template file not available'}), 404

        # Update DOCX content (preserving as much formatting as possible)
        try:
            from docx import Document
            from bs4 import BeautifulSoup
            import re

            print(f"📝 Updating template content...")

            doc = Document(local_template_path)
            soup = BeautifulSoup(content_html, 'html.parser')

            # Extract plain text from HTML
            new_text_content = soup.get_text(separator='\n').strip()

            # Try to match paragraphs and update them in place (preserves formatting)
            new_lines = [line.strip() for line in new_text_content.split('\n') if line.strip()]

            # Update existing paragraphs without removing them (preserves formatting)
            para_index = 0
            for paragraph in doc.paragraphs:
                if para_index < len(new_lines):
                    # Update text while keeping formatting
                    if paragraph.text.strip():  # Only update non-empty paragraphs
                        paragraph.text = new_lines[para_index]
                        para_index += 1

            # If there are more new lines than existing paragraphs, add them
            while para_index < len(new_lines):
                doc.add_paragraph(new_lines[para_index])
                para_index += 1

            # Save updated document
            doc.save(local_template_path)
            print(f"✅ Template content updated ({para_index} paragraphs)")

            # Upload to persistent storage
            upload_success = persistent_db.upload_template_file(template_id, local_template_path, template_filename)

            if upload_success:
                print(f"✅ Template content updated: {template_id}")

                # Delete cached thumbnail to force regeneration
                try:
                    storage_manager.delete_thumbnail(template_id)
                    print(f"✅ Thumbnail cache cleared for updated template")
                except:
                    pass

                return jsonify({
                    'success': True,
                    'message': 'Template updated successfully'
                })
            else:
                print(f"⚠️ Failed to upload updated template to persistent storage")
                return jsonify({'success': False, 'message': 'Failed to save changes'}), 500

        except Exception as e:
            print(f"❌ Error updating template content: {e}")
            traceback.print_exc()
            return jsonify({'success': False, 'message': f'Failed to update content: {str(e)}'}), 500

    except Exception as e:
        print(f"❌ Error updating template: {e}")
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/onlyoffice/status', methods=['GET'])
def onlyoffice_status():
    """Check OnlyOffice Document Server status"""
    try:
        import requests
        # Check local Docker container
        onlyoffice_url = Config.ONLYOFFICE_URL.rstrip('/')
        response = requests.get(f'{onlyoffice_url}/healthcheck', timeout=2)
        if response.status_code == 200:
            return jsonify({
                'success': True,
                'status': 'running',
                'available': True,
                'message': 'OnlyOffice Document Server is running'
            })
        else:
            return jsonify({
                'success': False,
                'status': 'error',
                'available': False,
                'message': f'OnlyOffice returned status code {response.status_code}'
            })
    except requests.exceptions.ConnectionError:
        return jsonify({
            'success': False,
            'status': 'offline',
            'available': False,
            'message': 'OnlyOffice Document Server is not running. Start it with: docker start onlyoffice-docs'
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'status': 'error',
            'available': False,
            'message': str(e)
        })

# Add debug logging for all requests
@app.before_request
def log_request_path():
    print(f"📡 Request path: {request.path}")

# Serve React Frontend (for single container deployment)
@app.route('/')
def index():
    """Serve React index.html"""
    return send_from_directory(app.template_folder, 'index.html')

# Serve logo.png from frontend root
@app.route('/logo.png')
def serve_logo():
    """Serve logo.png from frontend build root"""
    print(f"🎨 Serving logo from: {app.template_folder}")
    return send_from_directory(app.template_folder, 'logo.png')

# Serve static files manually (React CSS/JS)
@app.route('/static/<path:filename>')
def serve_static_files(filename):
    """Serve static files from React build"""
    print(f"🎨 Serving static file: {filename}")
    print(f"📁 Static folder: {app.static_folder}")
    return send_from_directory(app.static_folder, filename)

# SPA fallback handled by 404 error handler below

@app.route('/api/cleanup', methods=['POST'])
def cleanup_old_files():
    """Cleanup old output files to prevent storage buildup"""
    try:
        import time
        from datetime import datetime, timedelta
        
        # Remove files older than 1 hour
        cutoff_time = time.time() - 3600  # 1 hour
        deleted_count = 0
        
        for filename in os.listdir(Config.OUTPUT_FOLDER):
            file_path = os.path.join(Config.OUTPUT_FOLDER, filename)
            if os.path.isfile(file_path):
                file_mtime = os.path.getmtime(file_path)
                if file_mtime < cutoff_time:
                    try:
                        os.remove(file_path)
                        deleted_count += 1
                        print(f"🗑️  Deleted old file: {filename}")
                    except Exception as e:
                        print(f"⚠️  Failed to delete {filename}: {e}")
        
        return jsonify({
            'success': True,
            'deleted_count': deleted_count,
            'message': f'Cleaned up {deleted_count} old file(s)'
        })
    except Exception as e:
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)}), 500

def cleanup_on_startup():
    """Clean up all output files on startup"""
    try:
        deleted_count = 0
        for filename in os.listdir(Config.OUTPUT_FOLDER):
            file_path = os.path.join(Config.OUTPUT_FOLDER, filename)
            if os.path.isfile(file_path) and filename.startswith('formatted_'):
                try:
                    os.remove(file_path)
                    deleted_count += 1
                except Exception as e:
                    print(f"⚠️  Failed to delete {filename}: {e}")
        
        if deleted_count > 0:
            print(f"🗑️  Startup cleanup: Removed {deleted_count} old output file(s)")
    except Exception as e:
        print(f"⚠️  Startup cleanup failed: {e}")

# ===== SPA Fallback (MUST BE LAST ROUTE) =====
@app.route('/<path:path>')
def serve_spa(path):
    """Catch-all route for React SPA - serve index.html for all non-API routes"""
    if path.startswith('api/'):
        return '', 404
    try:
        return send_from_directory(frontend_dir, 'index.html')
    except Exception:
        return '<h1>Resume Formatter</h1><p>Frontend not available</p>', 200

if __name__ == '__main__':
    # CLEANUP OLD FILES ON STARTUP
    cleanup_on_startup()
    
    # PRE-WARM ML MODELS FOR INSTANT FIRST REQUEST 
    # Disabled for faster startup - models will load on first request
    # try:
    #     from utils.model_cache import prewarm_models
    #     prewarm_models()
    # except Exception as e:
    #     print(f"  Model pre-warming failed: {e}")
    #     print("   Models will load on first request instead")
    print("[QUICK] ML models will load on first request")
    
    import socket
    
    # Get local IP for display
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(("8.8.8.8", 80))
        local_ip = s.getsockname()[0]
        s.close()
    except:
        local_ip = "localhost"
    
    # Get port from environment (Azure uses PORT env variable)
    port = int(os.environ.get('PORT', 5000))
    
    # Demo route for new design
    @app.route('/demo')
    def demo_page():
        """Serve the perfect demo HTML page"""
        return send_from_directory('.', 'demo-perfect.html')
    
    # Check Azure Storage connection
    if os.getenv('AZURE_STORAGE_CONNECTION_STRING'):
        print("[AZURE] Storage: Connected")
        print(f"[STORAGE] Manager: {type(storage_manager).__name__}")
    else:
        print("[WARN] Azure Storage: Using local fallback")

    print("\n" + "="*70)
    print("RESUME FORMATTER - SINGLE CONTAINER (COST OPTIMIZED)")
    print("="*70)
    print(f"[OK] Server running on http://0.0.0.0:{port}")
    print(f"[OK] Network access: http://{local_ip}:{port}")
    print("[OK] Frontend: Served from /")
    print("[OK] API: Served from /api/*")
    print(f"[ENV] Environment: {'LOCAL' if Config.IS_LOCAL else 'PRODUCTION'}")
    print(f"[ENV] OnlyOffice URL: {Config.ONLYOFFICE_URL}")
    print(f"[ENV] Backend URL: {Config.BACKEND_URL}")
    print(f"[AZURE] Storage: {'Connected' if os.getenv('AZURE_STORAGE_CONNECTION_STRING') else 'Local Fallback'}")
    print("="*70)
    print("Main Routes:")
    print("   - / - React Frontend")
    print("   - /demo - New Design Demo")
    print("   - /api/format - Format resumes")
    print("   - /api/templates - Manage templates")
    print("   - /api/download/<filename> - Download files")
    print("="*70 + "\n")
    # CRITICAL: Bind to 0.0.0.0 to accept connections from Docker/Azure
    app.run(debug=False, host='0.0.0.0', port=port)
