"""
DOCX Image Extractor for OCR Processing
Extracts embedded images from DOCX files and prepares them for OCR

Handles:
- Images embedded in paragraphs
- Images in headers/footers
- Multiple images per document
- Image format conversion (PNG, JPG, BMP, etc.)
"""

import os
import io
from typing import List, Tuple, Optional
from PIL import Image
from docx import Document
from docx.oxml import parse_xml
import zipfile


class DOCXImageExtractor:
    """Extracts images from DOCX files for OCR processing"""

    def __init__(self):
        self.supported_formats = ['png', 'jpg', 'jpeg', 'bmp', 'tiff', 'gif']

    def extract_images_from_docx(self, docx_path: str) -> List[Tuple[Image.Image, dict]]:
        """
        Extract all images from a DOCX file

        Args:
            docx_path: Path to DOCX file

        Returns:
            List of (PIL Image, metadata dict) tuples
        """
        images = []

        try:
            # Method 1: Extract using python-docx
            doc = Document(docx_path)

            # Extract images from relationships
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_data = rel.target_part.blob
                        img = Image.open(io.BytesIO(image_data))

                        # Convert to RGB if needed
                        if img.mode not in ['RGB', 'L']:
                            img = img.convert('RGB')

                        metadata = {
                            'source': 'docx_relationship',
                            'format': img.format,
                            'size': img.size,
                            'mode': img.mode
                        }

                        images.append((img, metadata))
                        print(f"  [OK] Extracted image: {img.size[0]}x{img.size[1]} ({img.format})")

                    except Exception as e:
                        print(f"  [WARN] Error extracting image from relationship: {e}")
                        continue

            # Method 2: Extract using zipfile (backup method)
            if not images:
                images = self._extract_images_via_zip(docx_path)

        except Exception as e:
            print(f"  [ERROR] Error extracting images from DOCX: {e}")
            return []

        print(f"  [INFO] Total images extracted: {len(images)}")
        return images

    def _extract_images_via_zip(self, docx_path: str) -> List[Tuple[Image.Image, dict]]:
        """
        Extract images using zipfile method (DOCX is a ZIP archive)

        This is a fallback method if python-docx fails
        """
        images = []

        try:
            with zipfile.ZipFile(docx_path, 'r') as docx_zip:
                # Images are stored in word/media/ folder
                image_files = [f for f in docx_zip.namelist()
                             if f.startswith('word/media/') and
                             any(f.lower().endswith(fmt) for fmt in self.supported_formats)]

                for img_file in image_files:
                    try:
                        img_data = docx_zip.read(img_file)
                        img = Image.open(io.BytesIO(img_data))

                        # Convert to RGB if needed
                        if img.mode not in ['RGB', 'L']:
                            img = img.convert('RGB')

                        metadata = {
                            'source': 'docx_zip',
                            'filename': os.path.basename(img_file),
                            'format': img.format,
                            'size': img.size,
                            'mode': img.mode
                        }

                        images.append((img, metadata))
                        print(f"  [OK] Extracted image via ZIP: {img.size[0]}x{img.size[1]} ({img.format})")

                    except Exception as e:
                        print(f"  [WARN] Error extracting {img_file}: {e}")
                        continue

        except Exception as e:
            print(f"  [ERROR] Error extracting images via ZIP: {e}")

        return images

    def combine_images_vertically(self, images: List[Image.Image]) -> Image.Image:
        """
        Combine multiple images into one vertical image

        This is useful when a resume is split across multiple images

        Args:
            images: List of PIL Images

        Returns:
            Combined PIL Image
        """
        if not images:
            return None

        if len(images) == 1:
            return images[0]

        # Calculate total height and max width
        total_height = sum(img.size[1] for img in images)
        max_width = max(img.size[0] for img in images)

        # Create new image
        combined = Image.new('RGB', (max_width, total_height), color='white')

        # Paste images vertically
        y_offset = 0
        for img in images:
            # Center horizontally if image is narrower
            x_offset = (max_width - img.size[0]) // 2
            combined.paste(img, (x_offset, y_offset))
            y_offset += img.size[1]

        print(f"  [OK] Combined {len(images)} images into {max_width}x{total_height}")
        return combined

    def extract_and_combine(self, docx_path: str) -> Optional[Image.Image]:
        """
        Extract all images from DOCX and combine them into one image

        Args:
            docx_path: Path to DOCX file

        Returns:
            Combined PIL Image ready for OCR
        """
        print(f"\n[DOCX Image Extraction] Processing: {os.path.basename(docx_path)}")

        # Extract images
        image_tuples = self.extract_images_from_docx(docx_path)

        if not image_tuples:
            print("  [ERROR] No images found in DOCX")
            return None

        # Extract just the images (without metadata)
        images = [img for img, meta in image_tuples]

        # Combine images
        combined_image = self.combine_images_vertically(images)

        return combined_image

    def extract_text_from_docx(self, docx_path: str) -> str:
        """
        Extract any text content from DOCX (in case some text is not in images)

        Args:
            docx_path: Path to DOCX file

        Returns:
            Extracted text
        """
        try:
            doc = Document(docx_path)
            text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
            return text
        except Exception as e:
            print(f"  [WARN] Error extracting text from DOCX: {e}")
            return ""


# Utility functions
def extract_images_from_docx(docx_path: str) -> List[Tuple[Image.Image, dict]]:
    """Convenience function to extract images from DOCX"""
    extractor = DOCXImageExtractor()
    return extractor.extract_images_from_docx(docx_path)


def docx_to_ocr_image(docx_path: str) -> Optional[Image.Image]:
    """
    Convert DOCX with embedded images to a single OCR-ready image

    Args:
        docx_path: Path to DOCX file

    Returns:
        PIL Image ready for OCR processing
    """
    extractor = DOCXImageExtractor()
    return extractor.extract_and_combine(docx_path)


if __name__ == "__main__":
    import sys

    if len(sys.argv) > 1:
        docx_path = sys.argv[1]

        # Extract and combine
        combined_img = docx_to_ocr_image(docx_path)

        if combined_img:
            # Save for testing
            output_path = "docx_extracted.png"
            combined_img.save(output_path)
            print(f"\n[SUCCESS] Saved combined image to: {output_path}")
            print(f"   Size: {combined_img.size[0]}x{combined_img.size[1]}")
        else:
            print("\n[ERROR] Failed to extract images from DOCX")
    else:
        print("Usage: python docx_image_extractor.py <docx_file>")
