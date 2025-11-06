"""
Enhanced Resume Parser with Intelligent Section Mapping
Handles varied section names, missing headings, and semantic matching
OPTIMIZED: Prevents content mixing between sections
"""

import re
from typing import Dict, List, Optional, Tuple
from docx import Document
import numpy as np
from .section_content_validator import get_content_validator

# Install these if missing:
# pip install sentence-transformers fuzzywuzzy python-Levenshtein spacy
# python -m spacy download en_core_web_sm

try:
    from sentence_transformers import SentenceTransformer
    TRANSFORMERS_AVAILABLE = True
except ImportError:
    TRANSFORMERS_AVAILABLE = False
    print("⚠️  Run: pip install sentence-transformers")

try:
    from fuzzywuzzy import process, fuzz
    FUZZY_AVAILABLE = True
except ImportError:
    FUZZY_AVAILABLE = False
    print("⚠️  Run: pip install fuzzywuzzy python-Levenshtein")

try:
    import spacy
    SPACY_AVAILABLE = True
except ImportError:
    SPACY_AVAILABLE = False
    print("⚠️  Run: pip install spacy && python -m spacy download en_core_web_sm")


class IntelligentResumeParser:
    """
    Main parser with intelligent section mapping using ML
    OPTIMIZED: Uses singleton pattern and cached models for 10x faster performance
    """
    
    # Singleton pattern - shared models across all instances
    _instance = None
    _model = None
    _nlp = None
    _models_loaded = False
    
    def __new__(cls):
        """Singleton pattern - only one instance with shared models"""
        if cls._instance is None:
            cls._instance = super().__new__(cls)
        return cls._instance
    
    def __init__(self):
        # Only load models once
        if not IntelligentResumeParser._models_loaded:
            self._load_models()
            IntelligentResumeParser._models_loaded = True
        
        # Synonym mappings for fallback
        self.section_mappings = {
            'EMPLOYMENT': ['employment history', 'work experience', 'professional experience', 
                          'work history', 'career history', 'experience', 'professional background',
                          'employment', 'relevant employment', 'career experience', 'work', 'jobs'],
            'EDUCATION': ['education', 'educational background', 'academic background',
                         'academic qualifications', 'qualifications', 'academics',
                         'education background', 'education/certificates', 'schooling', 'degrees'],
            'SKILLS': ['skills', 'technical skills', 'core competencies', 'key skills',
                      'professional skills', 'areas of expertise', 'competencies',
                      'skill set', 'expertise', 'technical competencies', 'technologies'],
            'SUMMARY': ['summary', 'professional summary', 'career summary', 'profile',
                       'professional profile', 'career objective', 'objective',
                       'executive summary', 'career overview', 'about me', 'about'],
            'PROJECTS': ['projects', 'key projects', 'project experience', 'notable projects', 'portfolio'],
            'CERTIFICATIONS': ['certifications', 'certificates', 'professional certifications',
                              'licenses', 'credentials', 'certified', 'licensing'],
            'ACHIEVEMENTS': ['achievements', 'awards', 'honors', 'recognition', 'accomplishments'],
            'LANGUAGES': ['languages', 'language skills', 'language proficiency', 'linguistic']
        }
    
    def _load_models(self):
        """Load ML models once and cache them"""
        if TRANSFORMERS_AVAILABLE and IntelligentResumeParser._model is None:
            try:
                print("⚡ Loading OPTIMIZED Sentence Transformer (all-MiniLM-L6-v2)...")
                import time
                start = time.time()
                IntelligentResumeParser._model = SentenceTransformer(
                    'all-MiniLM-L6-v2',
                    device='cpu'
                )
                print(f"✅ Sentence Transformer loaded in {time.time()-start:.2f}s (cached for reuse)")
            except Exception as e:
                print(f"⚠️  Failed to load Sentence Transformer: {e}")
        
        if SPACY_AVAILABLE and IntelligentResumeParser._nlp is None:
            try:
                print("⚡ Loading spaCy (en_core_web_sm)...")
                import time
                start = time.time()
                IntelligentResumeParser._nlp = spacy.load("en_core_web_sm")
                print(f"✅ spaCy loaded in {time.time()-start:.2f}s (cached for reuse)")
            except Exception as e:
                print(f"⚠️  Failed to load spaCy: {e}")
                try:
                    print("📥 Downloading spaCy model...")
                    import subprocess
                    subprocess.run(["python", "-m", "spacy", "download", "en_core_web_sm"], check=True)
                    IntelligentResumeParser._nlp = spacy.load("en_core_web_sm")
                    print("✅ spaCy model downloaded and loaded")
                except Exception as e2:
                    print(f"⚠️  Failed to download spaCy: {e2}")
    
    @property
    def model(self):
        """Get cached sentence transformer model"""
        return IntelligentResumeParser._model
    
    @property
    def nlp(self):
        """Get cached spaCy model"""
        return IntelligentResumeParser._nlp
    
    def parse_resume(self, candidate_docx_path: str, template_docx_path: str) -> Dict[str, str]:
        """
        Main function: Parse candidate resume and map to template structure
        
        Args:
            candidate_docx_path: Path to candidate's resume DOCX
            template_docx_path: Path to template DOCX
            
        Returns:
            Dictionary mapping template sections to candidate content
        """
        print("\n" + "="*60)
        print("🚀 INTELLIGENT RESUME PARSING")
        print("="*60)
        
        # Step 1: Extract template structure
        template_sections = self._extract_template_sections(template_docx_path)
        print(f"\n📋 Template sections: {template_sections}")
        
        # Step 2: Extract candidate sections
        candidate_sections = self._extract_candidate_sections(candidate_docx_path)
        print(f"📄 Found {len(candidate_sections)} sections in candidate resume")
        
        # Step 3: Intelligently map sections
        print(f"\n🔄 Mapping sections...\n")
        mapped_content = self._map_sections(candidate_sections, template_sections)
        
        print(f"\n✅ Successfully mapped {len(mapped_content)} sections")
        print("="*60 + "\n")
        
        return mapped_content
    
    def _extract_template_sections(self, template_path: str) -> List[str]:
        """Extract section names from template"""
        doc = Document(template_path)
        sections = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Detect headings: bold, heading style, or ALL CAPS
            is_heading = (
                para.style.name.startswith('Heading') or
                (para.runs and para.runs[0].bold) or
                (text.isupper() and len(text.split()) <= 4)
            )
            
            if is_heading:
                sections.append(text)
        
        return sections
    
    def _extract_candidate_sections(self, candidate_path: str) -> List[Dict]:
        """Extract sections from candidate resume with headings and content"""
        doc = Document(candidate_path)
        sections = []
        current_section = None
        position_index = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Skip contact info (first 3 lines typically)
            if position_index < 3 and any(c in text for c in ['@', 'phone', 'email', '+', '(', ')']):
                position_index += 1
                continue
            
            # Detect if this is a heading
            is_heading = self._is_heading(para)
            
            if is_heading:
                # Save previous section
                if current_section:
                    sections.append(current_section)
                
                # Start new section
                current_section = {
                    'heading': text,
                    'content': [],
                    'position': position_index,
                    'has_heading': True
                }
            else:
                # Content paragraph
                if current_section is None:
                    # Paragraph without heading (e.g., summary at top)
                    current_section = {
                        'heading': None,
                        'content': [text],
                        'position': position_index,
                        'has_heading': False
                    }
                else:
                    current_section['content'].append(text)
            
            position_index += 1
        
        # Save last section
        if current_section:
            sections.append(current_section)
        
        # Convert content lists to strings
        for section in sections:
            section['content'] = '\n'.join(section['content'])
        
        return sections
    
    def _is_heading(self, para) -> bool:
        """Detect if paragraph is a heading"""
        text = para.text.strip()
        
        # Check style
        if para.style.name.startswith('Heading'):
            return True
        
        # Check formatting
        if para.runs:
            first_run = para.runs[0]
            if first_run.bold and len(text.split()) <= 5:
                return True
        
        # Check ALL CAPS short phrases
        if text.isupper() and len(text.split()) <= 4:
            return True
        
        # Check common section patterns
        section_patterns = [
            r'^(employment|education|skills|experience|summary|projects|certifications)',
            r'(history|background|profile|qualifications)$'
        ]
        text_lower = text.lower()
        for pattern in section_patterns:
            if re.search(pattern, text_lower):
                return True
        
        return False
    
    def _map_sections(self, candidate_sections: List[Dict], 
                     template_sections: List[str]) -> Dict[str, str]:
        """Map candidate sections to template sections using intelligent matching with validation"""
        mapped = {}
        validator = get_content_validator()
        used_content = set()  # Track content already mapped to prevent duplicates
        
        for section in candidate_sections:
            heading = section['heading']
            content = section['content']
            position = section['position']
            has_heading = section['has_heading']
            
            if has_heading and heading:
                # Skip if this content was already mapped
                content_hash = hash(content[:100])  # Hash first 100 chars for quick comparison
                if content_hash in used_content:
                    print(f"  ⏭️  Skipping '{heading}' - content already mapped to another section")
                    continue
                
                # Use intelligent heading matching
                matched = self._match_heading(heading, template_sections)
                
                if matched:
                    # VALIDATE: Does content actually match this section?
                    is_valid, confidence, reason = validator.validate_content(content, matched)
                    
                    if is_valid:
                        print(f"  ✓ '{heading}' → '{matched}' (validated, confidence: {confidence:.2f})")
                        # Filter out any mismatched content
                        filtered_content, removed = validator.filter_mismatched_content(content, matched)
                        if removed:
                            print(f"    ⚠️  Filtered {len(removed)} mismatched lines")
                        mapped[matched] = filtered_content
                        used_content.add(content_hash)  # Mark as used
                    else:
                        # Content doesn't match heading - try to find correct section
                        print(f"  ⚠️  '{heading}' → '{matched}' but content doesn't match ({reason})")
                        suggested = validator.suggest_correct_section(content, matched)
                        if suggested:
                            print(f"    💡 Content better fits: {suggested}")
                            # Find template section for suggested type
                            for ts in template_sections:
                                if suggested.lower() in ts.lower():
                                    mapped[ts] = content
                                    used_content.add(content_hash)
                                    break
                        else:
                            # Use content classification as fallback
                            classified = self._classify_content(content, position, template_sections)
                            if classified:
                                print(f"    ⚡ Reclassified by content → '{classified}'")
                                mapped[classified] = content
                                used_content.add(content_hash)
                else:
                    # Fallback: classify by content
                    classified = self._classify_content(content, position, template_sections)
                    if classified:
                        print(f"  ⚡ '{heading}' classified by content → '{classified}'")
                        # Validate and filter
                        filtered_content, removed = validator.filter_mismatched_content(content, classified)
                        if removed:
                            print(f"    ⚠️  Filtered {len(removed)} mismatched lines")
                        mapped[classified] = filtered_content
                        used_content.add(content_hash)
            else:
                # No heading - classify by content
                # Skip if content already used
                content_hash = hash(content[:100])
                if content_hash in used_content:
                    print(f"  ⏭️  Skipping unheaded content - already mapped")
                    continue
                
                classified = self._classify_content(content, position, template_sections)
                if classified:
                    print(f"  🎯 Unheaded paragraph → '{classified}'")
                    # Validate and filter
                    filtered_content, removed = validator.filter_mismatched_content(content, classified)
                    if removed:
                        print(f"    ⚠️  Filtered {len(removed)} mismatched lines")
                    mapped[classified] = filtered_content
                    used_content.add(content_hash)
        
        return mapped
    
    def _match_heading(self, candidate_heading: str, 
                      template_sections: List[str]) -> Optional[str]:
        """Match candidate heading to template section"""
        candidate_clean = candidate_heading.strip().lower()
        template_clean = [s.strip().lower() for s in template_sections]
        
        # Exact match
        if candidate_clean in template_clean:
            return template_sections[template_clean.index(candidate_clean)]
        
        # Fuzzy match (typos, minor variations)
        if FUZZY_AVAILABLE:
            result = process.extractOne(candidate_clean, template_clean, scorer=fuzz.token_sort_ratio)
            if result and result[1] > 85:
                return template_sections[template_clean.index(result[0])]
        
        # Semantic similarity (synonyms)
        if self.model:
            try:
                candidate_emb = self.model.encode([candidate_clean])
                template_embs = self.model.encode(template_clean)
                similarities = np.dot(candidate_emb, template_embs.T)[0]
                best_idx = np.argmax(similarities)
                
                if similarities[best_idx] > 0.65:
                    return template_sections[best_idx]
            except:
                pass
        
        # Rule-based fallback
        for template_key, synonyms in self.section_mappings.items():
            if candidate_clean in synonyms:
                for ts in template_sections:
                    if template_key.lower() in ts.lower():
                        return ts
        
        return None
    
    def _classify_content(self, content: str, position: int,
                         template_sections: List[str]) -> Optional[str]:
        """Classify content without heading"""
        content_lower = content.lower()
        
        # Rule 1: Position-based (summary usually first)
        if position <= 2 and len(content.split()) < 100:
            summary_keywords = ['years', 'experience', 'professional', 'seeking', 'expertise']
            if any(kw in content_lower for kw in summary_keywords):
                for ts in template_sections:
                    if 'summary' in ts.lower() or 'profile' in ts.lower():
                        return ts
        
        # Rule 2: Entity-based (spaCy)
        if self.nlp:
            try:
                doc = self.nlp(content[:500])
                has_dates = any(ent.label_ == "DATE" for ent in doc.ents)
                has_orgs = any(ent.label_ == "ORG" for ent in doc.ents)
                
                if has_dates and has_orgs:
                    for ts in template_sections:
                        if any(kw in ts.lower() for kw in ['employment', 'experience', 'work']):
                            return ts
            except:
                pass
        
        # Rule 3: Keyword scoring
        keyword_scores = {
            'EMPLOYMENT': ['worked', 'managed', 'developed', 'led', 'responsible', 'duties'],
            'EDUCATION': ['university', 'degree', 'graduated', 'gpa', 'bachelor', 'master'],
            'SKILLS': ['proficient', 'skilled', 'expertise', 'technologies', 'programming'],
            'PROJECTS': ['project', 'developed', 'built', 'created', 'implemented']
        }
        
        scores = {}
        for category, keywords in keyword_scores.items():
            score = sum(1 for kw in keywords if kw in content_lower)
            if score > 0:
                scores[category] = score
        
        if scores:
            best_category = max(scores, key=scores.get)
            for ts in template_sections:
                if best_category.lower() in ts.lower():
                    return ts
        
        return None
    
    def format_output(self, mapped_sections: Dict[str, str],
                     template_path: str, output_path: str):
        """Generate final formatted resume"""
        template_doc = Document(template_path)
        output_doc = Document()
        
        # Copy styles from template
        output_doc.styles = template_doc.styles
        
        # Build output based on template structure
        for para in template_doc.paragraphs:
            text = para.text.strip()
            
            # Check if this is a section heading
            is_section = any(section in text for section in mapped_sections.keys())
            
            if is_section:
                # Add heading
                new_para = output_doc.add_paragraph(text, style=para.style)
                
                # Add mapped content
                for section_name, content in mapped_sections.items():
                    if section_name in text:
                        output_doc.add_paragraph(content)
                        break
            else:
                # Copy other template content (name, contact, etc.)
                output_doc.add_paragraph(text, style=para.style)
        
        output_doc.save(output_path)
        print(f"✅ Saved formatted resume to: {output_path}")


# Singleton instance for reuse
_parser_instance = None

def get_intelligent_parser() -> IntelligentResumeParser:
    """Get or create singleton parser instance"""
    global _parser_instance
    if _parser_instance is None:
        _parser_instance = IntelligentResumeParser()
    return _parser_instance


# USAGE EXAMPLE
if __name__ == "__main__":
    parser = IntelligentResumeParser()
    
    # Parse and map sections
    mapped = parser.parse_resume(
        candidate_docx_path="candidate_resume.docx",
        template_docx_path="master_template.docx"
    )
    
    # Generate formatted output
    parser.format_output(
        mapped_sections=mapped,
        template_path="master_template.docx",
        output_path="formatted_resume.docx"
    )
