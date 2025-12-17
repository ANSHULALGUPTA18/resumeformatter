"""
DOC to DOCX Converter Utility
Converts old .doc files to .docx format while preserving structure
Also supports ODT and RTF files using Python libraries
"""

import os
import subprocess
import tempfile
from pathlib import Path

# Import Python-based converters
try:
    from odf import text, teletype
    from odf.opendocument import load as load_odt
    HAS_ODFPY = True
except ImportError:
    HAS_ODFPY = False

try:
    from striprtf.striprtf import rtf_to_text
    HAS_STRIPRTF = True
except ImportError:
    HAS_STRIPRTF = False

try:
    from docx import Document
    from docx.shared import Pt
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

def _convert_odt_to_docx_python(odt_path, docx_path):
    """Convert ODT to DOCX using Python libraries"""
    try:
        if not HAS_ODFPY or not HAS_PYTHON_DOCX:
            return False

        # Load ODT file
        odt_doc = load_odt(odt_path)

        # Create new DOCX document
        docx_doc = Document()

        # Extract text from ODT
        all_paragraphs = odt_doc.getElementsByType(text.P)
        for para in all_paragraphs:
            para_text = teletype.extractText(para)
            if para_text.strip():
                docx_doc.add_paragraph(para_text)

        # Save DOCX
        docx_doc.save(docx_path)
        return True
    except Exception as e:
        print(f"  Python ODT conversion failed: {e}")
        return False

def _convert_rtf_to_docx_python(rtf_path, docx_path):
    """Convert RTF to DOCX using Python libraries"""
    try:
        if not HAS_STRIPRTF or not HAS_PYTHON_DOCX:
            return False

        # Read RTF file
        with open(rtf_path, 'r', encoding='utf-8', errors='ignore') as f:
            rtf_content = f.read()

        # Convert RTF to plain text
        plain_text = rtf_to_text(rtf_content)

        # Create new DOCX document
        docx_doc = Document()

        # Add paragraphs
        for line in plain_text.split('\n'):
            if line.strip():
                docx_doc.add_paragraph(line.strip())

        # Save DOCX
        docx_doc.save(docx_path)
        return True
    except Exception as e:
        print(f"  Python RTF conversion failed: {e}")
        return False

def convert_doc_to_docx(doc_file_path):
    """
    Convert a .doc, .odt, or .rtf file to .docx format while preserving structure

    Args:
        doc_file_path (str): Path to the .doc/.odt/.rtf file

    Returns:
        str: Path to the converted .docx file, or None if conversion failed
    """
    try:
        # Check if the file exists
        if not os.path.exists(doc_file_path):
            print(f"❌ File not found: {doc_file_path}")
            return None

        # Check if the file needs conversion
        ext = doc_file_path.lower()
        if not (ext.endswith('.doc') or ext.endswith('.odt') or ext.endswith('.rtf')):
            print(f"❌ Not a convertible file type: {doc_file_path}")
            return None
            
        # Create output path with .docx extension
        doc_path = Path(doc_file_path)
        docx_path = doc_path.with_suffix('.docx')
        
        # Get file extension for display
        file_ext = doc_path.suffix.upper()
        print(f"🔄 Converting {file_ext} to DOCX: {doc_path.name} → {docx_path.name} (preserving structure)...")

        # Special handling for ODT and RTF - try Python libraries first (no external dependencies)
        if ext.endswith('.odt'):
            print(f"  Trying Python-based ODT conversion...")
            if _convert_odt_to_docx_python(doc_file_path, str(docx_path)):
                print(f"✅ Successfully converted using Python (odfpy): {docx_path}")
                return str(docx_path)

        if ext.endswith('.rtf'):
            print(f"  Trying Python-based RTF conversion...")
            if _convert_rtf_to_docx_python(doc_file_path, str(docx_path)):
                print(f"✅ Successfully converted using Python (striprtf): {docx_path}")
                return str(docx_path)

        # Method 1: Try using LibreOffice (best for preserving structure)
        if _convert_with_libreoffice(doc_file_path, str(docx_path)):
            print(f"✅ Successfully converted using LibreOffice: {docx_path}")
            return str(docx_path)

        # Method 2: Try using pandoc (good structure preservation)
        if _convert_with_pandoc(doc_file_path, str(docx_path)):
            print(f"✅ Successfully converted using pandoc: {docx_path}")
            return str(docx_path)

        # Method 3: Try using unoconv (if available)
        if _convert_with_unoconv(doc_file_path, str(docx_path)):
            print(f"✅ Successfully converted using unoconv: {docx_path}")
            return str(docx_path)
            
        print(f"❌ All conversion methods failed for: {doc_file_path}")
        print(f"💡 Suggestion: Install LibreOffice in the container for better .doc/.odt/.rtf support")
        return None
        
    except Exception as e:
        print(f"❌ Error converting {doc_file_path}: {str(e)}")
        return None

def _convert_with_libreoffice(doc_path, docx_path):
    """Try converting using LibreOffice headless mode"""
    try:
        # Try common LibreOffice paths
        libreoffice_paths = [
            'libreoffice',
            '/usr/bin/libreoffice',
            '/opt/libreoffice/program/soffice',
            'soffice'
        ]
        
        for lo_path in libreoffice_paths:
            try:
                # Create temp directory for output
                with tempfile.TemporaryDirectory() as temp_dir:
                    # Run LibreOffice conversion
                    cmd = [
                        lo_path,
                        '--headless',
                        '--convert-to', 'docx',
                        '--outdir', temp_dir,
                        doc_path
                    ]
                    
                    result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
                    
                    if result.returncode == 0:
                        # Find the converted file
                        doc_name = Path(doc_path).stem
                        temp_docx = os.path.join(temp_dir, f"{doc_name}.docx")
                        
                        if os.path.exists(temp_docx):
                            # Move to final location
                            os.rename(temp_docx, docx_path)
                            return True
                            
            except (subprocess.TimeoutExpired, FileNotFoundError, subprocess.SubprocessError):
                continue
                
        return False
        
    except Exception:
        return False

def _convert_with_unoconv(doc_path, docx_path):
    """Try converting using unoconv (LibreOffice command-line tool)"""
    try:
        cmd = ['unoconv', '-f', 'docx', '-o', docx_path, doc_path]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
        
        if result.returncode == 0 and os.path.exists(docx_path):
            return True
            
        return False
        
    except (subprocess.TimeoutExpired, FileNotFoundError, subprocess.SubprocessError):
        return False

def _convert_with_pandoc(doc_path, docx_path):
    """Try converting using pandoc"""
    try:
        cmd = ['pandoc', '-f', 'doc', '-t', 'docx', '-o', docx_path, doc_path]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        
        if result.returncode == 0 and os.path.exists(docx_path):
            return True
            
        return False
        
    except (subprocess.TimeoutExpired, FileNotFoundError, subprocess.SubprocessError):
        return False

def is_doc_file(filename):
    """Check if a file is a .doc file"""
    return filename.lower().endswith('.doc')

def needs_conversion(filename):
    """Check if a file needs conversion from .doc/.odt/.rtf to .docx"""
    ext = filename.lower()
    return ext.endswith('.doc') or ext.endswith('.odt') or ext.endswith('.rtf')
