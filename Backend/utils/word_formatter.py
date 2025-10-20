"""
Enhanced Word Document Formatter
Handles both .doc and .docx templates
Preserves all formatting, images, headers, footers
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
import os
import re
import shutil
import traceback
import json

# Try to import win32com for .doc support
try:
    import win32com.client
    HAS_WIN32 = True
except ImportError:
    HAS_WIN32 = False
    print("⚠️  win32com not available - .doc files will have limited support")

class WordFormatter:
    """Enhanced Word document formatting"""
    
    def __init__(self, resume_data, template_analysis, output_path):
        self.resume_data = resume_data
        self.template_analysis = template_analysis
        self.output_path = output_path
        self.template_path = template_analysis.get('template_path')
        self.template_type = template_analysis.get('template_type')
        
    def format(self):
        """Main formatting method"""
        print(f"\n{'='*70}")
        print(f"📝 WORD DOCUMENT FORMATTING")
        print(f"{'='*70}\n")
        
        print(f"📄 Template: {os.path.basename(self.template_path)}")
        print(f"👤 Candidate: {self.resume_data['name']}")
        print(f"📁 Output: {os.path.basename(self.output_path)}\n")
        
        try:
            # Handle .doc files
            if self.template_path.lower().endswith('.doc'):
                return self._format_doc_file()
            else:
                return self._format_docx_file()
                
        except Exception as e:
            print(f"❌ Error formatting Word document: {e}")
            traceback.print_exc()
            return False
    
    def _format_doc_file(self):
        """Handle .doc files (old Word format)"""
        print("📋 Processing .doc file (old Word format)...")
        
        if HAS_WIN32:
            # Convert .doc to .docx first
            print("✓ Converting .doc to .docx...")
            docx_path = self._convert_doc_to_docx(self.template_path)
            
            if docx_path:
                # Update template path temporarily
                original_path = self.template_path
                self.template_path = docx_path
                
                # Format the docx
                result = self._format_docx_file()
                
                # Cleanup
                try:
                    os.remove(docx_path)
                except:
                    pass
                
                self.template_path = original_path
                return result
            else:
                print("❌ Failed to convert .doc to .docx")
                return False
        else:
            print("⚠️  Cannot process .doc files without win32com")
            print("💡 Please convert template to .docx format or install pywin32")
            return False
    
    def _convert_doc_to_docx(self, doc_path):
        """Convert .doc to .docx using Word COM"""
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            # Open .doc file
            doc = word.Documents.Open(os.path.abspath(doc_path))
            
            # Save as .docx
            docx_path = doc_path.replace('.doc', '_temp.docx')
            doc.SaveAs2(os.path.abspath(docx_path), FileFormat=16)  # 16 = docx format
            
            doc.Close()
            word.Quit()
            
            print(f"✓ Converted to: {docx_path}")
            return docx_path
            
        except Exception as e:
            print(f"❌ Conversion error: {e}")
            return None

    def _postprocess_with_word_com(self, docx_path):
        """Final pass using Word COM to replace placeholders that may live in shapes/text boxes.
        Only affects placeholders that still exist (i.e., were not handled by python-docx).
        """
        if not HAS_WIN32:
            return
        try:
            import pythoncom
            pythoncom.CoInitialize()  # Initialize COM for this thread
            
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(os.path.abspath(docx_path))

            def find_replace_in_range(rng, find_text, replace_text, wildcard=False):
                find = rng.Find
                find.ClearFormatting()
                find.Replacement.ClearFormatting()
                find.Text = find_text
                find.Replacement.Text = replace_text
                find.Forward = True
                find.Wrap = 1  # wdFindContinue
                find.MatchCase = False
                find.MatchWholeWord = False
                find.MatchWildcards = bool(wildcard)
                find.Execute(Replace=2)  # wdReplaceAll

            # Build replacement strings from resume_data
            # SUMMARY
            summary_lines = self._find_matching_resume_section('summary', self.resume_data.get('sections', {})) or []
            summary_text = (self.resume_data.get('summary') or '').strip()
            summary_replace = ''
            if summary_lines:
                summary_replace = '\r'.join(['• ' + s.strip().lstrip('•–—-*● ') for s in summary_lines if s.strip()])
            elif summary_text:
                summary_replace = summary_text

            # SKILLS - limit to prevent "string too long" error
            skills_list = self.resume_data.get('skills', []) or []
            skill_lines = []
            for s in skills_list[:10]:  # Limit to 10 skills
                skill_name = (s if isinstance(s, str) else s.get('name', '')).strip()
                if skill_name and len(skill_name) < 50:  # Skip very long skill names
                    skill_lines.append('• ' + skill_name)
            skills_replace = '\r'.join(skill_lines)
            # Limit to 255 chars
            if len(skills_replace) > 255:
                skills_replace = skills_replace[:252] + '...'

            # EDUCATION
            education = self.resume_data.get('education', []) or []
            if not education:
                sect = self._find_matching_resume_section('education', self.resume_data.get('sections', {})) or []
                if sect:
                    education = self._build_education_from_bullets(sect)
            edu_lines = []
            for edu in education[:3]:  # Limit to 3 entries to avoid "string too long" error
                deg = (edu.get('degree') or '').strip()
                inst = (edu.get('institution') or '').strip()
                yr = self._clean_duration((edu.get('year') or '').strip())
                
                # Keep it short for COM
                if deg and len(deg) > 50:
                    deg = deg[:47] + '...'
                if inst and len(inst) > 40:
                    inst = inst[:37] + '...'
                
                parts = []
                if deg:
                    parts.append(deg)
                if inst:
                    parts.append(inst)
                if yr:
                    parts.append(yr)
                line = ' - '.join(parts[:-1]) if len(parts) > 1 else (parts[0] if parts else '')
                if yr and line:
                    line = f"{line} {yr}"
                if line and len(line) < 200:  # Safety check
                    edu_lines.append('• ' + line)
            education_replace = '\r'.join(edu_lines)
            
            # Limit total length to 255 chars (Word COM limit)
            if len(education_replace) > 255:
                education_replace = education_replace[:252] + '...'
            
            print(f"  📝 COM replacement strings prepared:")
            print(f"     - Education: {len(education_replace)} chars, {len(edu_lines)} lines")
            print(f"     - Skills: {len(skills_replace)} chars")
            print(f"     - Summary: {len(summary_replace)} chars")

            # Build candidate name replace string
            candidate_name = (self.resume_data.get('name') or '').strip()
            bracket_name = f"<{candidate_name}>" if candidate_name else ''

            # Story ranges include shapes/text frames and headers/footers
            story = doc.StoryRanges(1)  # wdMainTextStory = 1
            while story is not None:
                # Candidate name → bracketed name
                if bracket_name:
                    # Exact case and uppercase variants
                    find_replace_in_range(story, candidate_name, bracket_name, wildcard=False)
                    find_replace_in_range(story, candidate_name.upper(), bracket_name, wildcard=False)

                if summary_replace:
                    for pat in ["<summary>", "<professional summary>", "<profile>", "professional summary", "<summary*>"]:
                        find_replace_in_range(story, pat, summary_replace, wildcard='*' in pat)
                if skills_replace:
                    for pat in ["<skills>", "<technical skills>", "<list skills>", "<skills*>"]:
                        find_replace_in_range(story, pat, skills_replace, wildcard='*' in pat)
                if education_replace:
                    edu_pats = [
                        "<List candidate’s education background>",
                        "<List candidate’s education background>",
                        "<list candidate’s education background>",
                        "<list candidate’s education background>",
                        "list candidate’s education background",
                        "list candidate’s education background",
                        "education background",
                        "<education background>",
                        "<education>",
                        "<academic background>",
                        "<academic qualifications>",
                    ]
                    for pat in edu_pats:
                        find_replace_in_range(story, pat, education_replace, wildcard=False)
                story = story.NextStoryRange

            # Also traverse shapes explicitly (in case some shapes are not part of StoryRanges loop)
            def replace_in_shapes(shapes):
                for shp in shapes:
                    try:
                        if shp.TextFrame.HasText:
                            rng = shp.TextFrame.TextRange
                            # Candidate name
                            if bracket_name:
                                find_replace_in_range(rng, candidate_name, bracket_name)
                                find_replace_in_range(rng, candidate_name.upper(), bracket_name)
                            if summary_replace:
                                for pat in ["<summary>", "<professional summary>", "<profile>"]:
                                    find_replace_in_range(rng, pat, summary_replace)
                            if skills_replace:
                                for pat in ["<skills>", "<technical skills>", "<list skills>"]:
                                    find_replace_in_range(rng, pat, skills_replace)
                            if education_replace:
                                for pat in [
                                    "<List candidate’s education background>",
                                    "<List candidate’s education background>",
                                    "<list candidate’s education background>",
                                    "<list candidate’s education background>",
                                    "list candidate’s education background",
                                    "list candidate’s education background",
                                    "education background",
                                ]:
                                    find_replace_in_range(rng, pat, education_replace)
                    except Exception:
                        continue

            replace_in_shapes(doc.Shapes)
            for sec in doc.Sections:
                replace_in_shapes(sec.Headers(1).Shapes)
                replace_in_shapes(sec.Headers(2).Shapes)
                replace_in_shapes(sec.Headers(3).Shapes)
                replace_in_shapes(sec.Footers(1).Shapes)
                replace_in_shapes(sec.Footers(2).Shapes)
                replace_in_shapes(sec.Footers(3).Shapes)

            doc.Save()
            doc.Close(False)
            word.Quit()
            pythoncom.CoUninitialize()  # Clean up COM
            print("✓ COM post-processing complete (shapes/text boxes handled)")
        except Exception as e:
            print(f"⚠️  COM post-processing error: {e}")
            try:
                import pythoncom
                pythoncom.CoUninitialize()
            except:
                pass

    def _scan_primary_anchors(self, doc):
        """Scan the template once to locate primary anchors for SUMMARY, SKILLS, EMPLOYMENT, EDUCATION.
        If multiple EDUCATION headings exist and one is embedded immediately after EMPLOYMENT
        (likely sample content), pick the later EDUCATION heading as the primary.
        Returns (primary_anchors, all_anchors).
        """
        def is_heading_text(t):
            if not t:
                return False
            t = t.strip()
            return len(t) < 50

        keys = {
            'EMPLOYMENT': ['EMPLOYMENT HISTORY', 'WORK HISTORY', 'PROFESSIONAL EXPERIENCE', 'WORK EXPERIENCE', 'CAREER HISTORY', 'EMPLOYMENT'],
            'EDUCATION': ['EDUCATION', 'ACADEMIC BACKGROUND', 'EDUCATIONAL BACKGROUND', 'ACADEMIC QUALIFICATIONS', 'QUALIFICATIONS', 'EDUCATION BACKGROUND'],
            'SKILLS': ['SKILLS', 'TECHNICAL SKILLS'],
            'SUMMARY': ['SUMMARY', 'PROFESSIONAL SUMMARY', 'PROFILE', 'OBJECTIVE', 'CAREER SUMMARY', 'EXECUTIVE SUMMARY'],
        }
        all_anchors = {k: [] for k in keys}
        for idx, p in enumerate(doc.paragraphs):
            txt = (p.text or '').strip().upper()
            if not is_heading_text(txt):
                continue
            for k, aliases in keys.items():
                if any(a == txt or txt.startswith(a) for a in aliases):
                    all_anchors[k].append(idx)
                    break

        primary = {k: (v[0] if v else None) for k, v in all_anchors.items()}

        # Heuristic: if an EDUCATION heading appears immediately (<= 8 paras) after EMPLOYMENT
        # and there exists another EDUCATION later, treat the first as embedded placeholder
        emp_idx = primary.get('EMPLOYMENT')
        edu_list = all_anchors.get('EDUCATION') or []
        if emp_idx is not None and len(edu_list) > 1:
            first_edu = edu_list[0]
            if 0 <= (first_edu - emp_idx) <= 8:
                primary['EDUCATION'] = edu_list[-1]

        print("\n🔎 Anchor scan:")
        for k in ['SUMMARY', 'SKILLS', 'EMPLOYMENT', 'EDUCATION']:
            print(f"  - {k}: primary={primary.get(k)} all={all_anchors.get(k)}")

        return primary, all_anchors
    
    def _format_docx_file(self):
        """Format .docx file"""
        print("📋 Processing .docx file...")
        
        # Open template
        doc = Document(self.template_path)
        
        print(f"✓ Template loaded: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
        
        # Pre-scan anchors so we always insert into the correct template sections
        self._primary_anchors, self._all_anchors = self._scan_primary_anchors(doc)
        
        # Initialize section tracking flags
        self._summary_inserted = False
        self._experience_inserted = False
        self._skills_inserted = False
        self._education_inserted = False
        
        # DEBUG: Show what education data we have
        print(f"\n📊 Resume data check:")
        print(f"   - education list: {len(self.resume_data.get('education', []))} entries")
        print(f"   - sections.education: {len(self.resume_data.get('sections', {}).get('education', []))} lines")
        if self.resume_data.get('education'):
            for i, edu in enumerate(self.resume_data.get('education', [])[:3]):
                print(f"      {i+1}. {edu.get('degree', 'N/A')[:40]} | {edu.get('institution', 'N/A')[:30]} | {edu.get('year', 'N/A')}")
        
        # Ensure CAI CONTACT section is inserted with persistent data
        try:
            self._ensure_cai_contact(doc)
        except Exception as e:
            print(f"⚠️  CAI Contact insertion error: {e}")
        
        # Show what data we have from resume
        print(f"\n📊 Resume Data Available:")
        print(f"  • Name: {self.resume_data.get('name', 'NOT FOUND')}")
        print(f"  • Email: {self.resume_data.get('email', 'NOT FOUND')}")
        print(f"  • Phone: {self.resume_data.get('phone', 'NOT FOUND')}")
        print(f"  • Experience entries: {len(self.resume_data.get('experience', []))}")
        print(f"  • Education entries: {len(self.resume_data.get('education', []))}")
        print(f"  • Skills: {len(self.resume_data.get('skills', []))}")
        print(f"  • Sections: {list(self.resume_data.get('sections', {}).keys())}")
        
        # Create comprehensive replacement map
        replacements = self._create_replacement_map()
        print(f"\n📝 Created {len(replacements)} replacement mappings")
        
        # CRITICAL: Process skills tables in-place based on table headers (respect template order)
        table_replaced = 0
        print(f"\n🔍 STEP 1: Scanning {len(doc.tables)} tables...")
        for table_idx, table in enumerate(doc.tables):
            # Check if this is a skills table
            if self._is_skills_table(table):
                print(f"  📊 Found skills table at index {table_idx}")
                skills_filled = self._fill_skills_table(table)
                print(f"  ✅ Filled {skills_filled} skill rows")
                table_replaced += skills_filled
                # Mark as inserted if we actually filled rows
                if skills_filled > 0:
                    self._skills_inserted = True
        
        print(f"\n✓ Processed {table_replaced} table entries")
        
        # STEP 2: Replace in all paragraphs
        replaced_count = 0
        print(f"\n🔍 STEP 2: Scanning {len(doc.paragraphs)} paragraphs for placeholders...")
        
        # Prepare bracketed name and compute name anchor index
        candidate_name = self.resume_data.get('name', '').strip()
        bracketed_name = f"<{candidate_name}>" if candidate_name else '<Candidate Name>'
        self._name_anchor_idx = None
        try:
            # Look for name or name placeholder in main content area (skip early CAI CONTACT section)
            name_placeholder_patterns = [
                r'<\s*[Cc]andidate[^>]*[Nn]ame[^>]*>',
                r'<\s*[Nn]ame\s*>',
                r'<\s*[Ff]ull\s*[Nn]ame\s*>',
                r'<\s*YOUR\s*NAME\s*>',
            ]
            for idx, p in enumerate(doc.paragraphs[:40]):
                t = (p.text or '').strip()
                # Skip very early paragraphs (likely CAI CONTACT)
                if idx < 5:
                    continue
                # Check for actual name or bracketed name
                if (candidate_name and t == candidate_name) or (bracketed_name and t == bracketed_name):
                    self._name_anchor_idx = idx
                    print(f"  📍 Name anchor found at paragraph {idx}: '{t}'")
                    break
                # Check for name placeholder patterns
                for pat in name_placeholder_patterns:
                    if re.search(pat, t):
                        self._name_anchor_idx = idx
                        print(f"  📍 Name placeholder anchor found at paragraph {idx}: '{t[:50]}'")
                        break
                if self._name_anchor_idx is not None:
                    break
        except Exception as e:
            print(f"  ⚠️  Name anchor detection error: {e}")
            self._name_anchor_idx = None
        
        # PRE-PASS: Remove all SKILLS sections before EMPLOYMENT to prevent left-panel fills
        print(f"\n🔍 PRE-PASS: Removing any SKILLS sections in CAI CONTACT/left panel only...")
        try:
            emp_idx = self._primary_anchors.get('EMPLOYMENT')
            name_idx = getattr(self, '_name_anchor_idx', None)
            # Limit removal strictly to very early left-panel region (CAI CONTACT), not the main content
            left_boundary = name_idx if name_idx is not None else 5
            skills_removed = 0
            # Determine scan limit: before EMPLOYMENT but never beyond the left boundary
            scan_limit = min(emp_idx, left_boundary) if emp_idx is not None else left_boundary
            if scan_limit and scan_limit > 0:
                for idx in range(scan_limit - 1, -1, -1):
                    if idx >= len(doc.paragraphs):
                        continue
                    para = doc.paragraphs[idx]
                    t = (para.text or '').strip().upper()
                    if t in ('SKILLS', 'TECHNICAL SKILLS') or ('SKILLS' in t and len(t) < 30):
                        print(f"  🗑️  Removing SKILLS heading at para {idx} (CAI CONTACT area)")
                        # Clear content after this heading until next section or for ~20 lines
                        j = idx + 1
                        cleared = 0
                        while j < len(doc.paragraphs) and cleared < 20:
                            para_j = doc.paragraphs[j]
                            txt = (para_j.text or '').strip().upper()
                            # Stop at next major section
                            if len(txt) < 50 and any(h in txt for h in ['EMPLOYMENT', 'WORK HISTORY', 'EDUCATION', 'SUMMARY', 'CAI CONTACT', 'CERTIFICATIONS']):
                                break
                            try:
                                for r in para_j.runs:
                                    r.text = ''
                            except Exception:
                                pass
                            j += 1
                            cleared += 1
                        # Delete the SKILLS heading itself
                        try:
                            self._delete_paragraph(para)
                            skills_removed += 1
                        except Exception:
                            pass
            print(f"  ✅ Removed {skills_removed} SKILLS section(s) from CAI CONTACT area")
        except Exception as e:
            print(f"  ⚠️  Pre-pass error: {e}")
        
        for para_idx, paragraph in enumerate(doc.paragraphs):
            if not paragraph.text.strip():
                continue
                
            # Check each replacement
            for key, value in replacements.items():
                if self._text_contains(paragraph.text, key):
                    print(f"  📍 Found '{key}' in paragraph {para_idx}: '{paragraph.text[:50]}...'")
                    count = self._replace_in_paragraph(paragraph, key, value)
                    if count > 0:
                        print(f"  ✅ Replaced with: '{value[:50]}...'")
                    else:
                        print(f"  ⚠️  Found but couldn't replace (might be in multiple runs)")

            # Regex-driven fallback for angle bracket placeholders with variations
            # Candidate name generic patterns - very flexible to catch all variations
            name_patterns = [
                r"<\s*Candidate['']?s?\s+full\s+name\s*>",
                r"<\s*Candidate['']?s?\s+name\s*>",
                r"<\s*Name\s*>",
                r"<\s*Full\s+Name\s*>",
                r"<\s*YOUR\s+NAME\s*>",
            ]
            for pat in name_patterns:
                if re.search(pat, paragraph.text, re.IGNORECASE):
                    before = paragraph.text
                    # Replace with angle-bracketed name format: <CANDIDATE NAME>
                    self._regex_replace_paragraph(paragraph, pat, bracketed_name)
                    if paragraph.text != before:
                        print(f"  ✅ Regex replaced candidate name in paragraph {para_idx} with {bracketed_name}")
                        replaced_count += 1

            # Generic catch-all: any <...> containing both 'candidate' and 'name' (any order)
            generic_name_pat = r"<[^>]*(?:candidate[^>]*name|name[^>]*candidate)[^>]*>"
            if re.search(generic_name_pat, paragraph.text, re.IGNORECASE):
                before = paragraph.text
                self._regex_replace_paragraph(paragraph, generic_name_pat, bracketed_name)
                if paragraph.text != before:
                    print(f"  ✅ Generic regex replaced candidate name in paragraph {para_idx} with {bracketed_name}")
                    replaced_count += 1

            # CRITICAL: Check if this is an EMPLOYMENT HISTORY section heading
            if not self._experience_inserted:
                para_upper = paragraph.text.strip().upper()
                is_emp_heading = any(h in para_upper for h in [
                    'EMPLOYMENT HISTORY', 'WORK HISTORY', 'PROFESSIONAL EXPERIENCE',
                    'WORK EXPERIENCE', 'CAREER HISTORY', 'EMPLOYMENT'
                ])
                
                # Gate by primary EMPLOYMENT anchor if present
                is_primary_anchor = (self._primary_anchors.get('EMPLOYMENT') is None) or (para_idx == self._primary_anchors.get('EMPLOYMENT'))
                if is_emp_heading and len(paragraph.text.strip()) < 50 and is_primary_anchor:
                    print(f"  💼 Found EMPLOYMENT HISTORY heading at paragraph {para_idx}: '{paragraph.text[:60]}'")
                    
                    experience_data = self.resume_data.get('experience', [])
                    if not experience_data or len(experience_data) == 0:
                        print(f"     ⚠️  No experience data to insert, skipping")
                    else:
                        print(f"     → Will insert {len(experience_data)} experience entries after heading")
                        
                        # Look for instructional text in the NEXT paragraph
                        next_para = None
                        is_instruction = False
                        
                        if para_idx + 1 < len(doc.paragraphs):
                            next_para = doc.paragraphs[para_idx + 1]
                            next_text = next_para.text.strip().lower()
                            
                            # Check if next paragraph is instructional text
                            is_instruction = any(phrase in next_text for phrase in [
                                'please list', 'list the candidate', 'please provide',
                                'insert employment', 'add employment', 'employment details'
                            ])
                        
                        if is_instruction:
                            print(f"     → Found instructional text: '{next_para.text[:60]}'")
                            
                            experience_data = self.resume_data.get('experience', [])
                            if experience_data and len(experience_data) > 0:
                                print(f"     → Will insert {len(experience_data)} experience entries after heading")
                                # DEBUG: Show what data we have
                                for idx, exp in enumerate(experience_data[:3]):
                                    print(f"        Entry {idx+1}: {exp.get('company', 'N/A')[:30]} | {exp.get('role', 'N/A')[:30]} | {len(exp.get('details', []))} details")
                                
                                # Clear the instructional paragraph
                                for run in next_para.runs:
                                    run.text = ''
                                
                                # CRITICAL: Clear any existing employment content after instructional text
                                # Look ahead and clear paragraphs UNTIL the next section heading.
                                # DO NOT clear the next section heading itself (e.g., EDUCATION).
                                paras_to_clear = []
                                stop_headings = [
                                    'EDUCATION', 'ACADEMIC BACKGROUND', 'EDUCATIONAL BACKGROUND',
                                    'ACADEMIC QUALIFICATIONS', 'QUALIFICATIONS',
                                    'SKILLS', 'TECHNICAL SKILLS',
                                    'SUMMARY', 'PROFESSIONAL SUMMARY', 'PROFILE', 'OBJECTIVE',
                                    'CERTIFICATIONS', 'PROJECTS'
                                ]
                                for check_idx in range(para_idx + 2, min(para_idx + 50, len(doc.paragraphs))):
                                    check_para = doc.paragraphs[check_idx]
                                    check_text = check_para.text.strip().upper()
                                    # If we hit an EDUCATION heading that is NOT the primary one, clear it as placeholder
                                    if 'EDUCATION' in check_text and len(check_text) < 50:
                                        primary_edu = self._primary_anchors.get('EDUCATION')
                                        if primary_edu is not None and check_idx != primary_edu:
                                            print(f"     → Clearing embedded EDUCATION placeholder at {check_idx}")
                                            paras_to_clear.append(check_para)
                                            continue
                                    # Stop at the next section heading (including primary EDUCATION)
                                    if any(h in check_text for h in stop_headings) and len(check_text) < 50:
                                        print(f"     → Stopped clearing at next section: {check_text[:30]}")
                                        break
                                    # Clear only non-heading paragraphs in between
                                    paras_to_clear.append(check_para)
                                
                                print(f"     → Clearing {len(paras_to_clear)} old content paragraphs")
                                for p in paras_to_clear:
                                    try:
                                        for run in p.runs:
                                            run.text = ''
                                    except:
                                        pass
                                
                                # Insert employment blocks after the cleared paragraph
                                last_element = next_para
                                for idx, exp in enumerate(experience_data[:10]):
                                    print(f"        → Inserting job {idx+1}: {exp.get('company', 'N/A')[:25]} with {len(exp.get('details', []))} bullets")
                                    block = self._insert_experience_block(doc, last_element, exp)
                                    if block:
                                        last_element = block
                                        print(f"           ✓ Inserted successfully")
                                    else:
                                        print(f"           ✗ Failed to insert")
                                
                                self._experience_inserted = True
                                # Remember tail paragraph to place SKILLS after EMPLOYMENT
                                self._employment_tail_para = last_element
                                print(f"  ✅ Inserted employment data after EMPLOYMENT HISTORY heading")
                                replaced_count += 1
                        else:
                            # No instructional text - there's existing employment content
                            # Clear all content after heading until next section
                            print(f"     → No instructional text found, clearing existing employment content")
                            
                            paras_to_clear = []
                            stop_headings = [
                                'EDUCATION', 'ACADEMIC BACKGROUND', 'EDUCATIONAL BACKGROUND',
                                'ACADEMIC QUALIFICATIONS', 'QUALIFICATIONS',
                                'SKILLS', 'TECHNICAL SKILLS',
                                'SUMMARY', 'PROFESSIONAL SUMMARY', 'PROFILE', 'OBJECTIVE',
                                'CERTIFICATIONS', 'PROJECTS'
                            ]
                            for check_idx in range(para_idx + 1, min(para_idx + 50, len(doc.paragraphs))):
                                check_para = doc.paragraphs[check_idx]
                                check_text = check_para.text.strip().upper()
                                # If we hit an EDUCATION heading that is NOT the primary one, clear it as placeholder
                                if 'EDUCATION' in check_text and len(check_text) < 50:
                                    primary_edu = self._primary_anchors.get('EDUCATION')
                                    if primary_edu is not None and check_idx != primary_edu:
                                        print(f"     → Clearing embedded EDUCATION placeholder at {check_idx}")
                                        paras_to_clear.append(check_para)
                                        continue
                                # Stop if we hit another section heading (do not clear the heading itself)
                                if any(h in check_text for h in stop_headings) and len(check_text) < 50:
                                    print(f"     → Stopped clearing at section: {check_text[:30]}")
                                    break
                                # Clear this paragraph (it's old employment content)
                                paras_to_clear.append(check_para)
                            
                            print(f"     → Clearing {len(paras_to_clear)} old employment paragraphs")
                            for p in paras_to_clear:
                                try:
                                    for run in p.runs:
                                        run.text = ''
                                except:
                                    pass
                            
                            # Insert employment blocks after the heading
                            last_element = paragraph
                            for exp in experience_data[:10]:
                                block = self._insert_experience_block(doc, last_element, exp)
                                if block:
                                    last_element = block
                            
                            self._experience_inserted = True
                            self._employment_tail_para = last_element
                            print(f"  ✅ Inserted employment data after EMPLOYMENT HISTORY heading (no instruction text)")
                            replaced_count += 1
            
            # Employment placeholder generic patterns (very flexible)
            if not self._experience_inserted:
                # Avoid inserting EXP content inside table cells (e.g., SKILLS table headers/columns)
                try:
                    if self._paragraph_in_table(paragraph):
                        pass
                    else:
                        emp_patterns = [
                            r"<[^>]*list[^>]*candidate['']?s?[^>]*employment[^>]*history[^>]*>",
                            r"<[^>]*employment[^>]*history[^>]*>",
                            r"<[^>]*work[^>]*history[^>]*>",
                            r"<[^>]*professional[^>]*experience[^>]*>",
                            r"<[^>]*career[^>]*(history|experience)[^>]*>",
                            r"<[^>]*history[^>]*(employ|employer|work|career)[^>]*>",
                            r"<[^>]*list[^>]*employment[^>]*history[^>]*>",
                        ]
                        for emp_pat in emp_patterns:
                            if re.search(emp_pat, paragraph.text, re.IGNORECASE):
                                print(f"  💼 Found employment placeholder in paragraph {para_idx}: '{paragraph.text[:60]}'")
                                
                                # Use structured experience data (not sections)
                                experience_data = self.resume_data.get('experience', [])
                                
                                if experience_data and len(experience_data) > 0:
                                    print(f"     → Will replace with {len(experience_data)} experience entries")
                                    
                                    # Clear the placeholder paragraph
                                    self._regex_replace_paragraph(paragraph, emp_pat, '')
                                    
                                    # Insert properly formatted experience blocks
                                    last_element = paragraph
                                    for exp in experience_data[:10]:  # Limit to 10 entries
                                        block = self._insert_experience_block(doc, last_element, exp)
                                        if block:
                                            last_element = block
                                    
                                    # CRITICAL: Set flag to prevent duplicate insertion in _add_sections_content
                                    self._experience_inserted = True
                                    # Store tail paragraph for SKILLS placement
                                    self._employment_tail_para = last_element
                                    
                                    print(f"  ✅ Replaced employment placeholder with structured blocks")
                                    replaced_count += 1
                                    break
                                else:
                                    # Fallback: try to use sections data
                                    content = self._find_matching_resume_section('experience', self.resume_data.get('sections', {}))
                                    if content:
                                        bullets = []
                                        for item in content[:10]:
                                            if item.strip():
                                                bullets.append('• ' + item.strip().lstrip('•').strip())
                                        self._regex_replace_paragraph(paragraph, emp_pat, '\n'.join(bullets))
                                        # Set flag even in fallback to prevent duplication
                                        self._experience_inserted = True
                                        print(f"  ✅ Regex replaced experience placeholder (fallback)")
                                        replaced_count += 1
                                        break
                except Exception:
                    pass

            # CRITICAL: Check if this is a SUMMARY section heading (only before EMPLOYMENT)
            if not self._summary_inserted:
                para_upper = paragraph.text.strip().upper()
                is_summary_heading = any(h in para_upper for h in [
                    'SUMMARY', 'PROFESSIONAL SUMMARY', 'PROFILE', 'OBJECTIVE',
                    'CAREER SUMMARY', 'EXECUTIVE SUMMARY'
                ])
                
                emp_anchor_idx = self._primary_anchors.get('EMPLOYMENT')
                name_idx = getattr(self, '_name_anchor_idx', None)
                # Accept summary headings only if they are near the name (within 5 paras after name) and before employment
                is_position_ok = False
                if name_idx is not None:
                    if emp_anchor_idx is not None:
                        is_position_ok = (para_idx <= name_idx + 5) and (para_idx < emp_anchor_idx)
                    else:
                        is_position_ok = (para_idx <= name_idx + 5)
                else:
                    # Fallback: very early paragraphs only
                    is_position_ok = (emp_anchor_idx is None and para_idx < 8) or (emp_anchor_idx is not None and para_idx < min(emp_anchor_idx, 8))
                
                if is_summary_heading and len(paragraph.text.strip()) < 50 and is_position_ok:
                    print(f"  📝 Found SUMMARY heading at paragraph {para_idx}: '{paragraph.text[:60]}'")
                    
                    summary_text = (self.resume_data.get('summary') or '').strip()
                    summary_lines = self._find_matching_resume_section('summary', self.resume_data.get('sections', {}))
                    
                    if summary_text or summary_lines:
                        print(f"     → Will insert summary content after heading")
                        
                        # Clear any existing summary content after heading
                        paras_to_clear = []
                        for check_idx in range(para_idx + 1, min(para_idx + 10, len(doc.paragraphs))):
                            check_para = doc.paragraphs[check_idx]
                            check_text = check_para.text.strip().upper()
                            
                            # Stop at next section
                            if any(h in check_text for h in ['EMPLOYMENT', 'WORK HISTORY', 'EDUCATION', 'SKILLS']) and len(check_text) < 50:
                                break
                            
                            paras_to_clear.append(check_para)
                        
                        for p in paras_to_clear:
                            try:
                                for run in p.runs:
                                    run.text = ''
                            except:
                                pass
                        
                        # Insert summary content
                        if summary_lines:
                            self._insert_skills_bullets(doc, paragraph, summary_lines)
                        else:
                            summary_para = self._insert_paragraph_after(paragraph, summary_text)
                            if summary_para:
                                for run in summary_para.runs:
                                    run.font.size = Pt(10)
                        
                        self._summary_inserted = True
                        print(f"  ✅ Inserted summary after SUMMARY heading")
                        replaced_count += 1
            
            # Summary placeholder patterns - flexible (we clear them and insert after name later)
            if not self._summary_inserted:
                summary_patterns = [
                r"<[^>]*summary[^>]*>",
                r"<[^>]*professional[^>]*summary[^>]*>",
                r"<[^>]*profile[^>]*>",
            ]
            for sum_pat in summary_patterns:
                if re.search(sum_pat, paragraph.text, re.IGNORECASE):
                    print(f"  📝 Found summary placeholder in paragraph {para_idx} — clearing and deferring insertion after name")
                    self._regex_replace_paragraph(paragraph, sum_pat, '')
                    replaced_count += 1
                    break

            # Skills placeholder patterns - flexible
            skills_patterns = [
                r"<[^>]*skills[^>]*>",
                r"<[^>]*technical[^>]*skills[^>]*>",
                r"<[^>]*list[^>]*skills[^>]*>",
            ]
            for skl_pat in skills_patterns:
                if re.search(skl_pat, paragraph.text, re.IGNORECASE):
                    skills_list = self.resume_data.get('skills', [])
                    # Always clear the placeholder, but only insert content AFTER employment is inserted
                    print(f"  🧰 Found skills placeholder in paragraph {para_idx} — clearing; will insert after EMPLOYMENT")
                    self._regex_replace_paragraph(paragraph, skl_pat, '')
                    if skills_list and self._experience_inserted:
                        self._insert_skills_bullets(doc, paragraph, skills_list)
                        self._skills_inserted = True
                    replaced_count += 1
                    break

            # SKILLS section heading (respect template order; do not force after EMPLOYMENT)
            if not self._skills_inserted:
                para_upper = paragraph.text.strip().upper()
                is_skills_heading = any(h in para_upper for h in [
                    'SKILLS', 'TECHNICAL SKILLS'
                ])
                is_primary_anchor = (self._primary_anchors.get('SKILLS') is None) or (para_idx == self._primary_anchors.get('SKILLS'))
                if is_skills_heading and len(paragraph.text.strip()) < 50 and is_primary_anchor:
                    print(f"  🧰 Found SKILLS heading at paragraph {para_idx}: '{paragraph.text[:60]}'")
                    skills_list = self.resume_data.get('skills', [])
                    if skills_list:
                        paras_to_clear = []
                        stop_headings = ['EMPLOYMENT', 'WORK HISTORY', 'PROFESSIONAL EXPERIENCE', 'WORK EXPERIENCE', 'CAREER HISTORY', 'EDUCATION', 'SUMMARY', 'CERTIFICATIONS', 'PROJECTS']
                        for check_idx in range(para_idx + 1, min(para_idx + 30, len(doc.paragraphs))):
                            check_para = doc.paragraphs[check_idx]
                            check_text = check_para.text.strip().upper()
                            if any(h in check_text for h in stop_headings) and len(check_text) < 50:
                                break
                            paras_to_clear.append(check_para)
                        for p in paras_to_clear:
                            try:
                                for r in p.runs:
                                    r.text = ''
                            except:
                                pass
                        self._insert_skills_bullets(doc, paragraph, skills_list)
                        self._skills_inserted = True
                        replaced_count += 1

            # CRITICAL: Check if this is an EDUCATION section heading
            if not self._education_inserted:
                para_upper = paragraph.text.strip().upper()
                is_edu_heading = any(h in para_upper for h in [
                    'EDUCATION', 'ACADEMIC BACKGROUND', 'EDUCATIONAL BACKGROUND',
                    'ACADEMIC QUALIFICATIONS', 'QUALIFICATIONS', 'EDUCATION BACKGROUND'
                ])
                # Gate by primary EDUCATION anchor if present
                is_primary_anchor = (self._primary_anchors.get('EDUCATION') is None) or (para_idx == self._primary_anchors.get('EDUCATION'))
                
                if is_edu_heading and len(paragraph.text.strip()) < 50 and is_primary_anchor:
                    print(f"  🎓 Found EDUCATION heading at paragraph {para_idx}: '{paragraph.text[:60]}'")
                    
                    education_data = self.resume_data.get('education', [])
                    if not education_data or len(education_data) == 0:
                        print(f"     ⚠️  No education data to insert, skipping")
                    else:
                        print(f"     → Will insert {len(education_data)} education entries after heading")
                        
                        # Look for instructional text in the NEXT paragraph
                        next_para = None
                        is_instruction = False
                        
                        if para_idx + 1 < len(doc.paragraphs):
                            next_para = doc.paragraphs[para_idx + 1]
                            next_text = next_para.text.strip().lower()
                            
                            # Check if next paragraph is instructional text
                            is_instruction = any(phrase in next_text for phrase in [
                                'please list', 'list the candidate', 'please provide',
                                'insert education', 'add education', 'education details',
                                'educational background', 'academic background'
                            ])
                        
                        if is_instruction:
                            print(f"     → Found instructional text: '{next_para.text[:60]}'")
                            
                            education_data = self.resume_data.get('education', [])
                            if education_data and len(education_data) > 0:
                                print(f"     → Will insert {len(education_data)} education entries after heading")
                                
                                # Clear the instructional paragraph
                                for run in next_para.runs:
                                    run.text = ''
                                
                                # CRITICAL: Clear any existing education content after instructional text
                                paras_to_clear = []
                                for check_idx in range(para_idx + 2, min(para_idx + 50, len(doc.paragraphs))):
                                    check_para = doc.paragraphs[check_idx]
                                    check_text = check_para.text.strip().upper()
                                    
                                    # Stop if we hit another section heading
                                    if any(h in check_text for h in ['EMPLOYMENT', 'WORK HISTORY', 'SKILLS', 'SUMMARY', 'CERTIFICATIONS', 'PROJECTS']) and len(check_text) < 50:
                                        print(f"     → Stopped clearing at section: {check_text[:30]}")
                                        break
                                    
                                    # Clear this paragraph (it's old education content)
                                    paras_to_clear.append(check_para)
                                
                                print(f"     → Clearing {len(paras_to_clear)} old content paragraphs")
                                for p in paras_to_clear:
                                    try:
                                        for run in p.runs:
                                            run.text = ''
                                    except:
                                        pass
                                
                                # Insert education blocks after the cleared paragraph
                                last_element = next_para
                                simple_entries = [e for e in education_data if not (e.get('institution') or (e.get('details') or []))]
                                if simple_entries and len(simple_entries) == len(education_data):
                                    self._insert_education_bullets(doc, next_para, education_data[:5])
                                else:
                                    for edu in education_data[:5]:
                                        block = self._insert_education_block(doc, last_element, edu)
                                        if block:
                                            last_element = block
                                
                                self._education_inserted = True
                                print(f"  ✅ Inserted education data after EDUCATION heading")
                                replaced_count += 1
                        else:
                            # No instructional text - there's existing education content
                            # Clear all content after heading until next section
                            print(f"     → No instructional text found, clearing existing education content")
                            
                            paras_to_clear = []
                            for check_idx in range(para_idx + 1, min(para_idx + 50, len(doc.paragraphs))):
                                check_para = doc.paragraphs[check_idx]
                                check_text = check_para.text.strip().upper()
                                
                                # Stop if we hit another section heading or end of document
                                if any(h in check_text for h in ['EMPLOYMENT', 'WORK HISTORY', 'SKILLS', 'SUMMARY', 'CERTIFICATIONS', 'PROJECTS']) and len(check_text) < 50:
                                    print(f"     → Stopped clearing at section: {check_text[:30]}")
                                    break
                                
                                # Clear this paragraph (it's old education content)
                                paras_to_clear.append(check_para)
                            
                            print(f"     → Clearing {len(paras_to_clear)} old education paragraphs")
                            for p in paras_to_clear:
                                try:
                                    for run in p.runs:
                                        run.text = ''
                                except:
                                    pass
                            
                            # Insert education blocks after the heading
                            last_element = paragraph
                            simple_entries = [e for e in education_data if not (e.get('institution') or (e.get('details') or []))]
                            if simple_entries and len(simple_entries) == len(education_data):
                                self._insert_education_bullets(doc, paragraph, education_data[:5])
                            else:
                                for edu in education_data[:5]:
                                    block = self._insert_education_block(doc, last_element, edu)
                                    if block:
                                        last_element = block
                            
                            self._education_inserted = True
                            print(f"  ✅ Inserted education data after EDUCATION heading (no instruction text)")
                            replaced_count += 1
            
            # Education placeholder generic patterns - very flexible matching
            if not self._education_inserted:
                edu_patterns = [
                r"<[^>]*list[^>]*candidate['’]?s?[^>]*education[^>]*background[^>]*>",
                r"<[^>]*education[^>]*background[^>]*>",
                r"<[^>]*education[^>]*history[^>]*>",
                r"<[^>]*candidate['’]?s?[^>]*education[^>]*>",
                r"<[^>]*educational[^>]*background[^>]*>",
                r"<[^>]*academic[^>]*background[^>]*>",
                r"<[^>]*academic[^>]*qualifications[^>]*>",
                r"<[^>]*qualifications[^>]*>",
                r"<[^>]*(education|academic)[^>]*>",
                r"\blist\s*candidate(?:['’]s)?\s*education\s*background\b",
            ]
            for edu_pat in edu_patterns:
                if re.search(edu_pat, paragraph.text, re.IGNORECASE):
                    # If we detected a primary EDUCATION anchor, avoid replacing placeholders that
                    # are far away from the anchor (prevents inserting inside EMPLOYMENT region).
                    if self._primary_anchors.get('EDUCATION') is not None:
                        anchor_idx = self._primary_anchors.get('EDUCATION')
                        # Only allow placeholder replacement near or after the anchor
                        if para_idx < anchor_idx - 2:
                            print(f"  ⏭️  Skipping education placeholder at {para_idx} (before primary anchor {anchor_idx})")
                            continue
                    print(f"  🎓 Found education placeholder in paragraph {para_idx}: '{paragraph.text[:60]}'")
                    
                    # Use structured education data (not sections)
                    education_data = self.resume_data.get('education', [])
                    
                    if education_data and len(education_data) > 0:
                        print(f"     → Will replace with {len(education_data)} education entries")
                        
                        # Clear the placeholder paragraph
                        self._regex_replace_paragraph(paragraph, edu_pat, '')
                        
                        # If entries are simple (no institution/details), insert as bullets to match layout
                        simple_entries = [e for e in education_data if not (e.get('institution') or (e.get('details') or []))]
                        if simple_entries and len(simple_entries) == len(education_data):
                            self._insert_education_bullets(doc, paragraph, education_data[:5])
                        else:
                            # Insert properly formatted education blocks
                            last_element = paragraph
                            for edu in education_data[:5]:  # Limit to 5 entries
                                block = self._insert_education_block(doc, last_element, edu)
                                if block:
                                    last_element = block
                        
                        # CRITICAL: Set flag to prevent duplicate insertion in _add_sections_content
                        self._education_inserted = True
                        
                        print(f"  ✅ Replaced education placeholder with structured blocks")
                        replaced_count += 1
                        break
                    else:
                        # Fallback: try to use sections data
                        content = self._find_matching_resume_section('education', self.resume_data.get('sections', {}))
                        if content:
                            bullets = []
                            for item in content[:10]:
                                if item.strip():
                                    bullets.append('• ' + item.strip().lstrip('•').strip())
                            self._regex_replace_paragraph(paragraph, edu_pat, '\n'.join(bullets))
                            # Set flag even in fallback to prevent duplication
                            self._education_inserted = True
                            print(f"  ✅ Regex replaced education placeholder (fallback)")
                            replaced_count += 1
                            break
        
        print(f"\n Replaced {replaced_count} placeholders in paragraphs")
        
        # Fallback creation: SUMMARY (only if template has no SUMMARY heading)
        if not self._summary_inserted:
            if self._primary_anchors.get('SUMMARY') is not None:
                print(f"  SUMMARY heading exists in template; skipping fallback creation")
            else:
                summary_lines = self._find_matching_resume_section('summary', self.resume_data.get('sections', {})) or []
                summary_text = (self.resume_data.get('summary') or '').strip()
                if summary_lines or summary_text:
                    print(f"  Creating SUMMARY section after candidate name")
                    
                    # Use the pre-computed name anchor (avoids CAI CONTACT section)
                    anchor_idx = getattr(self, '_name_anchor_idx', None)
                    
                    # Fallback: search for name placeholder patterns if anchor wasn't found
                    if anchor_idx is None:
                        print(f"  Name anchor not found, searching for placeholder...")
                        name_placeholder_patterns = [
                            r'<\s*[Cc]andidate[^>]*[Nn]ame[^>]*>',
                            r'<\s*[Nn]ame\s*>',
                            r'<\s*[Ff]ull\s*[Nn]ame\s*>',
                            r'<\s*YOUR\s*NAME\s*>',
                            r'<[^>]*LAWSON[^>]*>',
                            r'<[^>]*PAULA[^>]*>',
                        ]
                        for idx, p in enumerate(doc.paragraphs):
                            if idx < 5:  # Skip CAI CONTACT area
                                continue
                            t = (p.text or '').strip()
                            for pat in name_placeholder_patterns:
                                if re.search(pat, t):
                                    anchor_idx = idx
                                    print(f"  Found name placeholder at paragraph {idx}: '{t}'")
                                    break
                            if anchor_idx is not None:
                                break
                
                    # Strategy: Place SUMMARY right after the candidate name placeholder
                    if anchor_idx is not None and anchor_idx >= 5 and anchor_idx < len(doc.paragraphs):
                        # Name anchor found in main content area (not CAI CONTACT)
                        anchor_para = doc.paragraphs[anchor_idx]
                        print(f"  Inserting SUMMARY after candidate name at paragraph {anchor_idx}")
                    else:
                        # Fallback: use paragraph before EMPLOYMENT if name not found
                        emp_idx = self._primary_anchors.get('EMPLOYMENT')
                        if emp_idx is not None and emp_idx > 0:
                            anchor_para = doc.paragraphs[emp_idx - 1]
                            print(f"  Fallback: Inserting SUMMARY before EMPLOYMENT at paragraph {emp_idx - 1}")
                        else:
                            # Skip SUMMARY insertion if no safe anchor found
                            print(f"  No safe anchor found for SUMMARY - skipping insertion")
                            self._summary_inserted = True  # Mark as inserted to prevent retry
                            anchor_para = None
                
                    if anchor_para is None:
                        # Don't insert if no safe location found
                        pass
                    else:
                        # Insert blank line for spacing
                        blank = self._insert_paragraph_after(anchor_para, '')
                        if blank is None:
                            blank = anchor_para
                        
                        # Insert SUMMARY heading with underline
                        heading = self._insert_paragraph_after(blank, 'SUMMARY')
                        if heading is None:
                            heading = anchor_para
                        for r in heading.runs:
                            r.bold = True
                            r.underline = True
                            r.font.size = Pt(11)
                        
                        # Insert summary content below the heading
                        if summary_text:
                            # Single paragraph summary
                            spara = self._insert_paragraph_after(heading, summary_text)
                            if spara:
                                for r in spara.runs:
                                    r.font.size = Pt(10)
                                print(f"  Inserted SUMMARY heading + paragraph")
                        elif summary_lines:
                            # Multiple lines - insert as bullets
                            self._insert_skills_bullets(doc, heading, summary_lines)
                            print(f"  Inserted SUMMARY heading + bullets")
                        
                        self._summary_inserted = True

        # Fallback creation: SKILLS (place AFTER EMPLOYMENT)
        if not self._skills_inserted:
            skills_list = self.resume_data.get('skills', []) or []
            if skills_list:
                # Anchor: after the last EMPLOYMENT paragraph if available; else after EMPLOYMENT heading; else end
                if hasattr(self, '_employment_tail_para') and self._employment_tail_para is not None:
                    anchor_para = self._employment_tail_para
                elif self._primary_anchors.get('EMPLOYMENT') is not None:
                    anchor_para = doc.paragraphs[self._primary_anchors.get('EMPLOYMENT')]
                else:
                    anchor_para = doc.paragraphs[-1] if doc.paragraphs else doc.add_paragraph('')
                # Clean up any stray SKILLS headings that occur before anchor
                try:
                    anchor_idx = None
                    for idx, p in enumerate(doc.paragraphs):
                        if p is anchor_para:
                            anchor_idx = idx
                            break
                    if anchor_idx is not None:
                        for idx in range(anchor_idx - 1, -1, -1):
                            t = (doc.paragraphs[idx].text or '').strip().upper()
                            if t in ('SKILLS', 'TECHNICAL SKILLS'):
                                self._delete_paragraph(doc.paragraphs[idx])
                except Exception:
                    pass
                heading = self._insert_paragraph_after(anchor_para, 'SKILLS')
                if heading is None:
                    heading = anchor_para
                for r in heading.runs:
                    r.bold = True
                    r.underline = True
                    r.font.size = Pt(11)
                self._insert_skills_bullets(doc, heading, skills_list)
                self._skills_inserted = True

        # CRITICAL: Final cleanup - remove any orphaned bullets that appear after section headings
        # This handles cases where resume parsing left stray content
        print(f"\n🧹 Final cleanup: Removing orphaned content...")
        sections_found = {}
        for para_idx, paragraph in enumerate(doc.paragraphs):
            para_text = paragraph.text.strip().upper()
            
            # Track section positions
            if any(h in para_text for h in ['EMPLOYMENT HISTORY', 'EDUCATION', 'SKILLS', 'SUMMARY']) and len(para_text) < 50:
                sections_found[para_text[:20]] = para_idx
        
        # If we inserted employment and education, clear any bullets that appear after education
        if self._experience_inserted and self._education_inserted and 'EDUCATION' in str(sections_found):
            print(f"     → Checking for orphaned bullets after EDUCATION section...")
            # This is handled by the section clearing logic above
        
        # STEP 3: Process remaining table content (non-skills tables)
        print(f"\n🔍 STEP 3: Processing non-skills table content...")
        other_table_replaced = 0
        
        for table_idx, table in enumerate(doc.tables):
            # Skip skills tables (already processed in STEP 1)
            if self._is_skills_table(table):
                print(f"  ⏭️  Skipping already-processed skills table at index {table_idx}")
                continue
            else:
                # Regular placeholder replacement in non-skills tables
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            # 1) Simple replacements
                            for key, value in replacements.items():
                                if self._text_contains(paragraph.text, key):
                                    other_table_replaced += self._replace_in_paragraph(paragraph, key, value)

                            # 1.5) EDUCATION heading inside table (check if heading OR placeholder exists)
                            if not self._education_inserted:
                                para_text = (paragraph.text or '').strip()
                                heading_text = para_text.upper()
                                
                                # Check if this cell contains EDUCATION heading or placeholder
                                has_heading = any(h in heading_text for h in ['EDUCATION', 'ACADEMIC BACKGROUND', 'EDUCATIONAL BACKGROUND', 'ACADEMIC QUALIFICATIONS', 'QUALIFICATIONS', 'EDUCATION BACKGROUND', 'ACADEMICS'])
                                has_placeholder = bool(re.search(r"<[^>]*list[^>]*candidate['']?s?[^>]*education[^>]*background[^>]*>", para_text, re.IGNORECASE))
                                
                                if has_heading or has_placeholder:
                                    print(f"  🎓 Found EDUCATION in TABLE cell (heading={has_heading}, placeholder={has_placeholder})")
                                    print(f"     Cell text: '{para_text[:80]}'")
                                    
                                    education_data = self.resume_data.get('education', [])
                                    print(f"     → resume_data.education has {len(education_data)} entries")
                                    if not education_data:
                                        sect = self._find_matching_resume_section('education', self.resume_data.get('sections', {}))
                                        print(f"     → sections.education has {len(sect) if sect else 0} lines")
                                        if sect:
                                            education_data = self._build_education_from_bullets(sect)
                                            print(f"     → Built {len(education_data)} education entries from sections")
                                    
                                    if education_data:
                                        print(f"     → Will insert {len(education_data)} education entries")
                                        # Clear the entire paragraph text (heading + placeholder)
                                        for run in paragraph.runs:
                                            run.text = ''
                                        # Rewrite just the heading
                                        if paragraph.runs:
                                            paragraph.runs[0].text = 'EDUCATION'
                                            paragraph.runs[0].bold = True
                                            paragraph.runs[0].font.size = Pt(11)
                                        
                                        # Clear any following raw content within the cell
                                        self._delete_following_bullets(paragraph, max_scan=80)
                                        self._delete_next_table(paragraph)
                                        
                                        # Insert content
                                        simple_entries = [e for e in education_data if not (e.get('institution') or (e.get('details') or []))]
                                        if simple_entries and len(simple_entries) == len(education_data):
                                            self._insert_education_bullets(doc, paragraph, education_data[:5])
                                            print(f"     ✅ Inserted {len(education_data[:5])} education bullets")
                                        else:
                                            last_element = paragraph
                                            for edu in education_data[:5]:
                                                blk = self._insert_education_block(doc, last_element, edu)
                                                if blk:
                                                    last_element = blk
                                            print(f"     ✅ Inserted {len(education_data[:5])} education blocks")
                                        
                                        self._education_inserted = True
                                        other_table_replaced += 1
                                    else:
                                        print(f"     ⚠️  No education data available to insert")
                                        # Still clear the placeholder even if no data
                                        for run in paragraph.runs:
                                            run.text = ''
                                        if paragraph.runs:
                                            paragraph.runs[0].text = 'EDUCATION'
                                            paragraph.runs[0].bold = True
                                        self._education_inserted = True

                            # 2) SUMMARY placeholder inside table
                            if not self._summary_inserted:
                                summary_patterns = [r"<[^>]*summary[^>]*>", r"<[^>]*professional[^>]*summary[^>]*>", r"<[^>]*profile[^>]*>"]
                                for sum_pat in summary_patterns:
                                    if re.search(sum_pat, paragraph.text, re.IGNORECASE):
                                        summary_lines = self._find_matching_resume_section('summary', self.resume_data.get('sections', {}))
                                        summary_text = (self.resume_data.get('summary') or '').strip()
                                        if summary_lines or summary_text:
                                            print(f"  📝 Found summary placeholder in TABLE cell")
                                            self._regex_replace_paragraph(paragraph, sum_pat, '')
                                            if summary_lines:
                                                self._insert_skills_bullets(doc, paragraph, summary_lines)
                                            else:
                                                sp = self._insert_paragraph_after(paragraph, summary_text)
                                                if sp:
                                                    for run in sp.runs:
                                                        run.font.size = Pt(10)
                                            self._summary_inserted = True
                                            table_replaced += 1
                                            break

                            # 3) SKILLS placeholder inside table
                            if not self._skills_inserted:
                                skills_patterns = [r"<[^>]*skills[^>]*>", r"<[^>]*technical[^>]*skills[^>]*>", r"<[^>]*list[^>]*skills[^>]*>"]
                                for skl_pat in skills_patterns:
                                    if re.search(skl_pat, paragraph.text, re.IGNORECASE):
                                        skills_list = self.resume_data.get('skills', [])
                                        if skills_list:
                                            print(f"  🧰 Found skills placeholder in TABLE cell")
                                            self._regex_replace_paragraph(paragraph, skl_pat, '')
                                            self._insert_skills_bullets(doc, paragraph, skills_list)
                                            self._skills_inserted = True
                                            other_table_replaced += 1
                                        else:
                                            self._regex_replace_paragraph(paragraph, skl_pat, '')
                                        break

                            # 4) EDUCATION placeholder inside table
                            if not self._education_inserted:
                                edu_patterns = [
                                    r"<[^>]*list[^>]*candidate['’]?s?[^>]*education[^>]*background[^>]*>",
                                    r"<[^>]*education[^>]*background[^>]*>",
                                    r"<[^>]*education[^>]*history[^>]*>",
                                    r"<[^>]*candidate['’]?s?[^>]*education[^>]*>",
                                    r"<[^>]*educational[^>]*background[^>]*>",
                                    r"<[^>]*academic[^>]*background[^>]*>",
                                    r"<[^>]*academic[^>]*qualifications[^>]*>",
                                    r"<[^>]*qualifications[^>]*>",
                                    r"<[^>]*(education|academic)[^>]*>",
                                ]
                                for edu_pat in edu_patterns:
                                    if re.search(edu_pat, paragraph.text, re.IGNORECASE):
                                        print(f"  🎓 Found education placeholder in TABLE cell")
                                        education_data = self.resume_data.get('education', [])
                                        if not education_data:
                                            sect = self._find_matching_resume_section('education', self.resume_data.get('sections', {}))
                                            if sect:
                                                education_data = self._build_education_from_bullets(sect)
                                        if education_data:
                                            self._regex_replace_paragraph(paragraph, edu_pat, '')
                                            simple_entries = [e for e in education_data if not (e.get('institution') or (e.get('details') or []))]
                                            if simple_entries and len(simple_entries) == len(education_data):
                                                self._insert_education_bullets(doc, paragraph, education_data[:5])
                                            else:
                                                last_element = paragraph
                                                for edu in education_data[:5]:
                                                    blk = self._insert_education_block(doc, last_element, edu)
                                                    if blk:
                                                        last_element = blk
                                            self._education_inserted = True
                                            other_table_replaced += 1
                                            break
        
        print(f"✓ Processed {other_table_replaced} non-skills table entries")
        print(f"\n{'='*70}")
        print(f"📊 FORMATTING SUMMARY")
        print(f"{'='*70}")
        print(f"  • Skills table entries: {table_replaced}")
        print(f"  • Paragraph replacements: {replaced_count}")
        print(f"  • Other table entries: {other_table_replaced}")
        
        # Replace in headers/footers
        header_footer_replaced = 0
        for section in doc.sections:
            # Header
            for paragraph in section.header.paragraphs:
                for key, value in replacements.items():
                    if self._text_contains(paragraph.text, key):
                        header_footer_replaced += self._replace_in_paragraph(paragraph, key, value)
            
            # Footer
            for paragraph in section.footer.paragraphs:
                for key, value in replacements.items():
                    if self._text_contains(paragraph.text, key):
                        header_footer_replaced += self._replace_in_paragraph(paragraph, key, value)
        
        if header_footer_replaced > 0:
            print(f"✓ Replaced {header_footer_replaced} placeholders in headers/footers")
        
        # Add sections content
        sections_added = self._add_sections_content(doc)
        print(f"✓ Added {sections_added} sections")
        
        # Final sweep: clear leftover instructional phrases globally
        self._clear_instruction_phrases(doc)

        # Save output
        output_docx = self.output_path.replace('.pdf', '.docx')
        doc.save(output_docx)
        
        # Post-process with Word COM to handle placeholders inside shapes/text boxes
        if HAS_WIN32:
            try:
                self._postprocess_with_word_com(output_docx)
            except Exception as e:
                print(f"⚠️  COM post-processing skipped due to error: {e}")
        else:
            print("ℹ️  COM post-processing unavailable (pywin32 not installed)")

        print(f"\n✅ Successfully created formatted document!")
        print(f"📁 Saved to: {output_docx}\n")
        
        # Optionally convert to PDF
        if self.output_path.endswith('.pdf'):
            print("📄 Converting to PDF...")
            if self._convert_to_pdf(output_docx, self.output_path):
                print(f"✓ PDF created: {self.output_path}")
                # Keep both docx and pdf
            else:
                print("⚠️  PDF conversion failed, keeping .docx file")
        
        return True

    # Helper: insert a new paragraph directly after a given paragraph
    def _insert_paragraph_after(self, paragraph, text):
        try:
            new_p = OxmlElement('w:p')
            paragraph._p.addnext(new_p)
            new_para = Paragraph(new_p, paragraph._parent)
            new_para.add_run(text)
            return new_para
        except Exception:
            # Fallback: append to document if direct insert fails
            return paragraph._parent.add_paragraph(text)
    
    def _add_right_tab(self, paragraph, pos_twips=9360):
        """Add a right-aligned tab stop to a paragraph at the given twips position (1 inch = 1440 twips)."""
        try:
            pPr = paragraph._p.get_or_add_pPr()
            tabs = pPr.find(qn('w:tabs'))
            if tabs is None:
                tabs = OxmlElement('w:tabs')
                pPr.append(tabs)
            tab = OxmlElement('w:tab')
            tab.set(qn('w:val'), 'right')
            tab.set(qn('w:pos'), str(pos_twips))
            tabs.append(tab)
        except Exception:
            # If this fails, the text will still render; right text just won't align via tab stop
            pass

    def _delete_paragraph(self, paragraph):
        """Safely delete a paragraph from the document body."""
        try:
            p = paragraph._element
            parent = p.getparent()
            parent.remove(p)
        except Exception:
            pass

    def _insert_skills_bullets(self, doc, after_paragraph, skills_list):
        """Insert a simple bullet list of skills after the given paragraph.
        Accepts list of strings or objects with 'name' field.
        Returns the last inserted paragraph.
        """
        try:
            last = after_paragraph
            count = 0
            # Limit to 15 skills for readability
            for skill in (skills_list or [])[:15]:
                name = skill if isinstance(skill, str) else (skill.get('name', '') if isinstance(skill, dict) else str(skill))
                name = (name or '').strip()
                if not name:
                    continue
                p = self._insert_paragraph_after(last, '')
                if p is None:
                    break
                p.paragraph_format.left_indent = Inches(0.25)
                run = p.add_run('• ' + name.lstrip('•–—-*● '))
                run.font.size = Pt(10)
                p.paragraph_format.space_after = Pt(2)
                last = p
                count += 1
            if count > 0:
                print(f"    ✓ Inserted {count} skill bullets")
            return last
        except Exception as e:
            print(f"    ⚠️  Error inserting skills bullets: {e}")
            return after_paragraph

    def _insert_education_bullets(self, doc, after_paragraph, education_list):
        """Insert education as bullets like 'High School Graduation 1986.'
        Used when entries have no institution/details to match desired layout.
        """
        try:
            last = after_paragraph
            count = 0
            for edu in (education_list or [])[:5]:
                degree = (edu.get('degree') or '').strip()
                year = self._clean_duration((edu.get('year') or '').strip())
                inst = (edu.get('institution') or '').strip()
                details = edu.get('details') or []
                # Only bullet-render when institution and details are empty (simple case)
                if inst or (details and len(details) > 0):
                    continue
                text_parts = []
                if degree:
                    text_parts.append(degree)
                if year:
                    # Append year after a space to match example
                    text_parts.append(year)
                line = ' '.join(text_parts).strip()
                if not line:
                    continue
                p = self._insert_paragraph_after(last, '')
                if p is None:
                    break
                p.paragraph_format.left_indent = Inches(0.25)
                run = p.add_run('• ' + line)
                run.font.size = Pt(10)
                p.paragraph_format.space_after = Pt(2)
                last = p
                count += 1
            if count > 0:
                print(f"    ✓ Inserted {count} education bullets (simple mode)")
            return last
        except Exception as e:
            print(f"    ⚠️  Error inserting education bullets: {e}")
            return after_paragraph

    # ===== CAI CONTACT PERSISTENCE AND INSERTION =====
    def _cai_store_path(self):
        """Return a stable file path to store CAI contact details."""
        home = os.path.expanduser("~")
        return os.path.join(home, ".resume_formatter_cai_contact.json")

    def _load_cai_contact(self, proposed=None, edit=False):
        """Load CAI contact from disk. If edit=True and proposed provided, overwrite and save.
        Structure: {"name": str, "phone": str, "email": str}
        """
        path = self._cai_store_path()
        stored = {}
        try:
            if os.path.exists(path):
                with open(path, 'r', encoding='utf-8') as f:
                    stored = json.load(f) or {}
        except Exception:
            stored = {}

        if edit and isinstance(proposed, dict) and any(proposed.get(k) for k in ("name", "phone", "email")):
            data = {
                "name": (proposed.get("name") or stored.get("name") or ""),
                "phone": (proposed.get("phone") or stored.get("phone") or ""),
                "email": (proposed.get("email") or stored.get("email") or ""),
            }
            try:
                with open(path, 'w', encoding='utf-8') as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)
            except Exception:
                pass
            return data

        # No edit: fall back to stored, else proposed, else empty
        if stored:
            return stored
        if isinstance(proposed, dict):
            return {
                "name": proposed.get("name", ""),
                "phone": proposed.get("phone", ""),
                "email": proposed.get("email", ""),
            }
        return {"name": "", "phone": "", "email": ""}

    def _ensure_cai_contact(self, doc):
        """Ensure the CAI CONTACT section exists and is filled from persistent storage.
        Will not change stored values unless an explicit edit flag is provided via
        resume_data['edit_cai_contact'] or template_analysis['edit_cai_contact'].
        """
        edit_flag = bool(self.resume_data.get('edit_cai_contact') or (self.template_analysis or {}).get('edit_cai_contact'))
        proposed = self.resume_data.get('cai_contact', {})
        cai = self._load_cai_contact(proposed=proposed, edit=edit_flag)

        # Find existing CAI CONTACT heading
        heading_idx = None
        for idx, p in enumerate(doc.paragraphs):
            if (p.text or '').strip().upper() == 'CAI CONTACT':
                heading_idx = idx
                break

        if heading_idx is None:
            # Do NOT create CAI CONTACT unless explicitly requested via edit flag
            if not edit_flag:
                print("  ⏭️  No 'CAI CONTACT' heading in template; skipping CAI contact insertion")
                return
            # Explicitly requested: create near the top
            anchor = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph("")
            heading = self._insert_paragraph_after(anchor, 'CAI CONTACT')
            if heading is None:
                heading = doc.add_paragraph('CAI CONTACT')
            for r in heading.runs:
                r.bold = True
                r.font.size = Pt(11)
            # Write lines under heading
            self._write_cai_contact_block(heading, cai)
        else:
            # CAI CONTACT heading exists in template
            heading = doc.paragraphs[heading_idx]
            
            if not edit_flag:
                # No edit requested - leave template CAI CONTACT section completely unchanged
                print("  ⏭️  CAI CONTACT exists in template but edit_flag not set; leaving unchanged")
                return
            
            # Edit flag is set - ADD new contact info BELOW existing template content
            print("  ✏️  Edit CAI Contact enabled - adding new contact info below template defaults")
            
            # Find the last paragraph of the CAI CONTACT section
            # Scan only within CAI CONTACT - stop at empty line or candidate name placeholder
            last_cai_para = heading
            for j in range(1, 15):  # Scan up to 15 paragraphs after heading
                k = heading_idx + j
                if k >= len(doc.paragraphs):
                    break
                txt = (doc.paragraphs[k].text or '').strip()
                upper = txt.upper()
                
                # Stop at candidate name placeholder (indicates start of main content)
                if '<' in txt and '>' in txt and any(word in upper for word in ['NAME', 'CANDIDATE', 'PAULA', 'LAWSON']):
                    print(f"  🛑 Stopped at candidate name placeholder at para {k}")
                    break
                
                # Stop at next major section headings
                if any(kw in upper for kw in ['EMPLOYMENT HISTORY', 'WORK EXPERIENCE', 'EDUCATION', 'SUMMARY', 'SKILLS']) and len(txt) < 50:
                    print(f"  🛑 Stopped at section heading '{txt}' at para {k}")
                    break
                
                # Stop at multiple consecutive empty lines (indicates section break)
                if not txt:
                    # Check if next line is also empty or a major section
                    if k + 1 < len(doc.paragraphs):
                        next_txt = (doc.paragraphs[k + 1].text or '').strip()
                        if not next_txt or any(kw in next_txt.upper() for kw in ['EMPLOYMENT', 'SUMMARY', 'EDUCATION']):
                            print(f"  🛑 Stopped at empty line break at para {k}")
                            break
                
                # This is still part of CAI CONTACT section
                last_cai_para = doc.paragraphs[k]
            
            # Insert new contact block after the last CAI CONTACT paragraph
            print(f"  📍 Inserting edited CAI contact after paragraph {heading_idx + (doc.paragraphs.index(last_cai_para) - heading_idx if last_cai_para in doc.paragraphs else 0)}")
            self._write_cai_contact_block(last_cai_para, cai)

    def _write_cai_contact_block(self, heading_para, cai):
        """Write CAI contact lines under the given heading paragraph."""
        name = (cai.get('name') or '').strip()
        phone = (cai.get('phone') or '').strip()
        email = (cai.get('email') or '').strip()

        p_name = self._insert_paragraph_after(heading_para, name or '')
        if p_name is not None:
            for r in p_name.runs:
                r.bold = True
                r.font.size = Pt(10)

        if phone:
            p_phone = self._insert_paragraph_after(p_name or heading_para, f"Phone:  {phone}")
            if p_phone is not None:
                for r in p_phone.runs:
                    r.font.size = Pt(10)

        if email:
            p_email = self._insert_paragraph_after(p_phone or p_name or heading_para, f"Email:  {email}")
            if p_email is not None:
                for r in p_email.runs:
                    r.font.size = Pt(10)
    
    def _insert_experience_block(self, doc, after_paragraph, exp_data):
        """Insert a structured 2-column experience block"""
        try:
            # Get parsed data from resume parser
            company = exp_data.get('company', '')
            role = exp_data.get('role', '')
            duration = exp_data.get('duration', '')
            details = exp_data.get('details', [])
            
            # Fallback: if company/role not parsed, try to extract from title
            if not company and not role:
                title = exp_data.get('title', '')
                company, role = self._parse_company_role(title)
            
            # CRITICAL: Remove date fragments from company/role
            # Sometimes dates like "City – 08/ 06/" end up in company field
            import re
            if company:
                # Remove patterns like "City – 08/ 06/" or "– 04/" etc
                company = re.sub(r'\s*[–-]\s*\d{2}/\s*\d{2}/?\s*.*?$', '', company)
                company = re.sub(r'\s*City\s*[–-]\s*\d{2}/.*?$', '', company)
                company = company.strip(' ,–-')
            
            if role:
                role = re.sub(r'\s*[–-]\s*\d{2}/\s*\d{2}/?/\s*.*?$', '', role)
                role = re.sub(r'\s*City\s*[–-]\s*\d{2}/.*?$', '', role)
                role = role.strip(' ,–-')
            
            # Clean up duration format
            duration_clean = self._clean_duration(duration)
            
            # Build header line using a right-aligned tab instead of a table
            header_para = self._insert_paragraph_after(after_paragraph, '')
            # Right tab at ~6.5" (letter page width minus 1" margins), 1 inch = 1440 twips
            self._add_right_tab(header_para, pos_twips=9360)

            # CRITICAL: Company name ALWAYS goes on first line if available
            # Role goes on second line if both exist
            # If only one exists, it goes on first line
            
            if company and role:
                # CASE 1: Both company and role exist
                # Line 1: Company (bold) + dates
                # Truncate if too long to prevent date wrapping
                display_company = company
                if len(display_company) > 70:
                    display_company = display_company[:67] + '...'
                
                left_run = header_para.add_run(display_company)
                left_run.bold = True
                left_run.font.size = Pt(10)
                
                if duration_clean:
                    header_para.add_run('\t')
                    dur_run = header_para.add_run(duration_clean)
                    dur_run.bold = False
                    dur_run.font.size = Pt(9)
                header_para.paragraph_format.space_after = Pt(0)
                header_para.paragraph_format.keep_together = True
                
                # Line 2: Role (bold)
                role_para = self._insert_paragraph_after(header_para, '')
                role_run = role_para.add_run(role)
                role_run.bold = True
                role_run.font.size = Pt(10)
                role_para.paragraph_format.space_after = Pt(0)
                last_para = role_para
                
            elif company:
                # CASE 2: Only company exists
                # Truncate if too long to prevent date wrapping
                display_company = company
                if len(display_company) > 70:
                    display_company = display_company[:67] + '...'
                
                left_run = header_para.add_run(display_company)
                left_run.bold = True
                left_run.font.size = Pt(10)
                
                if duration_clean:
                    header_para.add_run('\t')
                    dur_run = header_para.add_run(duration_clean)
                    dur_run.bold = False
                    dur_run.font.size = Pt(9)
                header_para.paragraph_format.space_after = Pt(0)
                header_para.paragraph_format.keep_together = True
                last_para = header_para
                
            elif role:
                # CASE 3: Only role exists
                # Truncate if too long to prevent date wrapping
                display_role = role
                if len(display_role) > 70:
                    display_role = display_role[:67] + '...'
                
                left_run = header_para.add_run(display_role)
                left_run.bold = True
                left_run.font.size = Pt(10)
                
                if duration_clean:
                    header_para.add_run('\t')
                    dur_run = header_para.add_run(duration_clean)
                    dur_run.bold = False
                    dur_run.font.size = Pt(9)
                header_para.paragraph_format.space_after = Pt(0)
                header_para.paragraph_format.keep_together = True
                last_para = header_para
                
            else:
                # CASE 4: Neither exists (fallback)
                left_run = header_para.add_run('Experience')
                left_run.bold = True
                left_run.font.size = Pt(10)
                
                if duration_clean:
                    header_para.add_run('\t')
                    dur_run = header_para.add_run(duration_clean)
                    dur_run.bold = False
                    dur_run.font.size = Pt(9)
                header_para.paragraph_format.space_after = Pt(0)
                last_para = header_para

            # Add details as individual bullet paragraphs
            if details:
                # Don't limit bullets - include ALL details from resume
                for detail in details:
                    txt = (detail or '').strip()
                    if not txt:
                        continue
                    p = self._insert_paragraph_after(last_para, '')
                    p.paragraph_format.left_indent = Inches(0.25)
                    run = p.add_run('• ' + txt.lstrip('•–—-*● '))
                    run.font.size = Pt(9)
                    last_para = p
            
            return last_para
            
        except Exception as e:
            print(f"  ⚠️  Error inserting experience block: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def _insert_education_block(self, doc, after_paragraph, edu_data):
        """Insert a structured 2-column education block"""
        try:
            # Get parsed data from resume parser
            degree = edu_data.get('degree', '').strip()
            institution = edu_data.get('institution', '').strip()
            year = edu_data.get('year', '').strip()
            details = edu_data.get('details', [])
            
            # DEBUG: Show what we received
            print(f"      📚 Education data: degree='{degree[:50] if degree else 'EMPTY'}', institution='{institution[:30] if institution else 'EMPTY'}', year='{year or 'EMPTY'}'")
            
            # CRITICAL: If we have neither degree nor institution, skip this entry
            if not degree and not institution:
                print(f"      ⚠️  Skipping empty education entry")
                return after_paragraph
            
            # Fallback: if institution not parsed, try to extract from degree or details
            if not institution:
                institution = self._extract_institution(degree, details)
                if institution:
                    print(f"      ✓ Extracted institution from degree: '{institution[:30]}'")
            
            # Clean up year format
            year_clean = self._clean_duration(year)
            
            # Build header line as paragraph with right-aligned tab (no tables)
            header_para = self._insert_paragraph_after(after_paragraph, '')
            if header_para is None:
                print(f"      ❌ Failed to insert paragraph")
                return after_paragraph
                
            self._add_right_tab(header_para, pos_twips=9360)
            
            # Parse degree to separate degree type from field
            # Handle both formats:
            # 1. "Master of Science : Leadership" (colon separator)
            # 2. "Master of Science in Data Science" (in separator)
            # 3. "Bachelor of Technology in Computer Science" (in separator)
            
            degree_type = degree
            field_and_institution = institution
            field = ''
            
            if ':' in degree:
                # Format: "Master of Science : Leadership"
                parts = degree.split(':', 1)
                degree_type = parts[0].strip()  # "Master of Science"
                field = parts[1].strip() if len(parts) > 1 else ''  # "Leadership"
                print(f"      ✂️  Split at colon: LEFT='{degree_type}' | Field='{field}'")
            
            elif ' in ' in degree.lower():
                # Format: "Master of Science in Data Science"
                # Find the position of " in " (case-insensitive)
                lower_degree = degree.lower()
                in_pos = lower_degree.find(' in ')
                if in_pos > 0:
                    degree_type = degree[:in_pos].strip()  # "Master of Science"
                    field = degree[in_pos + 4:].strip()     # "Data Science"
                    print(f"      ✂️  Split at 'in': LEFT='{degree_type}' | Field='{field}'")
            
            else:
                print(f"      ℹ️  No split: Using full degree as LEFT='{degree_type}'")
            
            # Combine field with institution
            if field and institution:
                field_and_institution = f"{field} - {institution}"
            elif field:
                field_and_institution = field
            
            print(f"      📐 Format: LEFT='{degree_type or '(no degree)'}' | RIGHT='{year_clean}'")
            print(f"      📐 Second line: '{field_and_institution or '(none)'}'")
            
            # CRITICAL FIX: Truncate degree_type if too long to prevent date wrapping
            # Max ~70 chars to ensure tab stop works properly (leaves room for date on right)
            display_degree = degree_type or institution or 'Education'
            if len(display_degree) > 70:
                display_degree = display_degree[:67] + '...'
                print(f"      ✂️  Truncated degree to: '{display_degree}'")
            
            # Degree type on the left (bold), year on the right
            deg_run = header_para.add_run(display_degree)
            deg_run.bold = True
            deg_run.font.size = Pt(10)
            if year_clean:
                header_para.add_run('\t')
                yr_run = header_para.add_run(year_clean)
                yr_run.bold = False
                yr_run.font.size = Pt(9)
            header_para.paragraph_format.space_after = Pt(0)
            # Prevent line from breaking
            header_para.paragraph_format.keep_together = True

            # Field + Institution on the next line (normal)
            last_para = header_para
            if field_and_institution:
                fi_para = self._insert_paragraph_after(header_para, field_and_institution)
                if fi_para is not None:
                    for run in fi_para.runs:
                        run.font.size = Pt(10)
                    fi_para.paragraph_format.space_after = Pt(2)
                    last_para = fi_para
                    print(f"      ✅ Inserted field/institution line")
                else:
                    print(f"      ⚠️  Failed to insert field/institution paragraph")

            # Add details as bullet paragraphs
            if details:
                opt_details = self._optimize_details(details, max_bullets=3, max_words=18, max_chars=120)
                detail_count = 0
                for detail in opt_details:
                    txt = (detail or '').strip()
                    if not txt or txt.lower() == (institution or '').lower():
                        continue
                    p = self._insert_paragraph_after(last_para, '')
                    if p is not None:
                        p.paragraph_format.left_indent = Inches(0.25)
                        run = p.add_run('• ' + txt.lstrip('•–—-*● '))
                        run.font.size = Pt(9)
                        p.paragraph_format.space_after = Pt(2)
                        last_para = p
                        detail_count += 1
                if detail_count > 0:
                    print(f"      ✅ Inserted {detail_count} detail bullets")
            
            # Add a blank line after each education entry for spacing
            blank = self._insert_paragraph_after(last_para, '')
            if blank is not None:
                blank.paragraph_format.space_after = Pt(6)
                last_para = blank
            
            print(f"      ✅ Education block inserted successfully")
            return last_para
            
        except Exception as e:
            print(f"  ⚠️  Error inserting education block: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def _clean_duration(self, duration):
        """Normalize duration.
        Prefer 'Mon YYYY-Mon YYYY' (e.g., 'Nov 2011-Sept 2025').
        Fallback to 'YYYY-YYYY' or single 'YYYY'.
        """
        if not duration:
            return ''

        t = (duration or '').strip()
        if not t:
            return ''

        # Normalize connectors
        t = re.sub(r'[–—]', '-', t)
        t = re.sub(r'\s*(to|–|—|-)\s*', '-', t, flags=re.IGNORECASE)

        present = bool(re.search(r'\b(current|present)\b', t, re.IGNORECASE))

        # Month map with 'Sept' spelling
        month_map = {
            'january': 'Jan', 'jan': 'Jan',
            'february': 'Feb', 'feb': 'Feb',
            'march': 'Mar', 'mar': 'Mar',
            'april': 'Apr', 'apr': 'Apr',
            'may': 'May',
            'june': 'Jun', 'jun': 'Jun',
            'july': 'Jul', 'jul': 'Jul',
            'august': 'Aug', 'aug': 'Aug',
            'september': 'Sept', 'sept': 'Sept', 'sep': 'Sept',
            'october': 'Oct', 'oct': 'Oct',
            'november': 'Nov', 'nov': 'Nov',
            'december': 'Dec', 'dec': 'Dec',
        }

        def abbr(m):
            return month_map.get(m.lower(), m[:3].title())

        # Find month-year tokens
        my = [m.groups() for m in re.finditer(r'\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\s*(?:,\s*)?((?:19|20)\d{2})\b', t, flags=re.IGNORECASE)]
        if my:
            start_m, start_y = my[0]
            start_m = abbr(start_m)
            if len(my) >= 2:
                end_m, end_y = my[-1]
                end_m = abbr(end_m)
                return f"{start_m} {start_y}-{end_m} {end_y}"
            # Single month-year; attach Present or end year if available
            if present:
                return f"{start_m} {start_y}-Present"
            end_years = re.findall(r'\b((?:19|20)\d{2})\b', t)
            if len(end_years) >= 2:
                return f"{start_m} {end_years[0]}-{end_years[-1]}"
            return f"{start_m} {start_y}"

        # Fallback to years-only
        years = re.findall(r'\b(?:19|20)\d{2}\b', t)
        if len(years) >= 2:
            return f"{years[0]}-{years[-1]}"
        if len(years) == 1:
            return years[0]
        return ''
    
    def _parse_company_role(self, title):
        """Parse company and role from title line"""
        # Common patterns: "Company Name - Role" or "Role at Company" or "Role, Company"
        if ' - ' in title:
            parts = title.split(' - ', 1)
            return parts[0].strip(), parts[1].strip()
        elif ' at ' in title.lower():
            parts = re.split(r'\s+at\s+', title, flags=re.IGNORECASE)
            return parts[1].strip() if len(parts) > 1 else '', parts[0].strip()
        elif ', ' in title:
            parts = title.split(', ', 1)
            return parts[1].strip(), parts[0].strip()
        else:
            # Assume entire line is company or role
            return title.strip(), ''
    
    def _extract_institution(self, degree, details):
        """Extract institution name from degree line or details"""
        # Check if degree line contains institution (common pattern: "Degree, Institution")
        if ', ' in degree:
            parts = degree.split(', ', 1)
            return parts[1].strip()
        
        # Look in details for institution keywords
        institution_keywords = ['university', 'college', 'institute', 'school', 'academy']
        for detail in details:
            if any(kw in detail.lower() for kw in institution_keywords):
                return detail.strip()
        
        return ''
    
    def _remove_cell_borders(self, cell):
        """Remove all borders from a table cell"""
        try:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                tcBorders.append(border)
            tcPr.append(tcBorders)
        except:
            pass
    
    def _insert_table_after(self, doc, anchor, rows=1, cols=2):
        """Create a table and position it immediately after the given anchor (Paragraph or Table)."""
        try:
            # Simply add table to document - python-docx will handle placement
            table = doc.add_table(rows=rows, cols=cols)
            
            # Try to move it after anchor, but don't fail if it doesn't work
            try:
                tbl = table._element
                if hasattr(anchor, '_element'):
                    anchor_elm = anchor._element
                elif hasattr(anchor, '_tbl'):
                    anchor_elm = anchor._tbl
                else:
                    anchor_elm = anchor
                
                # Only move if we have a valid anchor
                if anchor_elm is not None and hasattr(anchor_elm, 'addnext'):
                    anchor_elm.addnext(tbl)
            except:
                # If moving fails, table will just be at end of document
                pass

            return table

        except Exception as e:
            print(f"  ⚠️  Error creating table: {e}")
            return None

    def _cleanup_duplicate_bullets_after_section(self, doc, section_heading_para, next_section_name):
        """
        AGGRESSIVE cleanup: Scan entire document after inserting formatted content
        and delete ANY remaining bullet points between this section and next section.
        This ensures NO duplication of raw content.
        """
        try:
            print(f"    🧹 AGGRESSIVE cleanup: Removing ALL raw content until '{next_section_name}'...")
            
            # Find the section heading paragraph index
            heading_idx = None
            for idx, para in enumerate(doc.paragraphs):
                if para._element == section_heading_para._element:
                    heading_idx = idx
                    break
            
            if heading_idx is None:
                return
            
            # Now scan from heading to next section and delete EVERYTHING except tables and section headings
            deleted = 0
            paras_to_delete = []
            
            # List of all section keywords to preserve
            section_keywords = ['EDUCATION', 'SKILLS', 'SUMMARY', 'PROJECT', 'CERTIFICATION', 
                              'EXPERIENCE', 'WORK EXPERIENCE', 'EMPLOYMENT HISTORY', 
                              'PROFESSIONAL EXPERIENCE', 'CAREER HISTORY', 'QUALIFICATIONS',
                              'ACHIEVEMENTS', 'AWARDS', 'LANGUAGES']
            
            for idx in range(heading_idx + 1, len(doc.paragraphs)):
                para = doc.paragraphs[idx]
                text = para.text.strip().upper()
                
                # Stop at next section
                if next_section_name in text and len(text) < 50:
                    print(f"       Stopped at next section: {text[:40]}")
                    break
                
                # Skip if paragraph is empty
                if not text:
                    continue
                
                # PRESERVE section headings (don't delete them!)
                is_section_heading = any(keyword in text for keyword in section_keywords) and len(text) < 50
                if is_section_heading:
                    print(f"       Preserved section heading: '{text[:40]}'")
                    continue
                
                # DELETE this paragraph (it's duplicate raw content)
                paras_to_delete.append(para)
                deleted += 1
                if deleted <= 5:
                    print(f"       Removing duplicate: '{text[:60]}'")
            
            # Actually delete the paragraphs
            for para in paras_to_delete:
                p_element = para._element
                p_element.getparent().remove(p_element)
            
            print(f"    🧹 Cleanup complete: Removed {deleted} duplicate paragraphs")
            
        except Exception as e:
            print(f"    ⚠️  Cleanup error: {e}")
    
    def _delete_following_bullets(self, paragraph, max_scan=200):
        """Delete ALL content after a heading until next section - includes TABLES and paragraphs."""
        try:
            body = paragraph._element.getparent()
            node = paragraph._element.getnext()
            deleted_paras = 0
            deleted_tables = 0
            scanned = 0
            
            print(f"    🔍 Starting deletion scan (max {max_scan} items)...")
            
            while node is not None and scanned < max_scan:
                scanned += 1
                next_node = node.getnext()
                
                # DELETE TABLES (raw content might be in tables)
                if node.tag.endswith('tbl'):
                    print(f"       Deleting table #{deleted_tables + 1}")
                    body.remove(node)
                    deleted_tables += 1
                    node = next_node
                    continue
                
                # DELETE PARAGRAPHS
                if node.tag.endswith('p'):
                    # Extract plain text
                    text_nodes = node.xpath('.//w:t', namespaces=node.nsmap) if hasattr(node, 'xpath') else []
                    text = ''.join([t.text for t in text_nodes if t is not None and t.text is not None])
                    txt = (text or '').strip()
                    norm = txt.upper()
                    
                    # Stop ONLY at next section heading (not at tables or bullets)
                    section_keywords = ['EDUCATION', 'SKILLS', 'SUMMARY', 'PROJECT', 'CERTIFICATION', 
                                      'EXPERIENCE', 'WORK EXPERIENCE', 'EMPLOYMENT HISTORY', 
                                      'PROFESSIONAL EXPERIENCE', 'CAREER HISTORY', 'QUALIFICATIONS']
                    
                    if any(k in norm for k in section_keywords) and len(txt) < 50:
                        # This looks like a next section heading, stop deleting
                        print(f"       Stopped at next section: '{txt[:40]}'")
                        break
                    
                    # DELETE THIS PARAGRAPH (raw content)
                    if deleted_paras < 5:  # Log first 5 deletions
                        print(f"       Deleting para: '{txt[:60]}'...")
                    body.remove(node)
                    deleted_paras += 1
                
                node = next_node
            
            print(f"    🗑️  DELETED: {deleted_paras} paragraphs + {deleted_tables} tables (scanned {scanned} items)")
            
            if deleted_paras == 0 and deleted_tables == 0:
                print(f"    ⚠️  WARNING: Nothing was deleted! Content might still be there.")
                
        except Exception as e:
            print(f"    ⚠️  Error deleting content: {e}")
            import traceback
            traceback.print_exc()
    
    def _collect_bullets_after_heading(self, paragraph, max_scan=50):
        """Collect consecutive bullet-like paragraphs immediately after a heading/placeholder."""
        bullets = []
        try:
            node = paragraph._element.getnext()
            scanned = 0
            while node is not None and scanned < max_scan:
                scanned += 1
                if node.tag.endswith('tbl'):
                    # Collect all paragraph texts from the table as bullets
                    paras = node.xpath('.//w:p', namespaces=node.nsmap) if hasattr(node, 'xpath') else []
                    for p in paras:
                        tnodes = p.xpath('.//w:t', namespaces=p.nsmap) if hasattr(p, 'xpath') else []
                        text = ''.join([t.text for t in tnodes if t is not None and t.text is not None]).strip()
                        if text:
                            bullets.append(text.lstrip(' •–—-*●'))
                    break
                if node.tag.endswith('p'):
                    text_nodes = node.xpath('.//w:t', namespaces=node.nsmap) if hasattr(node, 'xpath') else []
                    text = ''.join([t.text for t in text_nodes if t is not None and t.text is not None])
                    txt = (text or '').strip()
                    norm = txt.upper()
                    if any(k in norm for k in ['EDUCATION', 'SKILLS', 'SUMMARY', 'PROJECT', 'CERTIFICATION', 'EXPERIENCE', 'WORK EXPERIENCE', 'EMPLOYMENT HISTORY']):
                        break
                    if txt.startswith(('•', '-', '–', '—', '*', '●')) or re.match(r'^\d+[\).\-\s]', txt):
                        bullets.append(txt.lstrip(' •–—-*●'))
                    else:
                        break
                node = node.getnext()
        except Exception:
            pass
        return bullets

    def _delete_next_table(self, paragraph):
        """Delete the immediate next table after a heading/placeholder (used when raw content is a table)."""
        try:
            node = paragraph._element.getnext()
            deleted = 0
            # Delete multiple tables if they exist
            while node is not None and node.tag.endswith('tbl'):
                next_node = node.getnext()
                parent = node.getparent()
                parent.remove(node)
                deleted += 1
                node = next_node
            if deleted > 0:
                print(f"    🗑️  Deleted {deleted} old table(s)")
        except Exception as e:
            print(f"    ⚠️  Error deleting tables: {e}")
    
    def _paragraph_in_table(self, paragraph):
        try:
            node = paragraph._element
            while node is not None:
                if str(getattr(node, 'tag', '')).endswith('tbl'):
                    return True
                node = node.getparent()
        except Exception:
            pass
        return False
    
    def _remove_instructional_until_table(self, paragraph, max_scan=40):
        try:
            node = paragraph._element.getnext()
            scanned = 0
            while node is not None and scanned < max_scan:
                scanned += 1
                if node.tag.endswith('tbl'):
                    break
                if node.tag.endswith('p'):
                    tnodes = node.xpath('.//w:t', namespaces=node.nsmap) if hasattr(node, 'xpath') else []
                    txt = ''.join([t.text for t in tnodes if t is not None and t.text is not None]).strip()
                    upper = txt.upper()
                    if any(h in upper for h in ['EMPLOYMENT', 'WORK HISTORY', 'PROFESSIONAL EXPERIENCE', 'WORK EXPERIENCE', 'CAREER HISTORY', 'EDUCATION', 'SUMMARY', 'CERTIFICATIONS', 'PROJECTS']) and len(txt) < 50:
                        break
                    is_instr = bool(re.search(r'\bplease\b', txt, re.IGNORECASE)) or \
                              bool(re.search(r'(use this table|add or delete rows|respond with the years|list the candidate|required/desired)', txt, re.IGNORECASE))
                    if is_instr or not txt:
                        parent = node.getparent()
                        parent.remove(node)
                        node = paragraph._element.getnext()
                        continue
                    else:
                        break
                node = node.getnext()
        except Exception:
            pass

    def _clear_instruction_phrases(self, doc):
        try:
            phrases = [
                'PLEASE USE THIS TABLE TO LIST THE SKILLS',
                'PLEASE LIST THE CANDIDATE’S RELEVANT EMPLOYMENT HISTORY',
                "PLEASE LIST THE CANDIDATE'S RELEVANT EMPLOYMENT HISTORY",
                'ADD OR DELETE ROWS AS NECESSARY'
            ]
            for p in doc.paragraphs:
                t = (p.text or '').strip().upper()
                if any(ph in t for ph in phrases):
                    for r in p.runs:
                        r.text = ''
        except Exception:
            pass
    
    def _build_experience_from_bullets(self, bullets):
        """Best-effort convert raw bullet lines into structured exp list when parser is empty."""
        exps = []
        i = 0
        while i < len(bullets):
            line = bullets[i]
            # Case: role + dates on this line
            if re.search(r'(?:19|20)\d{2}', line):
                duration = self._clean_duration(line)
                role = re.sub(r'\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-zA-Z]*\s+(?:19|20)\d{2}\b', '', line, flags=re.IGNORECASE)
                role = re.sub(r'\b(?:19|20)\d{2}\b', '', role)
                role = re.sub(r'\b(to|–|—|-)\b', '', role, flags=re.IGNORECASE).strip(' ,;:-')
                company = ''
                j = i + 1
                if i + 1 < len(bullets):
                    maybe_company = bullets[i+1].strip()
                    def _prob_company(text):
                        if not text:
                            return False
                        if len(text.split()) > 6:
                            return False
                        if text.endswith('.'):
                            return False
                        if re.search(r'\b(designed|managed|implemented|maintained|developed|led|created|configured|administered|tested|deployed|collaborated|supported|provided|responsible|oversaw|ownership|monitor|monitored|engineered)\b', text, re.IGNORECASE):
                            return False
                        if text[0].islower():
                            return False
                        return True
                    if _prob_company(maybe_company):
                        company = re.sub(r',[^,]*\b(?:city|state|india|usa|uk)\b.*$', '', maybe_company, flags=re.IGNORECASE).strip()
                        j = i + 2
                details = []
                while j < len(bullets):
                    if re.search(r'(?:19|20)\d{2}', bullets[j]):
                        break
                    details.append(bullets[j])
                    j += 1
                exps.append({'company': company, 'role': role, 'duration': duration, 'details': details})
                i = j
            else:
                i += 1
        return exps
    
    def _build_education_from_bullets(self, bullets):
        """Convert raw education bullets into degree/institution/year list when parser is empty."""
        edus = []
        i = 0
        while i < len(bullets):
            line = bullets[i]
            degree = ''
            institution = ''
            year = ''
            # Try to split by institution keyword
            m = re.search(r'(university|college|school|institute|academy)\b.*', line, flags=re.IGNORECASE)
            if m:
                degree = line[:m.start()].strip(' ,;:-')
                institution = line[m.start():].strip()
                year = self._clean_duration(line)
            else:
                # If next line is year, treat current as degree+institution
                if i + 1 < len(bullets) and re.search(r'(?:19|20)\d{2}', bullets[i+1]):
                    degree = line
                    year = self._clean_duration(bullets[i+1])
                    i += 1
                else:
                    degree = line
                    year = self._clean_duration(line)
            # Cleanup
            degree = degree.strip()
            institution = re.sub(r',[^,]*\b(?:city|state|india|usa|uk)\b.*$', '', institution, flags=re.IGNORECASE).strip()
            edus.append({'degree': degree, 'institution': institution, 'year': year, 'details': []})
            i += 1
        return edus

    def _optimize_details(self, details, max_bullets=12, max_words=22, max_chars=160):
        """Shorten and normalize bullet points while preserving meaning.
        - trims bullets, normalizes acronyms, removes duplicate entries
        - caps each bullet by words/chars
        """
        cleaned = []
        seen = set()
        for d in details:
            if not d or not isinstance(d, str):
                continue
            t = d.strip()
            if not t:
                continue
            # Remove leading bullet chars
            t = t.lstrip('•–—-*● \t-').strip()
            # Normalize common wording
            repl = [
                (r'\bas well as\b', 'and'),
                (r'\bin order to\b', 'to'),
                (r'\bkey performance indicators\s*\(([^\)]+)\)', r'\1'),
                (r'\bkey performance indicators\b', 'KPIs'),
                (r'\bquickbooks\b', 'QuickBooks'),
                (r'\bums?\s*worldship\b', 'UPS WorldShip'),
            ]
            for pattern, repl_to in repl:
                t = re.sub(pattern, repl_to, t, flags=re.IGNORECASE)

            t = self._shorten_text(t, max_words=max_words, max_chars=max_chars)
            t = self._normalize_acronyms(t)

            key = t.lower()
            if key and key not in seen:
                seen.add(key)
                cleaned.append(t.rstrip())

            if len(cleaned) >= max_bullets:
                break
        return cleaned

    def _normalize_acronyms(self, text):
        """Normalize common acronyms casing."""
        mapping = {
            'kpi': 'KPI', 'kpis': 'KPIs',
            'lms': 'LMS',
            'qa': 'QA',
        }
        def repl(m):
            return mapping.get(m.group(0).lower(), m.group(0))
        return re.sub(r'\b(kpis?|lms|qa)\b', repl, text, flags=re.IGNORECASE)

    def _shorten_text(self, text, max_words=22, max_chars=160):
        """Heuristic shortening: prefer cutting at clause boundaries, then word limit."""
        t = re.sub(r'\s+', ' ', text).strip()
        # Prefer to cut at clause markers if too long
        if len(t) > max_chars:
            for marker in ['; ', '. ', ' which ', ' that ', ' ensuring ', ' including ', ' while ', ' whereas ', ' whereby ']:
                idx = t.lower().find(marker)
                if 0 < idx <= max_chars:
                    t = t[:idx].rstrip('.; ,')
                    break
        # Enforce word cap
        words = t.split()
        if len(words) > max_words:
            t = ' '.join(words[:max_words]).rstrip(',;')
        # Ensure terminal period for readability
        if t and t[-1] not in '.!?':
            t = t + '.'
        return t
    
    def _create_replacement_map(self):
        """Create comprehensive replacement map"""
        replacements = {}
        
        # Personal information - Multiple formats
        # NOTE: Be specific to avoid replacing CAI contact manager info
        if self.resume_data.get('name'):
            display_name = f"<{self.resume_data['name']}>"
            replacements['[NAME]'] = display_name
            replacements['[CANDIDATE NAME]'] = display_name
            replacements['<CANDIDATE NAME>'] = display_name
            replacements["<Candidate's full name>"] = display_name
            replacements['<Candidate Name>'] = display_name
            replacements['<Name>'] = display_name
            replacements['Your Name'] = display_name
            replacements['CANDIDATE NAME'] = display_name
            # DO NOT replace "Insert name" as it might be in CAI contact section
        
        if self.resume_data.get('email'):
            # ONLY replace explicit placeholders, NOT actual email addresses
            replacements['[EMAIL]'] = self.resume_data['email']
            replacements['[Email]'] = self.resume_data['email']
            replacements['<EMAIL>'] = self.resume_data['email']
            replacements['<Email>'] = self.resume_data['email']
            replacements['<Candidate Email>'] = self.resume_data['email']
            # DO NOT replace example emails or "Email:" labels to avoid changing CAI contact info
        
        if self.resume_data.get('phone'):
            # ONLY replace explicit placeholders, NOT actual phone numbers
            replacements['[PHONE]'] = self.resume_data['phone']
            replacements['[Phone]'] = self.resume_data['phone']
            replacements['<PHONE>'] = self.resume_data['phone']
            replacements['<Phone>'] = self.resume_data['phone']
            replacements['<Candidate Phone>'] = self.resume_data['phone']
            # DO NOT replace example numbers or "Phone:" labels to avoid changing CAI contact info
        
        if self.resume_data.get('address'):
            replacements['[ADDRESS]'] = self.resume_data['address']
            replacements['[Address]'] = self.resume_data['address']
            replacements['<ADDRESS>'] = self.resume_data['address']
            replacements['<Address>'] = self.resume_data['address']
            replacements['Your Address'] = self.resume_data['address']
        
        if self.resume_data.get('linkedin'):
            replacements['[LINKEDIN]'] = self.resume_data['linkedin']
            replacements['[LinkedIn]'] = self.resume_data['linkedin']
            replacements['<LINKEDIN>'] = self.resume_data['linkedin']
            replacements['<LinkedIn>'] = self.resume_data['linkedin']
            replacements['linkedin.com/in/username'] = self.resume_data['linkedin']
        
        if self.resume_data.get('dob'):
            replacements['[DOB]'] = self.resume_data['dob']
            replacements['[Date of Birth]'] = self.resume_data['dob']
            replacements['<DOB>'] = self.resume_data['dob']
        
        return replacements
    
    def _text_contains(self, text, search_term):
        """Case-insensitive text search"""
        return search_term.lower() in text.lower()
    
    def _replace_in_paragraph(self, paragraph, search_term, replacement):
        """Replace text in paragraph while preserving formatting"""
        replaced = 0
        
        # First try: Replace in individual runs
        for run in paragraph.runs:
            if self._text_contains(run.text, search_term):
                # Case-insensitive replacement
                pattern = re.compile(re.escape(search_term), re.IGNORECASE)
                run.text = pattern.sub(replacement, run.text)
                replaced += 1
        
        # Second try: If not found in individual runs, text might be split
        # Combine all runs and check
        if replaced == 0 and self._text_contains(paragraph.text, search_term):
            # Text is split across runs - need to handle differently
            full_text = paragraph.text
            pattern = re.compile(re.escape(search_term), re.IGNORECASE)
            new_text = pattern.sub(replacement, full_text)
            
            if new_text != full_text:
                # Clear all runs and add new text
                for run in paragraph.runs:
                    run.text = ''
                
                # Add replacement text to first run
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                    replaced += 1
                else:
                    # No runs, add new run
                    paragraph.add_run(new_text)
                    replaced += 1
        
        return replaced

    def _regex_replace_paragraph(self, paragraph, pattern, replacement):
        """Regex-based replacement across runs: rebuilds paragraph text."""
        try:
            full_text = paragraph.text or ''
            new_text = re.sub(pattern, replacement, full_text, flags=re.IGNORECASE)
            if new_text != full_text:
                # clear runs and set new_text
                for run in paragraph.runs:
                    run.text = ''
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                else:
                    paragraph.add_run(new_text)
        except Exception:
            pass
    
    def _add_sections_content(self, doc):
        """Add resume sections to document and replace placeholders - SIMPLIFIED to prevent duplication"""
        sections_added = 0
        
        # Flags are initialized in _format_docx_file()
        # This method will check them to prevent duplicate insertion
        
        print(f"\n🔍 Scanning document for sections (SUMMARY, EXPERIENCE, EDUCATION, SKILLS)...")
        print(f"  📊 Section status: Summary={self._summary_inserted}, Experience={self._experience_inserted}, Education={self._education_inserted}")
        
        # SINGLE PASS: Look for headings only (ignore placeholders to avoid duplication)
        for para_idx, paragraph in enumerate(doc.paragraphs):
            para_text = paragraph.text.upper().strip()
            
            # SUMMARY SECTION
            if not self._summary_inserted and any(marker in para_text for marker in ['SUMMARY', 'OBJECTIVE', 'PROFILE', 'PROFESSIONAL SUMMARY']):
                summary = (self.resume_data.get('summary') or '').strip()
                summary_lines = self._find_matching_resume_section('summary', self.resume_data.get('sections', {}))
                if summary or summary_lines:
                    print(f"  ✓ Found SUMMARY at paragraph {para_idx}: '{paragraph.text[:50]}'")
                    
                    # Clear the heading paragraph (keep only the heading text)
                    for run in paragraph.runs:
                        run.text = ''
                    if paragraph.runs:
                        paragraph.runs[0].text = 'SUMMARY'
                        paragraph.runs[0].bold = True
                        paragraph.runs[0].font.size = Pt(12)
                    
                    # Delete any following content before inserting new
                    self._delete_following_bullets(paragraph, max_scan=20)
                    
                    # Prefer bullet-style summary if multiple lines available
                    if summary_lines:
                        self._insert_skills_bullets(doc, paragraph, summary_lines)
                    else:
                        summary_para = self._insert_paragraph_after(paragraph, summary)
                        if summary_para:
                            for run in summary_para.runs:
                                run.font.size = Pt(10)
                            summary_para.paragraph_format.space_after = Pt(10)
                    
                    self._summary_inserted = True
                    sections_added += 1
                    print(f"    → Inserted summary ({len(summary_lines) if summary_lines else len(summary)} items)")
                    continue
                else:
                    # No summary data: remove heading and any following content until next section
                    print(f"  ⚠️  SUMMARY heading found but no data; removing section")
                    self._delete_following_bullets(paragraph, max_scan=80)
                    self._delete_next_table(paragraph)
                    self._delete_paragraph(paragraph)
                    continue
            
            # EXPERIENCE SECTION - Check if it hasn't been inserted yet
            if (not self._experience_inserted \
                and any(marker in para_text for marker in ['EMPLOYMENT HISTORY', 'WORK EXPERIENCE', 'PROFESSIONAL EXPERIENCE', 'EXPERIENCE', 'WORK HISTORY', 'CAREER HISTORY']) \
                and len(paragraph.text.strip()) < 50 \
                and not self._paragraph_in_table(paragraph)):
                experiences = self.resume_data.get('experience', [])
                # Fallback: build structured experiences from the raw bullets beneath the heading
                if not experiences:
                    raw_bullets = self._collect_bullets_after_heading(paragraph, max_scan=120)
                    if raw_bullets:
                        experiences = self._build_experience_from_bullets(raw_bullets)
                        print(f"  🔄 Built {len(experiences)} experience entries from raw bullets")
                if experiences:
                    print(f"  ✓ Found EXPERIENCE at paragraph {para_idx}: '{paragraph.text[:50]}'")
                    
                    # STEP 1: Clear the heading paragraph (keep only the heading text)
                    original_heading = paragraph.text.strip()
                    for run in paragraph.runs:
                        run.text = ''
                    if paragraph.runs:
                        paragraph.runs[0].text = 'EMPLOYMENT HISTORY'
                        paragraph.runs[0].bold = True
                        paragraph.runs[0].font.size = Pt(12)
                    
                    # STEP 2: Delete old template content only if experiences were built from fallback
                    # (Don't delete if we have structured experience from resume parser)
                    if not self.resume_data.get('experience'):
                        # Only clear template placeholders/raw bullets when using fallback
                        self._delete_following_bullets(paragraph, max_scan=800)
                        self._delete_next_table(paragraph)
                    
                    # STEP 3: Insert clean structured blocks
                    last_element = paragraph
                    for exp in experiences[:10]:
                        table = self._insert_experience_block(doc, last_element, exp)
                        if table:
                            last_element = table
                    
                    # STEP 4: Skip aggressive cleanup to preserve newly inserted bullets
                    # self._cleanup_duplicate_bullets_after_section(doc, paragraph, 'EDUCATION')
                    
                    self._experience_inserted = True
                    sections_added += 1
                    print(f"    → Inserted {len(experiences[:10])} experience entries")
                    continue
            
            # SKILLS SECTION
            if not self._skills_inserted and any(marker in para_text for marker in ['SKILLS', 'TECHNICAL SKILLS', 'COMPETENCIES', 'EXPERTISE']):
                skills = self.resume_data.get('skills', [])
                if skills and len(skills) > 0:
                    print(f"  ✓ Found SKILLS at paragraph {para_idx}: '{paragraph.text[:50]}'")
                    
                    # Ensure heading formatting
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(12)

                    # Determine if a table follows this heading
                    node = paragraph._element.getnext()
                    has_table_next = bool(node is not None and str(getattr(node, 'tag', '')).endswith('tbl'))

                    if not has_table_next:
                        # No table available → insert bullets
                        print("    → No skills table after heading; inserting bullet list")
                        # Remove any raw content under heading first
                        self._delete_following_bullets(paragraph, max_scan=50)
                        self._insert_skills_bullets(doc, paragraph, skills)
                    else:
                        print("    → Skills table detected; will fill during table pass")
                        # Remove instructional text between heading and table
                        self._remove_instructional_until_table(paragraph, max_scan=20)
                    
                    self._skills_inserted = True
                    sections_added += 1
                    continue
                else:
                    # No skills: remove section heading and trailing content
                    print(f"  ⚠️  SKILLS heading found but resume has no skills; removing section")
                    self._delete_following_bullets(paragraph, max_scan=50)
                    self._delete_next_table(paragraph)
                    self._delete_paragraph(paragraph)
                    continue
            
            # EDUCATION SECTION - Check multiple variations and handle fallback
            if not self._education_inserted and any(marker in para_text for marker in ['EDUCATION', 'ACADEMIC BACKGROUND', 'EDUCATIONAL BACKGROUND', 'ACADEMIC QUALIFICATIONS', 'QUALIFICATIONS', 'EDUCATION BACKGROUND', 'ACADEMICS']):
                education = self.resume_data.get('education', [])
                # Fallback priority 1: derive from resume sections (not from template content)
                if not education:
                    section_lines = self._find_matching_resume_section('education', self.resume_data.get('sections', {}))
                    if section_lines:
                        education = self._build_education_from_bullets(section_lines)
                        print(f"  🔄 Built {len(education)} education entries from resume sections")
                # Fallback priority 2: derive from any raw bullets directly under the heading in template (rare)
                if not education:
                    raw_bullets = self._collect_bullets_after_heading(paragraph, max_scan=80)
                    if raw_bullets:
                        education = self._build_education_from_bullets(raw_bullets)
                        print(f"  🔄 Built {len(education)} education entries from raw bullets under heading")
                if education:
                    print(f"  ✓ Found EDUCATION at paragraph {para_idx}: '{paragraph.text[:50]}'")
                    print(f"  📚 Have {len(education)} education entries to insert")
                    
                    # DEBUG: Show education data
                    for i, edu in enumerate(education[:3]):
                        print(f"      {i+1}. Degree: '{edu.get('degree', '')[:40]}', Institution: '{edu.get('institution', '')[:30]}', Year: '{edu.get('year', '')}'")
                    
                    # STEP 1: Clear the heading paragraph (keep only the heading text)
                    for run in paragraph.runs:
                        run.text = ''
                    if paragraph.runs:
                        paragraph.runs[0].text = 'EDUCATION'
                        paragraph.runs[0].bold = True
                        paragraph.runs[0].font.size = Pt(12)
                    
                    # STEP 2: Delete ALL following content (tables + paragraphs)
                    self._delete_following_bullets(paragraph)
                    self._delete_next_table(paragraph)
                    
                    # STEP 3: Insert clean structured blocks (or bullets for simple entries)
                    simple_entries = [e for e in education if not (e.get('institution') or (e.get('details') or []))]
                    inserted_count = 0
                    if simple_entries and len(simple_entries) == len(education):
                        self._insert_education_bullets(doc, paragraph, education[:5])
                        inserted_count = min(5, len(education))
                    else:
                        last_element = paragraph
                        for edu in education[:5]:
                            block = self._insert_education_block(doc, last_element, edu)
                            if block:
                                last_element = block
                                inserted_count += 1
                    
                    # STEP 4: Skip aggressive cleanup to preserve newly inserted bullets
                    # self._cleanup_duplicate_bullets_after_section(doc, paragraph, 'SKILLS')
                    
                    self._education_inserted = True
                    sections_added += 1
                    print(f"    → Successfully inserted {inserted_count} education entries")
                    continue
        
        print(f"\n✅ Section insertion complete. Summary: {self._summary_inserted}, Experience: {self._experience_inserted}, Education: {self._education_inserted}")
        return sections_added
    
    def _detect_table_type(self, table):
        """
        DYNAMICALLY detect table type by analyzing column headers.
        Returns: 'skills', 'experience', 'education', or None
        """
        if len(table.rows) < 1:
            return None
        
        # Get all text from first few rows (headers might span multiple rows)
        header_text = ''
        for row_idx in range(min(3, len(table.rows))):
            for cell in table.rows[row_idx].cells:
                header_text += ' ' + cell.text.lower()
        
        # Skills table indicators
        skills_indicators = ['skill', 'technology', 'competency', 'expertise', 'proficiency', 
                           'years used', 'last used', 'technical']
        
        # Experience/Employment table indicators  
        experience_indicators = ['employment', 'company', 'employer', 'position', 'role', 
                                'job title', 'work history', 'experience', 'responsibilities']
        
        # Education table indicators
        education_indicators = ['education', 'degree', 'institution', 'university', 'college', 
                              'school', 'graduation', 'qualification']
        
        # Count matches for each type
        skills_score = sum(1 for ind in skills_indicators if ind in header_text)
        exp_score = sum(1 for ind in experience_indicators if ind in header_text)
        edu_score = sum(1 for ind in education_indicators if ind in header_text)
        
        # Return type with highest score
        if skills_score > 0 and skills_score >= exp_score and skills_score >= edu_score:
            return 'skills'
        elif exp_score > 0 and exp_score >= edu_score:
            return 'experience'
        elif edu_score > 0:
            return 'education'
        
        return None
    
    def _fill_dynamic_table(self, table, table_type):
        """
        DYNAMICALLY fill table based on column headers and data type.
        Works with ANY column structure!
        """
        if len(table.rows) < 1:
            return 0
        
        print(f"     🔍 Analyzing table structure...")
        
        # Get column headers (from first row)
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        print(f"     📋 Column headers: {headers}")
        
        # Map columns to data fields intelligently
        column_mapping = self._map_columns_to_fields(headers, table_type)
        print(f"     🗺️  Column mapping: {column_mapping}")
        
        # Get data to fill
        if table_type == 'skills':
            data_items = self._get_skills_data()
        elif table_type == 'experience':
            data_items = self._get_experience_data()
        elif table_type == 'education':
            data_items = self._get_education_data()
        else:
            return 0
        
        if not data_items:
            print(f"     ⚠️  No {table_type} data available")
            return 0
        
        print(f"     📊 Found {len(data_items)} {table_type} items to fill")
        
        # Clear existing rows (keep header)
        for i in reversed(range(1, len(table.rows))):
            table._element.remove(table.rows[i]._element)
        
        # Fill rows dynamically
        filled = 0
        for item in data_items[:15]:  # Limit to 15 rows
            new_row = table.add_row()
            
            for col_idx, field_name in column_mapping.items():
                if col_idx < len(new_row.cells) and field_name:
                    value = item.get(field_name, '')
                    new_row.cells[col_idx].text = str(value)
            
            filled += 1
            if filled <= 3:
                print(f"        ✓ Row {filled}: {list(item.values())[:3]}")
        
        return filled
    
    def _map_columns_to_fields(self, headers, table_type):
        """
        INTELLIGENT column mapping - maps table columns to resume data fields
        based on semantic understanding of column names.
        """
        mapping = {}  # {column_index: field_name}
        
        if table_type == 'skills':
            # Define possible column patterns for skills
            patterns = {
                'skill': ['skill', 'technology', 'competency', 'expertise', 'tool', 'name'],
                'years': ['year', 'experience', 'exp', 'yrs', 'duration', 'used'],
                'last_used': ['last', 'recent', 'current', 'latest', 'when']
            }
        elif table_type == 'experience':
            patterns = {
                'company': ['company', 'employer', 'organization', 'firm'],
                'role': ['role', 'position', 'title', 'job'],
                'duration': ['date', 'year', 'period', 'duration', 'from', 'to', 'when'],
                'location': ['location', 'city', 'state', 'place'],
                'responsibilities': ['responsibilit', 'duties', 'description', 'summary']
            }
        elif table_type == 'education':
            patterns = {
                'degree': ['degree', 'qualification', 'certificate', 'program'],
                'institution': ['institution', 'university', 'college', 'school'],
                'year': ['year', 'date', 'graduation', 'completion'],
                'field': ['field', 'major', 'specialization', 'subject'],
                'gpa': ['gpa', 'grade', 'marks', 'score']
            }
        else:
            return mapping
        
        # Map each column to best matching field
        for col_idx, header in enumerate(headers):
            header_lower = header.lower()
            best_match = None
            best_score = 0
            
            for field_name, keywords in patterns.items():
                score = sum(1 for kw in keywords if kw in header_lower)
                if score > best_score:
                    best_score = score
                    best_match = field_name
            
            if best_match and best_score > 0:
                mapping[col_idx] = best_match
                print(f"        Column {col_idx} ('{header}') → {best_match}")
        
        return mapping
    
    def _get_skills_data(self):
        """Get skills data in standardized format"""
        skills_list = []
        raw_skills = self.resume_data.get('skills', [])
        
        for skill in raw_skills[:15]:
            skill_name = skill if isinstance(skill, str) else skill.get('name', '')
            skills_list.append({
                'skill': skill_name,
                'years': '2+ years',  # Default
                'last_used': 'Recent'
            })
        
        return skills_list
    
    def _get_experience_data(self):
        """Get experience data in standardized format"""
        exp_list = []
        experiences = self.resume_data.get('experience', [])
        
        for exp in experiences[:10]:
            exp_list.append({
                'company': exp.get('company', ''),
                'role': exp.get('role', ''),
                'duration': exp.get('duration', ''),
                'location': '',  # Extract if available
                'responsibilities': '\n'.join(exp.get('details', [])[:3])
            })
        
        return exp_list
    
    def _get_education_data(self):
        """Get education data in standardized format"""
        edu_list = []
        education = self.resume_data.get('education', [])
        
        for edu in education[:5]:
            edu_list.append({
                'degree': edu.get('degree', ''),
                'institution': edu.get('institution', ''),
                'year': edu.get('year', ''),
                'field': '',  # Extract from degree if available
                'gpa': ''
            })
        
        return edu_list
    
    def _is_skills_table(self, table):
        """Check if table is a skills table by examining headers - FLEXIBLE detection"""
        if len(table.rows) < 1:  # Changed from 2 to 1 - just need header row
            print(f"       ⚠️  Table has <1 rows ({len(table.rows)}), skipping")
            return False
        
        # Get first row (header) text - check multiple rows in case header spans multiple
        header_texts = []
        rows_to_check = min(3, len(table.rows))  # Check first 3 rows for headers
        
        for row_idx in range(rows_to_check):
            row_texts = [cell.text.strip().lower() for cell in table.rows[row_idx].cells]
            # Skip completely empty rows
            if any(t for t in row_texts):
                header_texts.extend(row_texts)
        
        # Join all potential headers
        all_headers = ' '.join(header_texts)
        
        print(f"       🔍 Table has {len(table.rows)} rows, {len(table.columns)} columns")
        print(f"       🔍 First row cells: {[cell.text.strip() for cell in table.rows[0].cells]}")
        print(f"       🔍 All header candidates: {header_texts[:6]}")  # Show first 6
        print(f"       🔍 Combined text: '{all_headers[:100]}'")  # First 100 chars
        
        # Check for skills table indicators - VERY FLEXIBLE
        skills_keywords = ['skill', 'skills', 'technology', 'technologies', 'competency', 'competencies', 
                          'technical', 'proficiency', 'expertise', 'tool', 'tools', 'qualification']
        years_keywords = ['years', 'experience', 'years used', 'years of experience', 'exp', 'yrs', 
                         'year', 'duration']
        last_used_keywords = ['last used', 'last', 'recent', 'most recent', 'latest', 'when', 'current']
        
        has_skill_col = any(kw in all_headers for kw in skills_keywords)
        has_years_col = any(kw in all_headers for kw in years_keywords)
        has_last_used_col = any(kw in all_headers for kw in last_used_keywords)
        
        # Also check if table has exactly 3 columns (Skill, Years, Last Used pattern)
        has_three_cols = len(table.columns) == 3
        
        print(f"       📊 Detection results:")
        print(f"          - Has 3 columns: {has_three_cols} (actual: {len(table.columns)})")
        print(f"          - Has skill column: {has_skill_col}")
        print(f"          - Has years column: {has_years_col}")
        print(f"          - Has last_used column: {has_last_used_col}")
        
        # It's a skills table if:
        # 1. Has skill keyword AND (years OR last_used keyword)
        # 2. OR has 3 columns with years AND last_used (common pattern)
        is_skills = (has_skill_col and (has_years_col or has_last_used_col)) or \
                    (has_three_cols and has_years_col and has_last_used_col)
        
        print(f"       {'✅ IS SKILLS TABLE' if is_skills else '❌ NOT SKILLS TABLE'}")
        
        return is_skills
    
    def _fill_skills_table(self, table):
        """Fill skills table with candidate's skills data"""
        if len(table.rows) < 1:
            print(f"     ⚠️  Table has no rows, cannot fill")
            return 0
        
        # Get header row to identify columns
        header_row = table.rows[0]
        header_texts = [cell.text.strip().lower() for cell in header_row.cells]
        
        print(f"     📋 Filling skills table...")
        print(f"     📋 Table headers: {header_texts}")
        print(f"     📋 Table has {len(table.rows)} rows initially")
        
        # Find column indices - FLEXIBLE matching
        skill_col = None
        years_col = None
        last_used_col = None
        
        skill_keywords = ['skill', 'technology', 'competency', 'technical', 'tool', 'expertise', 'proficiency']
        years_keywords = ['years', 'experience', 'exp', 'yrs', 'years used']
        last_keywords = ['last', 'recent', 'latest', 'last used', 'most recent']
        
        for idx, header in enumerate(header_texts):
            # Match skill column
            if any(kw in header for kw in skill_keywords) and skill_col is None:
                skill_col = idx
                print(f"     ✓ Skill column: {idx} ('{header}')")
            # Match years column
            elif any(kw in header for kw in years_keywords) and years_col is None:
                years_col = idx
                print(f"     ✓ Years column: {idx} ('{header}')")
            # Match last used column
            elif any(kw in header for kw in last_keywords) and last_used_col is None:
                last_used_col = idx
                print(f"     ✓ Last Used column: {idx} ('{header}')")
        
        if skill_col is None:
            print(f"     ⚠️  No skill column found in headers: {header_texts}")
            return 0
        
        # Get skills from resume
        skills_data = self._extract_skills_with_details()
        
        print(f"     📊 Extracted {len(skills_data) if skills_data else 0} skills from resume")
        
        if not skills_data or len(skills_data) == 0:
            print(f"     ⚠️  No skills data to fill!")
            # Get raw skills list as fallback
            raw_skills = self.resume_data.get('skills', [])
            print(f"     ℹ️  Raw skills available: {len(raw_skills)}")
            if raw_skills:
                # Convert raw skills to format expected
                skills_data = [{'skill': s, 'years': '1+ years', 'last_used': 'Recent'} for s in raw_skills[:15]]
                print(f"     ✓ Using {len(skills_data)} raw skills as fallback")
            else:
                return 0
        
        # Clear existing data rows (keep header)
        rows_to_delete = []
        for i in range(1, len(table.rows)):
            rows_to_delete.append(i)
        
        # Delete from bottom to top to avoid index issues
        for i in reversed(rows_to_delete):
            table._element.remove(table.rows[i]._element)
        
        # Add skills rows
        filled_count = 0
        print(f"     🔄 Adding {min(15, len(skills_data))} skill rows to table...")
        
        for skill_info in skills_data[:15]:  # Limit to 15 skills
            # Add new row
            new_row = table.add_row()
            
            skill_name = skill_info.get('skill', '')
            
            # Fill skill name
            if skill_col is not None:
                new_row.cells[skill_col].text = skill_name
            
            # Fill years
            if years_col is not None:
                new_row.cells[years_col].text = skill_info.get('years', '')
            
            # Fill last used
            if last_used_col is not None:
                new_row.cells[last_used_col].text = skill_info.get('last_used', '')
            
            filled_count += 1
            if filled_count <= 3:
                print(f"        ✓ Added: {skill_name}")
        
        print(f"     ✅ Successfully filled {filled_count} skill rows")
        return filled_count
    
    def _extract_skills_with_details(self):
        """Extract skills with years and last used info from resume data"""
        skills_list = []
        
        # Get skills from resume data
        skills = self.resume_data.get('skills', [])
        experience = self.resume_data.get('experience', [])
        
        # Try to extract years from experience
        current_year = 2025
        
        for skill in skills[:15]:  # Limit to 15 skills
            skill_name = skill if isinstance(skill, str) else skill.get('name', '')
            
            # Try to find this skill in experience to get dates
            years_exp = ''
            last_used = ''
            
            # Search through experience for this skill
            for exp in experience:
                exp_text = str(exp).lower()
                if skill_name.lower() in exp_text:
                    # Try to extract years
                    duration = exp.get('duration', '') if isinstance(exp, dict) else ''
                    
                    # Parse years from duration like "2020-2023" or "2020-Present"
                    import re
                    year_matches = re.findall(r'(20\d{2})', str(duration))
                    if year_matches:
                        start_year = int(year_matches[0])
                        end_year = int(year_matches[-1]) if len(year_matches) > 1 else current_year
                        
                        if 'present' in str(duration).lower() or 'current' in str(duration).lower():
                            end_year = current_year
                        
                        years_count = end_year - start_year
                        if years_count > 0:
                            years_exp = f"{years_count}+ years"
                            last_used = str(end_year) if end_year < current_year else "Present"
                    
                    break
            
            # Default values if not found in experience
            if not years_exp:
                years_exp = "1+ years"
            if not last_used:
                last_used = "Recent"
            
            skills_list.append({
                'skill': skill_name,
                'years': years_exp,
                'last_used': last_used
            })
        
        return skills_list
    
    def _find_matching_resume_section(self, section_key, resume_sections):
        """Find matching resume section with synonyms"""
        # Direct match
        if section_key in resume_sections:
            return resume_sections[section_key]

        synonyms = {
            'experience': ['experience', 'employment', 'work', 'professional'],
            'education': ['education', 'academic', 'qualification', 'academics'],
            'skills': ['skills', 'technical', 'competencies', 'expertise'],
            'summary': ['summary', 'objective', 'profile', 'about'],
            'projects': ['projects', 'portfolio'],
            'certifications': ['certifications', 'certificates', 'licenses'],
            'awards': ['awards', 'achievements', 'honors']
        }

        patterns = synonyms.get(section_key, [section_key])
        for resume_key, content in resume_sections.items():
            key_lower = resume_key.lower()
            if any(p in key_lower for p in patterns):
                return content

        return []
    
    def _convert_to_pdf(self, docx_path, pdf_path):
        """Convert DOCX to PDF"""
        try:
            if HAS_WIN32:
                # Use Word COM to convert
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                
                doc = word.Documents.Open(os.path.abspath(docx_path))
                doc.SaveAs2(os.path.abspath(pdf_path), FileFormat=17)  # 17 = PDF format
                doc.Close()
                word.Quit()
                
                return True
            else:
                print("⚠️  PDF conversion requires Microsoft Word")
                return False
                
        except Exception as e:
            print(f"❌ PDF conversion error: {e}")
            return False


def format_word_document(resume_data, template_analysis, output_path):
    """Main function for Word document formatting"""
    formatter = WordFormatter(resume_data, template_analysis, output_path)
    return formatter.format()
