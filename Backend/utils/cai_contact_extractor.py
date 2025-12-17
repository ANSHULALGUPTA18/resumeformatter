"""
CAI Contact Extractor
Automatically detects CAI contact information from templates
"""

import re
from docx import Document
from typing import Dict, Any, Optional, List


def extract_cai_contact_from_template(template_path: str) -> Optional[Dict[str, Any]]:
    """
    Extract CAI contact information from a template DOCX file

    Looks for patterns like:
    - Name: John Doe
    - Phone: (123) 456-7890
    - Email: john@example.com
    - State: California or CA

    Args:
        template_path: Path to the template DOCX file

    Returns:
        Dict with {name, phone, email, state} or None if not found
    """
    try:
        doc = Document(template_path)

        # Extract all text from document
        full_text = ""
        for paragraph in doc.paragraphs:
            full_text += paragraph.text + "\n"

        # Also check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text += cell.text + "\n"

        print(f"📄 Analyzing template for CAI contact...")

        # Extract contact information
        contact_info = {
            "name": extract_name(full_text),
            "phone": extract_phone(full_text),
            "email": extract_email(full_text),
            "state": extract_state(full_text)
        }

        # Check if we found at least some information
        has_info = any(contact_info.values())

        if has_info:
            print(f"✅ CAI Contact detected:")
            print(f"   Name: {contact_info['name']}")
            print(f"   Phone: {contact_info['phone']}")
            print(f"   Email: {contact_info['email']}")
            print(f"   State: {contact_info['state']}")
            return contact_info
        else:
            print(f"⚠️ No CAI contact information found in template")
            return None

    except Exception as e:
        print(f"❌ Error extracting CAI contact: {e}")
        return None


def extract_name(text: str) -> str:
    """Extract CAI contact name from text"""
    patterns = [
        r'(?:CAI\s+)?(?:Contact\s+)?Name\s*[:\-]\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)',
        r'Contact\s*[:\-]\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)',
        r'CAI\s+Representative\s*[:\-]\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)',
        r'Representative\s*[:\-]\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)',
    ]

    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
        if match:
            name = match.group(1).strip()
            # Validate it's a real name (at least 2 words)
            if len(name.split()) >= 2:
                return name

    return ""


def extract_phone(text: str) -> str:
    """Extract phone number from text"""
    patterns = [
        r'(?:Phone|Tel|Telephone|Cell|Mobile)\s*[:\-]?\s*(\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4})',
        r'(\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4})',  # Generic phone pattern
    ]

    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            phone = match.group(1).strip()
            # Clean up phone number
            phone = re.sub(r'[^\d]', '', phone)  # Remove non-digits
            if len(phone) == 10:  # Valid US phone
                # Format as (XXX) XXX-XXXX
                return f"({phone[:3]}) {phone[3:6]}-{phone[6:]}"

    return ""


def extract_email(text: str) -> str:
    """Extract email address from text"""
    pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'

    match = re.search(pattern, text)
    if match:
        return match.group(0).strip()

    return ""


def extract_state(text: str) -> str:
    """Extract state information from text"""
    # Common US state names and abbreviations
    states = {
        # Full names
        'alabama': 'AL', 'alaska': 'AK', 'arizona': 'AZ', 'arkansas': 'AR',
        'california': 'CA', 'colorado': 'CO', 'connecticut': 'CT', 'delaware': 'DE',
        'florida': 'FL', 'georgia': 'GA', 'hawaii': 'HI', 'idaho': 'ID',
        'illinois': 'IL', 'indiana': 'IN', 'iowa': 'IA', 'kansas': 'KS',
        'kentucky': 'KY', 'louisiana': 'LA', 'maine': 'ME', 'maryland': 'MD',
        'massachusetts': 'MA', 'michigan': 'MI', 'minnesota': 'MN', 'mississippi': 'MS',
        'missouri': 'MO', 'montana': 'MT', 'nebraska': 'NE', 'nevada': 'NV',
        'new hampshire': 'NH', 'new jersey': 'NJ', 'new mexico': 'NM', 'new york': 'NY',
        'north carolina': 'NC', 'north dakota': 'ND', 'ohio': 'OH', 'oklahoma': 'OK',
        'oregon': 'OR', 'pennsylvania': 'PA', 'rhode island': 'RI', 'south carolina': 'SC',
        'south dakota': 'SD', 'tennessee': 'TN', 'texas': 'TX', 'utah': 'UT',
        'vermont': 'VT', 'virginia': 'VA', 'washington': 'WA', 'west virginia': 'WV',
        'wisconsin': 'WI', 'wyoming': 'WY'
    }

    # Abbreviations (reverse mapping)
    state_abbrevs = {v: v for v in states.values()}

    # Patterns to find state
    patterns = [
        r'State\s*[:\-]\s*([A-Za-z\s]+)',
        r'(?:State\s+of|for)\s+([A-Za-z\s]+)',
        r'\b(Georgia|Florida|Texas|California|New York|Virginia|Indiana|Idaho|Connecticut|North Dakota|Arkansas)\b',  # Common states in templates
        r'\b([A-Z]{2})\b',  # Two-letter state code
    ]

    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            state_text = match.group(1).strip().lower()

            # Check if it's a full state name
            if state_text in states:
                return states[state_text]

            # Check if it's already an abbreviation
            if state_text.upper() in state_abbrevs:
                return state_text.upper()

            # Check multi-word states
            for state_name, abbrev in states.items():
                if state_name in state_text:
                    return abbrev

    return ""


def extract_multiple_cai_contacts(template_path: str) -> List[Dict[str, Any]]:
    """
    Extract multiple CAI contacts from a template if it contains several

    Args:
        template_path: Path to the template DOCX file

    Returns:
        List of contact dictionaries
    """
    # For now, return single contact as a list
    # Can be extended to detect multiple contacts if needed
    contact = extract_cai_contact_from_template(template_path)
    if contact and any(contact.values()):
        return [contact]
    return []
